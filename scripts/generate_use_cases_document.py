from __future__ import annotations

import shutil
from pathlib import Path

import openpyxl
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MATRIX_FILE = ROOT / "artifacts" / "matrices" / "Matriz_Trazabilidad_Detallada_Tesis_Calzatura_Vilchez_CORREGIDA.xlsx"
OUT_DIR = ROOT / "artifacts" / "documentos"
OUTPUT_FILE = OUT_DIR / "Documento_3_Casos_de_Uso_Calzatura_Vilchez_COMPLETO.docx"
DOWNLOAD_FILE = Path.home() / "Downloads" / OUTPUT_FILE.name

PROJECT_TITLE = (
    "Sistema web de comercio electrónico con modelo de inteligencia artificial "
    "para la predicción del riesgo empresarial en la empresa Calzatura Vilchez"
)
AUTHOR = "Serpa Sedano Yeferson Wilson"
DOC_DATE = "30/05/2026"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.color.rgb = RGBColor(31, 78, 121)


def add_title(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("DOCUMENTO 3 — CASOS DE USO")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(PROJECT_TITLE).italic = True

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Autor: {AUTHOR} | Fecha: {DOC_DATE}").bold = True
    document.add_paragraph()


def add_table(document: Document, headers: list[str], rows: list[list[str]]):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0]
    for index, text in enumerate(headers):
        cell = header.cells[index]
        cell.text = text
        set_cell_shading(cell, "1F4E79")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(8)
    for row in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            cells[index].text = str(text)
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    document.add_paragraph()
    return table


def load_use_cases() -> list[dict[str, str]]:
    workbook = openpyxl.load_workbook(MATRIX_FILE, data_only=True)
    sheet = workbook["Casos_Uso"]
    cases: list[dict[str, str]] = []
    for row in sheet.iter_rows(min_row=2, values_only=True):
        if not row[0]:
            continue
        cases.append(
            {
                "codigo": str(row[0] or ""),
                "nombre": str(row[1] or ""),
                "actor": str(row[2] or ""),
                "requisitos": str(row[3] or ""),
                "precondicion": str(row[4] or ""),
                "flujo": str(row[5] or ""),
                "alternativo": str(row[6] or ""),
                "postcondicion": str(row[7] or ""),
                "evidencia": str(row[8] or ""),
            }
        )
    return cases


def add_case_detail(document: Document, case: dict[str, str]) -> None:
    document.add_heading(f"{case['codigo']} — {case['nombre']}", level=2)
    add_table(
        document,
        ["Elemento", "Descripción"],
        [
            ["Código del caso de uso", case["codigo"]],
            ["Nombre", case["nombre"]],
            ["Actor", case["actor"]],
            ["Requisito relacionado", case["requisitos"]],
            ["Descripción", case["nombre"]],
            ["Precondición", case["precondicion"]],
            ["Flujo principal", case["flujo"]],
            ["Flujo alternativo", case["alternativo"]],
            ["Postcondición", case["postcondicion"]],
            ["Evidencia esperada", case["evidencia"]],
        ],
    )


def main() -> None:
    if not MATRIX_FILE.exists():
        raise FileNotFoundError(MATRIX_FILE)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    cases = load_use_cases()

    document = Document()
    configure_document(document)
    add_title(document)

    document.add_heading("1. Propósito del documento", level=1)
    document.add_paragraph(
        "Este documento describe cómo interactúan los usuarios con el sistema web Calzatura Vilchez. "
        "Cada caso de uso está trazado a un requisito de la matriz corregida, con actor, precondición, flujo principal, "
        "flujo alternativo, postcondición y evidencia técnica esperada."
    )

    document.add_heading("2. Alcance", level=1)
    document.add_paragraph(
        "El alcance corresponde únicamente al sistema web comercial y sus servicios de soporte: tienda en línea, catálogo, "
        "carrito, checkout, pedidos, clientes, administración, ventas físicas, Excel, seguridad, auditoría, base de datos, IA, "
        "calidad, despliegue y continuidad."
    )

    document.add_heading("3. Estructura usada para cada caso de uso", level=1)
    add_table(
        document,
        ["Elemento", "Descripción"],
        [
            ["Código del caso de uso", "Identificador único: CU-001, CU-002, CU-003, etc."],
            ["Nombre", "Acción o proceso que representa el caso de uso."],
            ["Actor", "Usuario o componente que ejecuta o inicia la interacción."],
            ["Descripción", "Qué hace el caso de uso dentro del sistema."],
            ["Flujo principal", "Pasos correctos esperados del proceso."],
            ["Flujo alternativo", "Qué ocurre si existe error, datos inválidos o falta de autorización."],
            ["Precondición", "Condición previa necesaria para iniciar el caso de uso."],
            ["Postcondición", "Resultado final esperado después de ejecutar el caso de uso."],
            ["Evidencia", "Código, prueba, endpoint o artifact que confirma la implementación."],
        ],
    )

    document.add_heading("4. Resumen de casos de uso", level=1)
    add_table(
        document,
        ["Código", "Caso de uso", "Actor", "Requisito", "Evidencia"],
        [[case["codigo"], case["nombre"], case["actor"], case["requisitos"], case["evidencia"]] for case in cases],
    )

    document.add_heading("5. Detalle de casos de uso", level=1)
    for case in cases:
        add_case_detail(document, case)

    document.add_heading("6. Cierre de trazabilidad", level=1)
    document.add_paragraph(
        f"Total de casos de uso documentados: {len(cases)}. Todos provienen de la matriz de trazabilidad corregida y se encuentran "
        "asociados a requisitos implementados, pruebas ejecutadas y evidencia verificable en el repositorio o artifacts del proyecto."
    )

    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(f"Documento de casos de uso — Calzatura Vilchez — {DOC_DATE}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)

    document.save(OUTPUT_FILE)
    shutil.copy2(OUTPUT_FILE, DOWNLOAD_FILE)
    print(OUTPUT_FILE)
    print(DOWNLOAD_FILE)


if __name__ == "__main__":
    main()
