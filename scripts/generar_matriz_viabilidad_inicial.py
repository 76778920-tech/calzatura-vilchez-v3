#!/usr/bin/env python3
"""
Genera Matriz de Viabilidad Inicial (TELOS + Estratégica + Científica)
para Calzatura Vilchez. Fuente: CU-T13 + matriz-viabilidad-inicial-proyecto.md

Metodología: marco TELOS (Technical, Economic, Legal, Operational, Schedule)
ampliado con viabilidad estratégica de mercado y viabilidad científica-investigación,
respaldado por artículos Q1 de alto impacto (JCR/Scopus).

Uso: python scripts/generar_matriz_viabilidad_inicial.py
"""
from __future__ import annotations

import csv
import json
import sys
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
CSV_PATH = ROOT / "documentacion/cuadros-excel/CU-T13-matriz-viabilidad-inicial.csv"
REFS_43 = ROOT / "documentacion/referencias-estado-arte-43-verificadas.json"
REFS_COMP = ROOT / "documentacion/referencias-q1-complementarias-iso25000.json"
OUT_DOCX = ROOT / "Matriz_Viabilidad_Inicial_Calzatura_Vilchez.docx"
OUT_DOCX_COPY = ROOT / "documentacion/Matriz_Viabilidad_Inicial_Calzatura_Vilchez.docx"

TITULO = (
    "Sistema web de comercio electrónico con modelo de Inteligencia Artificial "
    "para la predicción del riesgo empresarial en la empresa Calzatura Vilchez, Huancayo, 2026"
)

DIMENSION_WEIGHTS: dict[str, float] = {
    "Tecnica": 0.25,
    "Economica": 0.15,
    "Legal": 0.10,
    "Operativa": 0.20,
    "Temporal": 0.10,
    "Estrategica": 0.10,
    "Cientifica": 0.10,
}

DIMENSION_LABELS: dict[str, str] = {
    "Tecnica": "T — Técnica",
    "Economica": "E — Económica",
    "Legal": "L — Legal / normativa",
    "Operativa": "O — Operativa / organizacional",
    "Temporal": "S — Temporal (Schedule)",
    "Estrategica": "Es — Estratégica / mercado",
    "Cientifica": "C — Científica / investigación",
}

ESCALA = [
    ("5", "Muy viable", "Evidencia implementada y verificable en repo/producción"),
    ("4", "Viable", "Evidencia sólida con reservas menores documentadas"),
    ("3", "Viable con reservas", "Factible pero requiere mitigación explícita"),
    ("2", "Poco viable", "Brechas significativas sin plan de cierre"),
    ("1", "No viable", "Imposibilidad técnica, legal o económica demostrada"),
]

BIBLIOGRAFIA_Q1 = [
    ("Marco TELOS / gestión proyectos", "Shen et al. (2018). Managing scope in IT projects.", "10.1016/j.infsof.2017.08.006"),
    ("Transformación digital", "Verhoef et al. (2019). Digital transformation: A multidisciplinary reflection.", "10.1016/j.jbusres.2019.09.022"),
    ("Retail / cadena valor", "Reinartz et al. (2019). Digital transformation and retailing value chain.", "10.1016/j.ijresmar.2018.12.002"),
    ("E-commerce + IA PYME", "Dai et al. (2024). AI and big data in cross-border e-commerce.", "10.1371/journal.pone.0305639"),
    ("Microservicios", "Jamshidi et al. (2018). Microservices: journey and challenges.", "10.1109/ms.2018.2141039"),
    ("Calidad microservicios", "Li et al. (2021). Quality attributes of microservices.", "10.1016/j.infsof.2020.106449"),
    ("ML en producción", "Paleyes et al. (2022). Challenges deploying ML.", "10.1145/3533378"),
    ("ML industrial", "Lwakatare et al. (2020). Large-scale ML in industry.", "10.1016/j.infsof.2020.106368"),
    ("IA para decisiones", "Duan et al. (2019). AI for decision making in Big Data era.", "10.1016/j.ijinfomgt.2019.01.021"),
    ("Big data operaciones", "Choi et al. (2018). Big data analytics in operations.", "10.1111/poms.12838"),
    ("Big data desempeño PYME", "Amba et al. (2017). Big data analytics and firm performance.", "10.1016/j.jbusres.2016.08.009"),
    ("Seguridad zero-trust", "Buck et al. (2021). Never trust, always verify.", "10.1016/j.cose.2021.102436"),
    ("Retail digitalización", "Hagberg et al. (2016). Digitalization of retailing.", "10.1108/ijrdm-09-2015-0140"),
    ("Decisiones organizacionales IA", "Shrestha et al. (2019). Organizational decision-making and AI.", "10.1177/0008125619862257"),
    ("Metodologías ágiles", "Dingsøyr et al. (2012). A decade of agile methodologies.", "10.1016/j.jss.2012.02.033"),
    ("Ciclo de vida IA", "Haakman et al. (2021). AI lifecycle models need revision.", "10.1007/s10664-021-09993-1"),
    ("Pronóstico / forecasting", "Makridakis et al. (2018). Statistical and ML forecasting.", "10.1371/journal.pone.0194889"),
    ("Retail forecasting", "Fildes et al. (2019). Retail forecasting research and practice.", "10.1016/j.ijforecast.2019.06.004"),
    ("E-SQ moda/calzado", "Gutiérrez-Rodríguez et al. (2020). E-SQ fashion e-retailers.", "10.1016/j.jretconser.2020.102201"),
    ("Marketing digital", "Kannan & Li (2017). Digital marketing framework.", "10.1016/j.ijresmar.2016.11.006"),
]

LIMITACIONES = [
    ("Validación empírica O₂", "Encuesta Likert y mediciones post-implementación en curso", "CU-T09; instrumento investigación"),
    ("Datos históricos PYME", "Series cortas limitan precisión inicial del Random Forest", "Fallback conservador; backtest documentado"),
    ("Fiabilidad producción", "Checklist fiabilidad 84% (no 100%) por DR/uptime pendientes", "fiabilidad-trazabilidad-iso25000.md"),
    ("Adopción operativa", "Capacitación admin y operación sostenida requieren seguimiento", "Runbook 10-operacion-y-seguridad.md"),
]

MATRIZ_HEADERS = [
    "Dim",
    "ID",
    "Criterio",
    "Indicador",
    "Evidencia",
    "Art. Q1",
    "DOI",
    "Pts",
    "Veredicto",
]


def set_cell_text(cell, text: str, bold: bool = False, font_size: float = 8.0) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text or "—")
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_widths(table, widths: list[float]) -> None:
    for row in table.rows:
        for index, width in enumerate(widths):
            if index < len(row.cells):
                row.cells[index].width = Inches(width)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    sizes = {1: 12, 2: 11, 3: 10}
    run.font.size = Pt(sizes.get(level, 9))
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)


def add_body(doc: Document, text: str, italic: bool = False, size: float = 9.0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.italic = italic


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float],
    header_fill: str = "D9EAF7",
) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_widths(table, widths)
    for col_index, header in enumerate(headers):
        cell = table.cell(0, col_index)
        set_cell_text(cell, header, bold=True, font_size=8)
        shade(cell, header_fill)
    for row_index, row in enumerate(rows, start=1):
        fill = "F9FBFD" if row_index % 2 == 0 else "FFFFFF"
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            set_cell_text(cell, value, font_size=7.5)
            shade(cell, fill)


def load_csv_rows() -> list[dict[str, str]]:
    with CSV_PATH.open(encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f))


def parse_score(value: str) -> float:
    try:
        score = float(value.strip())
    except ValueError as exc:
        raise ValueError(f"Puntaje inválido: {value!r}") from exc
    if score < 1 or score > 5:
        raise ValueError(f"Puntaje fuera de escala 1–5: {score}")
    return score


def parse_weight(value: str) -> float:
    try:
        return float(value.strip())
    except ValueError as exc:
        raise ValueError(f"Peso inválido: {value!r}") from exc


def compute_scores(rows: list[dict[str, str]]) -> tuple[dict[str, float], float]:
    by_dim: dict[str, list[tuple[float, float]]] = defaultdict(list)
    for row in rows:
        dim = row["Dimension"]
        by_dim[dim].append((parse_score(row["Puntaje_1a5"]), parse_weight(row["Peso_criterio_pct"])))

    dim_scores: dict[str, float] = {}
    for dim, pairs in by_dim.items():
        total_w = sum(w for _, w in pairs)
        if total_w <= 0:
            raise ValueError(f"Dimensión {dim} sin pesos válidos")
        dim_scores[dim] = sum(s * w for s, w in pairs) / total_w

    global_score = sum(dim_scores.get(dim, 0) * w for dim, w in DIMENSION_WEIGHTS.items())
    return dim_scores, global_score


def verdict_global(score: float) -> str:
    if score >= 4.5:
        return "VIABLE — proceder con el proyecto (alta confianza)"
    if score >= 4.0:
        return "VIABLE — proceder con reservas documentadas"
    if score >= 3.0:
        return "VIABLE CONDICIONAL — requiere plan de mitigación antes de escalar"
    if score >= 2.0:
        return "POCO VIABLE — replantear alcance o recursos"
    return "NO VIABLE — detener o reformular el proyecto"


def load_article_index() -> dict[str, dict]:
    index: dict[str, dict] = {}
    if REFS_43.exists():
        index.update(json.loads(REFS_43.read_text(encoding="utf-8")))
    if REFS_COMP.exists():
        index.update(json.loads(REFS_COMP.read_text(encoding="utf-8")))
    return index


def validate(rows: list[dict[str, str]]) -> list[str]:
    errors: list[str] = []
    if len(rows) < 30:
        errors.append(f"CU-T13 tiene pocas filas ({len(rows)}); se esperan ≥30.")
    dims_found = {r["Dimension"] for r in rows}
    for dim in DIMENSION_WEIGHTS:
        if dim not in dims_found:
            errors.append(f"Falta dimensión obligatoria: {dim}")
    articles = load_article_index()
    for row in rows:
        if not row.get("Criterio_ID"):
            errors.append("Fila sin Criterio_ID")
        num = row.get("Articulo_num", "").strip()
        doi = (row.get("DOI") or "").strip()
        if num == "Norma":
            continue
        if num not in articles:
            errors.append(f"{row.get('Criterio_ID')}: artículo {num} no está en corpus 43+complementos")
        elif doi and articles[num].get("doi") and articles[num]["doi"] not in doi:
            errors.append(f"{row.get('Criterio_ID')}: DOI no coincide con artículo {num}")
        score = parse_score(row["Puntaje_1a5"])
        veredicto = (row.get("Veredicto_criterio") or "").lower()
        if score >= 5 and "reserva" in veredicto:
            errors.append(f"{row.get('Criterio_ID')}: puntaje 5 con veredicto que menciona reserva")
        if score <= 3 and "alta viabilidad" in veredicto:
            errors.append(f"{row.get('Criterio_ID')}: puntaje ≤3 con veredicto Alta viabilidad")
    try:
        compute_scores(rows)
    except ValueError as exc:
        errors.append(str(exc))
    return errors


def build_docx(rows: list[dict[str, str]], dim_scores: dict[str, float], global_score: float) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MATRIZ DE VIABILIDAD INICIAL DEL PROYECTO")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(14)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(TITULO)
    run.font.name = "Arial"
    run.font.size = Pt(9)

    add_body(
        doc,
        "Metodología: marco TELOS (Technical, Economic, Legal, Operational, Schedule) ampliado "
        "con viabilidad estratégica de mercado y viabilidad científica-investigación. "
        "Cuadro trazable CU-T13 · documentacion/matriz-viabilidad-inicial-proyecto.md · "
        "formato-09-alcance-proyecto-software.md · 07-modulo-ia-riesgo-empresarial.md.",
        italic=True,
        size=8,
    )
    meta = doc.add_paragraph()
    run = meta.add_run(
        "Empresa: Calzatura Vilchez, Huancayo, Perú   |   Asesor: Dr. Maglioni Arana Caparachin   |   "
        "Versión: 1.0   |   Fecha: junio 2026"
    )
    run.font.name = "Arial"
    run.font.size = Pt(8)

    add_heading(doc, "1. Fundamento metodológico", 1)
    add_body(
        doc,
        "La viabilidad inicial evalúa si el proyecto de software es factible antes de comprometer "
        "recursos significativos. El acrónimo TELOS (Technical, Economic, Legal, Operational, Schedule) "
        "es el marco estándar en gestión de proyectos de sistemas de información (Shen et al., 2018; "
        "Kuhrmann et al., 2018). Para una tesis de ingeniería con componente de IA aplicada a una PYME "
        "retail, se añaden dos dimensiones: (Es) estratégica — alineación con transformación digital y "
        "ventaja competitiva (Verhoef et al., 2019; Reinartz et al., 2019); y (C) científica — contribución "
        "al estado del arte y replicabilidad empírica (Makridakis et al., 2018; Haakman et al., 2021).",
    )
    add_body(
        doc,
        "Cada criterio se puntúa en escala Likert 1–5, se pondera dentro de su dimensión y se respalda "
        "con evidencia del repositorio Calzatura Vilchez y al menos un artículo Q1 verificado (DOI) o "
        "norma aplicable (Ley 29571, Ley 29733).",
    )

    add_heading(doc, "1.1 Dimensiones y ponderación global", 2)
    add_table(
        doc,
        ["Dimensión", "Peso global", "Fundamento Q1 / normativo"],
        [
            ("T — Técnica", "25%", "Jamshidi et al. 2018; Li et al. 2021; Paleyes et al. 2022"),
            ("E — Económica", "15%", "Verhoef et al. 2019; Reinartz et al. 2019; Amba et al. 2017"),
            ("L — Legal", "10%", "Ley 29571; Ley 29733; Buck et al. 2021"),
            ("O — Operativa", "20%", "Choi et al. 2018; Hagberg et al. 2016; Dingsøyr et al. 2012"),
            ("S — Temporal", "10%", "Dingsøyr et al. 2012; Kuhrmann et al. 2018; Haakman et al. 2021"),
            ("Es — Estratégica", "10%", "Verhoef et al. 2019; Dai et al. 2024; Shrestha et al. 2019"),
            ("C — Científica", "10%", "Makridakis et al. 2018; Duan et al. 2019; Fildes et al. 2019"),
        ],
        [1.4, 0.8, 4.3],
    )

    add_heading(doc, "1.2 Escala de puntuación", 2)
    add_table(doc, ["Puntaje", "Etiqueta", "Criterio interpretación"], ESCALA, [0.6, 1.2, 4.7])

    add_heading(doc, "2. Resumen ejecutivo de viabilidad", 1)
    global_verdict = verdict_global(global_score)
    add_body(
        doc,
        f"Puntaje global ponderado: {global_score:.2f} / 5.00. Veredicto: {global_verdict}. "
        f"Total criterios evaluados: {len(rows)}. "
        "El proyecto presenta viabilidad técnica, económica, legal y operativa demostrada en el repositorio; "
        "las reservas principales corresponden a validación empírica post-implementación (O₂) y métricas "
        "de fiabilidad en producción aún parciales (84%).",
    )

    add_heading(doc, "2.1 Puntaje por dimensión", 2)
    summary_rows = []
    for dim, weight in DIMENSION_WEIGHTS.items():
        score = dim_scores.get(dim, 0)
        summary_rows.append(
            [
                DIMENSION_LABELS.get(dim, dim),
                f"{weight * 100:.0f}%",
                f"{score:.2f}",
                verdict_global(score).split("—")[0].strip(),
            ]
        )
    add_table(
        doc,
        ["Dimensión", "Peso", "Puntaje dim.", "Veredicto dim."],
        summary_rows,
        [2.2, 0.7, 0.9, 2.7],
        header_fill="E2F0D9",
    )

    for dim in DIMENSION_WEIGHTS:
        dim_rows = [r for r in rows if r["Dimension"] == dim]
        if not dim_rows:
            continue
        add_heading(doc, f"3.{list(DIMENSION_WEIGHTS.keys()).index(dim) + 1} {DIMENSION_LABELS.get(dim, dim)}", 2)
        table_data = [
            [
                r["Criterio_ID"],
                r["Criterio"],
                r["Indicador_evaluacion"][:80] + ("…" if len(r["Indicador_evaluacion"]) > 80 else ""),
                r["Puntaje_1a5"],
                r["Veredicto_criterio"],
            ]
            for r in dim_rows
        ]
        add_table(
            doc,
            ["ID", "Criterio", "Indicador", "Pts", "Veredicto"],
            table_data,
            [0.7, 1.8, 2.5, 0.5, 1.0],
        )
        add_body(
            doc,
            f"Puntaje dimensión {DIMENSION_LABELS.get(dim, dim)}: {dim_scores.get(dim, 0):.2f}/5.00.",
            italic=True,
            size=8,
        )

    add_heading(doc, "4. Limitaciones y reservas declaradas", 1)
    add_table(doc, ["Aspecto", "Limitación", "Mitigación"], LIMITACIONES, [1.4, 2.5, 2.5])

    add_heading(doc, "5. Respaldo bibliográfico Q1 (alto impacto)", 1)
    add_table(
        doc,
        ["Dimensión", "Referencia", "DOI"],
        [[a, b, f"https://doi.org/{c}"] for a, b, c in BIBLIOGRAFIA_Q1],
        [1.3, 3.6, 1.6],
    )

    add_heading(doc, "6. Referencias metodológicas", 1)
    refs = [
        "Shen, Y., Collier, R. E., & Wang, Y. (2018). Managing scope in IT projects. Information and Software Technology.",
        "Kuhrmann, M., et al. (2018). Hybrid software development approaches in practice. IEEE Software.",
        "Verhoef, P. C., et al. (2019). Digital transformation: A multidisciplinary reflection. Journal of Business Research.",
        "Reinartz, W., et al. (2019). The impact of digital transformation on the retailing value chain. International Journal of Research in Marketing.",
        "Dai, Y., et al. (2024). Revolutionizing cross-border e-commerce with AI and big data. PLOS ONE.",
        "Paleyes, A., et al. (2022). Challenges in deploying machine learning. ACM Computing Surveys.",
        "Project Management Institute (2021). A Guide to the Project Management Body of Knowledge (PMBOK Guide), 7th ed.",
    ]
    for ref in refs:
        add_body(doc, f"• {ref}", size=8)

    doc.add_page_break()
    landscape = doc.add_section()
    landscape.orientation = WD_ORIENT.LANDSCAPE
    landscape.page_width, landscape.page_height = landscape.page_height, landscape.page_width
    landscape.top_margin = Inches(0.35)
    landscape.bottom_margin = Inches(0.35)
    landscape.left_margin = Inches(0.35)
    landscape.right_margin = Inches(0.35)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run("ANEXO — MATRIZ CU-T13 VIABILIDAD INICIAL (TRAZABILIDAD COMPLETA)")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)

    matrix_rows = [
        [
            r["Dimension"][:3],
            r["Criterio_ID"],
            r["Criterio"][:45] + ("…" if len(r["Criterio"]) > 45 else ""),
            r["Indicador_evaluacion"][:55] + ("…" if len(r["Indicador_evaluacion"]) > 55 else ""),
            r["Evidencia_proyecto"][:40] + ("…" if len(r["Evidencia_proyecto"]) > 40 else ""),
            f"#{r['Articulo_num']} {r['Articulo_cita'][:35]}…" if len(r["Articulo_cita"]) > 35 else f"#{r['Articulo_num']} {r['Articulo_cita']}",
            r["DOI"][:28] if r["DOI"] else "Norma",
            r["Puntaje_1a5"],
            r["Veredicto_criterio"],
        ]
        for r in rows
    ]
    table = doc.add_table(rows=len(matrix_rows) + 1, cols=len(MATRIZ_HEADERS))
    table.style = "Table Grid"
    table.autofit = False
    set_widths(table, [0.35, 0.55, 1.05, 1.25, 1.05, 1.35, 0.95, 0.35, 0.75])
    for col_index, header in enumerate(MATRIZ_HEADERS):
        cell = table.cell(0, col_index)
        set_cell_text(cell, header, bold=True, font_size=7)
        shade(cell, "1F4E79")
    for row_index, row in enumerate(matrix_rows, start=1):
        score = float(row[7])
        fill = "E2EFDA" if score >= 4.5 else ("FFF2CC" if score >= 3.5 else "FCE4D6")
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            set_cell_text(cell, str(value), font_size=6.5)
            shade(cell, fill)

    doc.add_section()
    portrait = doc.sections[-1]
    portrait.orientation = WD_ORIENT.PORTRAIT
    portrait.page_width, portrait.page_height = Inches(8.27), Inches(11.69)
    portrait.top_margin = Inches(0.7)
    portrait.bottom_margin = Inches(0.7)
    portrait.left_margin = Inches(0.8)
    portrait.right_margin = Inches(0.8)

    add_heading(doc, "7. Resumen de validación documental", 1)
    add_body(
        doc,
        f"Total filas CU-T13: {len(rows)}. Puntaje global: {global_score:.2f}/5.00. "
        f"Veredicto: {global_verdict}. "
        f"Regeneración: python scripts/generar_matriz_viabilidad_inicial.py",
    )

    doc.save(OUT_DOCX)
    doc.save(OUT_DOCX_COPY)


def main() -> int:
    if not CSV_PATH.exists():
        print(f"ERROR: no existe {CSV_PATH}", file=sys.stderr)
        return 1
    rows = load_csv_rows()
    issues = validate(rows)
    if issues:
        for issue in issues:
            print(f"VALIDACIÓN: {issue}", file=sys.stderr)
        return 1
    dim_scores, global_score = compute_scores(rows)
    build_docx(rows, dim_scores, global_score)
    print(f"OK: {OUT_DOCX}")
    print(f"OK: {OUT_DOCX_COPY}")
    print(f"Filas CU-T13: {len(rows)}")
    print(f"Puntaje global: {global_score:.2f}/5.00 — {verdict_global(global_score)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
