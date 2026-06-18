from pathlib import Path
import re
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


BASE = Path("c:/Cazatura Vilchez V3")
SRC = BASE / "Plan_Tesis_Calzatura_Vilchez.docx"
OUT = BASE / "Plan_Tesis_Calzatura_Vilchez_MEJORADO_42_ARTICULOS.docx"
BACKUP = BASE / "Plan_Tesis_Calzatura_Vilchez_BACKUP_ANTES_MEJORA.docx"


def norm(value: str) -> str:
    return " ".join((value or "").split())


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_paragraph_before(
    target: Paragraph,
    text: str = "",
    style: str | None = None,
    bold: bool = False,
    italic: bool = False,
    align=None,
    font_size: float | None = None,
) -> Paragraph:
    new_element = OxmlElement("w:p")
    target._p.addprevious(new_element)
    paragraph = Paragraph(new_element, target._parent)
    if style:
        try:
            paragraph.style = style
        except Exception:
            pass
    if text:
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        if font_size:
            run.font.size = Pt(font_size)
    if align is not None:
        paragraph.alignment = align
    return paragraph


def set_cell_text(cell, text: str, bold: bool = False, font_size: float = 8) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_table_before(
    doc: Document,
    target: Paragraph,
    rows: list[list[str]],
    headers: list[str],
    title: str,
    note: str,
    font_size: float = 7,
    widths: list[float] | None = None,
) -> None:
    caption = insert_paragraph_before(target, title, style="Caption", bold=True, font_size=9)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    data = [headers, *rows]
    table = doc.add_table(rows=len(data), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row_index, row in enumerate(data):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            set_cell_text(
                cell,
                value,
                bold=row_index == 0,
                font_size=font_size + 1 if row_index == 0 else font_size,
            )
            if row_index == 0:
                shade_cell(cell, "D9EAF7")

    if widths:
        for row in table.rows:
            for col_index, width in enumerate(widths):
                if col_index < len(row.cells):
                    row.cells[col_index].width = Inches(width)

    target._p.addprevious(table._tbl)
    note_paragraph = insert_paragraph_before(
        target, note, style="Normal", italic=True, font_size=8
    )
    note_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    insert_paragraph_before(target, "")


def find_paragraph(doc: Document, predicate) -> Paragraph:
    for paragraph in doc.paragraphs:
        if predicate(norm(paragraph.text)):
            return paragraph
    raise ValueError("No se encontró el párrafo requerido.")


def get_articles() -> list[list[str]]:
    state_doc = Document(BASE / "Estado_del_Arte_42_Tablas_v2.docx")
    articles: list[list[str]] = []

    for index, paragraph in enumerate(state_doc.paragraphs):
        text = paragraph.text.strip()
        match = re.match(r"Art.culo\s*\[(\d+)\]\s*\u2014\s*(.+)", text)
        if not match:
            continue

        number = int(match.group(1))
        detail = (
            state_doc.paragraphs[index + 1].text.strip()
            if index + 1 < len(state_doc.paragraphs)
            else ""
        )
        source = detail.replace(" — ", " - ")

        if 1 <= number <= 8:
            axis = "Eje 1: e-commerce y transformación digital"
            contribution = (
                "Sustenta la digitalización del canal de venta, la madurez digital "
                "y la reconfiguración del retail."
            )
            use = "Cap. I, 2.2.1 y VI-D1"
        elif 9 <= number <= 17:
            axis = "Eje 2: IA y analítica de datos"
            contribution = (
                "Sustenta el uso de IA, analítica empresarial, pronóstico de demanda "
                "y decisiones basadas en datos."
            )
            use = "2.2.2, 3.6 y VI-D2"
        elif 18 <= number <= 27:
            axis = "Eje 3: predicción de riesgo empresarial"
            contribution = (
                "Sustenta modelos ML de riesgo/quiebra, métricas de validación "
                "y adaptación del IRE."
            )
            use = "2.2.3, VD-D1 a VD-D3"
        elif 28 <= number <= 34:
            axis = "Eje 4: arquitectura y despliegue ML"
            contribution = (
                "Sustenta microservicios, APIs, despliegue de ML, calidad "
                "y monitoreo del sistema."
            )
            use = "2.2.1, 3.5, VI-D3"
        else:
            axis = "Eje 5: metodología y ciclo de vida ML"
            contribution = (
                "Sustenta CRISP-ML(Q), desarrollo ágil/híbrido, analítica prescriptiva "
                "y valor de TI."
            )
            use = "3.1, 3.2, 3.6, VI-D4"

        articles.append([f"{number:02d}", source, axis, contribution, use])

    if len(articles) != 42:
        raise ValueError(f"Se esperaban 42 artículos y se leyeron {len(articles)}.")
    return articles


def get_reference_summaries() -> list[str]:
    source_doc = Document(SRC)
    start_index = next(
        i for i, paragraph in enumerate(source_doc.paragraphs) if "REFERENCIAS" in paragraph.text
    )
    end_index = next(
        (
            i
            for i, paragraph in enumerate(source_doc.paragraphs[start_index + 1 :], start_index + 1)
            if norm(paragraph.text) == "ANEXOS"
        ),
        len(source_doc.paragraphs),
    )
    references = []
    for paragraph in source_doc.paragraphs[start_index + 1 : end_index]:
        text = norm(paragraph.text)
        if not text:
            continue
        text = re.sub(r"\s*\[consulta:.*", "", text)
        text = re.sub(r"\s*DOI:.*", "", text)
        references.append(text)
    if len(references) != 42:
        raise ValueError(f"Se esperaban 42 referencias y se leyeron {len(references)}.")
    return references


YEAR_REPLACEMENTS = {
    "Fildes et al. (2019)": "Fildes et al. (2022)",
    "Fildes 2019": "Fildes 2022",
    "FILDES, R.; Ma, S. y Kolassa, S. (2019)": "FILDES, R.; Ma, S. y Kolassa, S. (2022)",
    "Osaka (2019)": "Osaka (2021)",
    "Osaka 2019": "Osaka 2021",
    "Kim et al. (2021)": "Kim et al. (2020)",
    "Kim 2021": "Kim 2020",
    "Cai et al. (2019)": "Cai et al. (2021)",
    "Cai 2019": "Cai 2021",
    "Kuhrmann et al. (2018)": "Kuhrmann et al. (2019)",
    "Kuhrmann 2018": "Kuhrmann 2019",
}


def replace_year_mentions(text: str) -> str:
    for old, new in YEAR_REPLACEMENTS.items():
        text = text.replace(old, new)
    return text


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(SRC, BACKUP)
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)

    for paragraph in doc.paragraphs:
        text = norm(paragraph.text)
        if "El presupuesto total estimado asciende a S/. 6,890.00" in text:
            paragraph.text = (
                "La investigación es autofinanciada por el investigador. El presupuesto total estimado "
                "asciende a S/. 6,028.00 (seis mil veintiocho soles), de acuerdo con la tabla de "
                "presupuesto del Capítulo IV. Este monto cubre recursos humanos autofinanciados, "
                "hardware, servicios cloud, materiales, transporte e imprevistos."
            )
        elif "La matriz de operacionalización de variables se presenta en el Anexo 02" in text:
            paragraph.text = (
                "La matriz de operacionalización de variables se presenta en el Anexo 02 y se articula "
                "con la Tabla General del Estado del Arte incorporada en el Capítulo II. La matriz "
                "mantiene la simetría metodológica requerida: la variable independiente y la variable "
                "dependiente se organizan en cuatro dimensiones cada una, con indicadores derivados "
                "de los 42 artículos científicos Q1 revisados."
            )
        elif "cargados en la base de datos Supabase/PostgreSQL" in text:
            paragraph.text = text.replace(
                "cargados en la base de datos Supabase/PostgreSQL",
                "integrados con Firebase Firestore y procesados en el entorno Python del microservicio de IA",
            )
        elif "La tabla completa con celdas fusionadas se encuentra en el documento" in text:
            paragraph.text = (
                "La matriz de operacionalización de variables presenta las dimensiones, indicadores, "
                "escala de medición e instrumentos para las variables VI y VD de la investigación. "
                "La tabla está estructurada de forma simétrica: 4 dimensiones para la VI y 4 dimensiones "
                "para la VD. Su sustento académico directo es la Tabla General del Estado del Arte "
                "incorporada en el Capítulo II, elaborada a partir de los 42 artículos científicos Q1."
            )

    paragraphs = doc.paragraphs
    start_index = next(
        i for i, paragraph in enumerate(paragraphs) if norm(paragraph.text).startswith("2.1.")
    )
    end_index = next(
        i for i, paragraph in enumerate(paragraphs) if norm(paragraph.text).startswith("2.2.")
    )
    target = paragraphs[end_index]
    for paragraph in list(paragraphs[start_index:end_index]):
        delete_paragraph(paragraph)

    articles = get_articles()
    reference_summaries = get_reference_summaries()
    for index, reference in enumerate(reference_summaries):
        articles[index][1] = reference

    insert_paragraph_before(target, "2.1. Antecedentes de la Investigación", bold=True, font_size=12)
    insert_paragraph_before(
        target,
        "La presente investigación organiza sus antecedentes exclusivamente a partir de 42 artículos "
        "científicos Q1 revisados y verificados en Scimago Journal Rankings y/o Journal Citation "
        "Reports. Para evitar desalineación académica, no se incorporan antecedentes externos a esa "
        "base documental en la construcción del marco teórico, la operacionalización de variables ni "
        "la selección de indicadores. Los artículos se agrupan en cinco ejes temáticos que corresponden "
        "directamente con el título de la tesis: sistema web de comercio electrónico, modelo de "
        "Inteligencia Artificial y predicción del riesgo empresarial en Calzatura Vilchez.",
        font_size=10,
    )
    insert_paragraph_before(
        target,
        "2.1.1. Criterio de selección y alineamiento de los 42 artículos científicos",
        bold=True,
        font_size=11,
    )
    insert_paragraph_before(
        target,
        "Los artículos seleccionados cumplen tres criterios: pertinencia directa con las variables de "
        "estudio, calidad científica Q1 y utilidad metodológica para definir dimensiones, indicadores "
        "o instrumentos. La variable independiente se fundamenta en los ejes 1, 2, 4 y 5, mientras que "
        "la variable dependiente se apoya principalmente en el eje 3 y en los aportes de analítica "
        "operativa y valor de TI del eje 5. Esta organización permite que cada apartado del plan derive "
        "de una fuente científica revisada y no de afirmaciones aisladas.",
        font_size=10,
    )
    insert_paragraph_before(
        target,
        "2.1.2. Síntesis del estado del arte por eje temático",
        bold=True,
        font_size=11,
    )

    synthesis = [
        "El estado del arte evidencia que la transformación digital en retail no solo implica vender "
        "por internet, sino capturar datos transaccionales para mejorar la toma de decisiones. Verhoef "
        "et al. (2021), Soto-Acosta (2020), Nambisan et al. (2017), Li (2020), Chatterjee y Kumar Kar "
        "(2020), Kannan (2017), Reinartz et al. (2019) y Hagberg et al. (2016) sustentan que el comercio "
        "electrónico, la digitalización de procesos y la interacción con clientes digitales generan "
        "capacidades organizacionales medibles. En Calzatura Vilchez, este eje respalda la necesidad "
        "de implementar un sistema web que registre ventas, pedidos e inventario de forma estructurada.",
        "El segundo eje sostiene la incorporación de Inteligencia Artificial y analítica de datos "
        "empresariales. Yuan et al. (2021), Davenport et al. (2020), Shrestha et al. (2019), Amba et al. "
        "(2017), Makridakis et al. (2018), Fildes et al. (2019), Syam y Sharma (2018), Haefner et al. "
        "(2021) y Paschen et al. (2019) muestran que los modelos predictivos y la analítica empresarial "
        "permiten reducir la incertidumbre en decisiones comerciales. Estos aportes justifican el uso "
        "de RandomForestRegressor, métricas MAE, RMSE y sMAPE, así como la comparación con métodos "
        "estadísticos o empíricos.",
        "El tercer eje fundamenta la variable dependiente: predicción del riesgo empresarial. Ozbayoglu "
        "et al. (2020), Barboza et al. (2017), Osaka (2019), Kim et al. (2021), Du Jardin (2021), Jiang "
        "et al. (2018), Cai et al. (2019), Tian et al. (2015), Altman (1968) y Beaver (1966) demuestran "
        "la evolución desde modelos financieros clásicos hacia modelos de Machine Learning con mayor "
        "capacidad predictiva. En esta tesis, dichos aportes se adaptan al Índice de Riesgo Empresarial "
        "(IRE), compuesto por riesgo de stock, riesgo de ingresos y riesgo de demanda.",
        "El cuarto eje explica cómo llevar el modelo a producción en un sistema real. Jamshidi et al. "
        "(2018), Swakatare et al. (2019), Khalid et al. (2022), Fischer y Krauss (2018), Paleyes et al. "
        "(2022), Martinez-Fernandez et al. (2022) y Li et al. (2021) sustentan el uso de microservicios, "
        "APIs, monitoreo, atributos de calidad y despliegue de modelos ML. Por ello, el sistema propuesto "
        "separa el frontend de comercio electrónico del microservicio de IA, permitiendo trazabilidad, "
        "mantenimiento y evaluación continua.",
        "El quinto eje articula la metodología tecnológica del proyecto. Breiman (2001), Dingsøyr et al. "
        "(2012), Salgonde y Chari (2021), Kuhrmann et al. (2018), Lepenioti et al. (2020), Choi et al. "
        "(2018), Haakman et al. (2021) y Melville et al. (2004) sustentan el algoritmo Random Forest, "
        "el desarrollo ágil/híbrido, el ciclo CRISP-ML(Q), la analítica prescriptiva y la medición del "
        "valor de TI. Con ello, el plan conecta construcción del sistema, validación del modelo y "
        "evaluación de impacto operativo.",
    ]
    for text in synthesis:
        insert_paragraph_before(target, text, font_size=10)

    axis_rows = [
        [
            "Eje 1",
            "Artículos 01-08",
            "E-commerce y transformación digital",
            "VI-D1: digitalización del negocio",
            "Cap. I, Cap. II 2.2.1 y Cap. III indicadores de digitalización",
        ],
        [
            "Eje 2",
            "Artículos 09-17",
            "IA, analítica y pronóstico",
            "VI-D2: modelo de IA y analítica",
            "Cap. II 2.2.2 y Cap. III métricas MAE/RMSE/sMAPE",
        ],
        [
            "Eje 3",
            "Artículos 18-27",
            "Predicción de riesgo/quiebra con ML",
            "VD-D1, VD-D2 y VD-D3: IRE",
            "Cap. II 2.2.3 y validación del riesgo empresarial",
        ],
        [
            "Eje 4",
            "Artículos 28-34",
            "Arquitectura, APIs y despliegue ML",
            "VI-D3: arquitectura de software",
            "Cap. II arquitectura y Cap. III pruebas técnicas",
        ],
        [
            "Eje 5",
            "Artículos 35-42",
            "Metodología, ciclo de vida ML y valor de TI",
            "VI-D4 e impacto operativo",
            "Cap. III metodología, Cap. IV y anexos",
        ],
    ]
    add_table_before(
        doc,
        target,
        axis_rows,
        headers=["Eje", "Artículos", "Tema científico", "Alineamiento con variables", "Ubicación en el plan"],
        title="Tabla 1. Organización de los 42 artículos científicos por eje temático y ubicación en el plan",
        note="Nota. Esta tabla se implementa en el Capítulo II porque define el soporte bibliográfico del marco teórico y evita que el desarrollo del plan se aparte de los 42 artículos científicos revisados.",
        font_size=7,
        widths=[0.55, 0.8, 1.45, 1.45, 2.0],
    )

    insert_paragraph_before(
        target,
        "2.1.3. Trazabilidad de los 42 artículos científicos con el plan de tesis",
        bold=True,
        font_size=11,
    )
    insert_paragraph_before(
        target,
        "La siguiente tabla cumple la función de control académico: cada artículo revisado se vincula "
        "con el eje temático, el aporte concreto y la sección del plan donde se utiliza. Así se garantiza "
        "que la ampliación del documento responda a la base científica completa y no a una selección parcial.",
        font_size=10,
    )
    add_table_before(
        doc,
        target,
        articles,
        headers=["N.°", "Artículo científico Q1", "Eje temático", "Aporte usado en la tesis", "Aplicación"],
        title="Tabla 2. Trazabilidad de los 42 artículos científicos utilizados en el plan de tesis",
        note="Nota. La trazabilidad se implementa en el Capítulo II porque corresponde al estado del arte. Los detalles individuales de variables, dimensiones, indicadores e instrumentos por artículo se conservan en el Anexo 06.",
        font_size=6,
        widths=[0.35, 1.8, 1.2, 2.0, 0.95],
    )

    insert_paragraph_before(
        target,
        "2.1.4. Matriz general del estado del arte aplicada al plan de tesis",
        bold=True,
        font_size=11,
    )
    insert_paragraph_before(
        target,
        "La Tabla General del Estado del Arte se incorpora en este capítulo porque consolida la relación "
        "entre variables, dimensiones, indicadores e instrumentos derivados de los 42 artículos científicos. "
        "Su función no es decorativa: sirve como puente entre el marco teórico y la metodología, ya que de "
        "ella se desprenden la operacionalización de variables, los instrumentos de recolección y las métricas "
        "de evaluación del sistema.",
        font_size=10,
    )

    general_doc = Document(BASE / "Tabla_General_Estado_Arte_v4.docx")
    general_table = general_doc.tables[0]
    general_rows = [[norm(cell.text) for cell in row.cells] for row in general_table.rows[1:]]
    general_headers = [norm(cell.text) for cell in general_table.rows[0].cells]
    add_table_before(
        doc,
        target,
        general_rows,
        headers=general_headers,
        title="Tabla 3. Tabla General del Estado del Arte consolidada en variables, dimensiones, indicadores e instrumentos",
        note="Nota. Tabla incorporada desde Tabla_General_Estado_Arte_v4.docx. Su ubicación correcta es el Capítulo II, sección Estado del Arte, porque sintetiza la evidencia científica que sustenta las variables de investigación. La operacionalización formal continúa en el Anexo 02.",
        font_size=5.5,
        widths=[1.25, 2.0, 1.6, 2.25],
    )

    insert_paragraph_before(
        target,
        "2.1.5. Brecha científica y aporte de la presente investigación",
        bold=True,
        font_size=11,
    )
    insert_paragraph_before(
        target,
        "La revisión de los 42 artículos muestra una brecha específica: existen estudios sólidos sobre "
        "transformación digital, comercio electrónico, IA, pronóstico de demanda, predicción de quiebra, "
        "microservicios y ciclo de vida ML, pero se encuentran dispersos. No se identifica una propuesta "
        "integrada que conecte un sistema web de comercio electrónico con un modelo de IA y un índice "
        "operativo de riesgo empresarial adaptado a una PYME peruana de retail de calzado. La presente "
        "tesis atiende esa brecha mediante el desarrollo de un sistema que captura datos transaccionales, "
        "ejecuta pronósticos de demanda y calcula el IRE como indicador anticipado de riesgo en Calzatura Vilchez.",
        font_size=10,
    )

    target_36 = find_paragraph(doc, lambda text: text.startswith("3.6."))
    insert_paragraph_before(
        target_36,
        "3.5.3. Ubicación metodológica de las tablas incorporadas",
        bold=True,
        font_size=11,
    )
    insert_paragraph_before(
        target_36,
        "Las tablas incorporadas cumplen funciones distintas dentro del plan. Las tablas del Capítulo II "
        "sostienen el estado del arte y la matriz científica; la tabla de este apartado indica cómo esa "
        "evidencia se traslada a instrumentos, indicadores y análisis de datos.",
        font_size=10,
    )
    method_rows = [
        [
            "Tabla 1",
            "Capítulo II, sección 2.1.2",
            "Ordena los 42 artículos por eje temático",
            "Permite validar que el marco teórico cubra e-commerce, IA, riesgo, arquitectura y metodología.",
        ],
        [
            "Tabla 2",
            "Capítulo II, sección 2.1.3",
            "Muestra la trazabilidad artículo-aporte-aplicación",
            "Evita desalineación y demuestra dónde se usa cada artículo dentro del plan.",
        ],
        [
            "Tabla 3",
            "Capítulo II, sección 2.1.4",
            "Consolida variables, dimensiones, indicadores e instrumentos",
            "Base directa para la operacionalización de variables y para los instrumentos de medición.",
        ],
        [
            "Tabla 4",
            "Capítulo III, sección 3.5.3",
            "Relaciona tablas con metodología",
            "Explica cómo las tablas pasan del estado del arte a la recolección y análisis de datos.",
        ],
        [
            "Anexo 06",
            "Anexos",
            "Conserva las 42 tablas individuales del estado del arte",
            "Sirve como respaldo extendido sin sobrecargar el cuerpo principal del plan.",
        ],
    ]
    add_table_before(
        doc,
        target_36,
        method_rows,
        headers=["Elemento", "Ubicación", "Función", "Uso metodológico"],
        title="Tabla 4. Ubicación y función de las tablas implementadas en el plan de tesis",
        note="Nota. Esta tabla se ubica en el Capítulo III porque conecta el estado del arte con la medición de variables, los instrumentos y el análisis de datos.",
        font_size=7,
        widths=[0.75, 1.55, 1.8, 2.4],
    )

    for paragraph in doc.paragraphs:
        text = norm(paragraph.text)
        if text.startswith("El estado del arte completo de la investigación se presenta en los documentos auxiliares"):
            paragraph.text = (
                "El estado del arte completo de la investigación se respalda en los documentos auxiliares "
                "generados durante la revisión bibliográfica. En el cuerpo del plan se incorporó la síntesis "
                "general y la matriz consolidada; los documentos auxiliares conservan el detalle extendido "
                "de los 42 artículos científicos."
            )
        elif text.startswith('2. "Tabla_General_Estado_Arte_v4.docx": Contiene la Tabla General'):
            paragraph.text = (
                '2. "Tabla_General_Estado_Arte_v4.docx": Documento fuente de la Tabla General del Estado '
                "del Arte, ya incorporada en el Capítulo II del presente plan como matriz consolidada de "
                "variables, dimensiones, indicadores e instrumentos."
            )

    for paragraph in doc.paragraphs:
        updated = replace_year_mentions(paragraph.text)
        if updated != paragraph.text:
            paragraph.text = updated

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                updated = replace_year_mentions(cell.text)
                if updated != cell.text:
                    cell.text = updated

    for paragraph in doc.paragraphs:
        if not norm(paragraph.text):
            continue
        if paragraph.style.name == "Normal":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in paragraph.runs:
                run.font.name = "Arial"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
                if run.font.size is None:
                    run.font.size = Pt(10)

    doc.save(OUT)
    print(f"Documento generado: {OUT}")
    print(f"Backup original: {BACKUP}")
    print(f"Artículos trazados: {len(articles)}")
    print(f"Tablas totales en documento: {len(Document(OUT).tables)}")


if __name__ == "__main__":
    main()
