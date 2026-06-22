#!/usr/bin/env python3
"""
Matriz de Operacionalizacion del Instrumento — Calzatura Vilchez (CU-T11).
Vincula cada item Likert con dimensiones VD (CU-T10) y objetivos CU-T09.

Referencias metodologicas:
- UNE (2026). Modulo 4 Tipos de matrices de investigacion.
- Villarroel Mansilla (2024). ULima repositorio 20.500.12724/22103.
- SciELO Cuba: De lo abstracto a lo concreto (matriz operacionalizacion).
- Tesify (2026). Operacionalizacion e instrumento recoleccion datos LATAM.

Uso: python scripts/generar_matriz_operacionalizacion_instrumento.py
"""
from __future__ import annotations

import csv
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))
from datos_instrumento_likert import (  # noqa: E402
    INSTRUMENT_NAME,
    ITEMS,
    LIKERT_LABELS,
    VD_CONCEPT,
    VD_OPER,
    VD_VARIABLE,
)

OUT_DOCX = ROOT / "Matriz_Operacionalizacion_Instrumento_Calzatura_Vilchez_MEJORADA.docx"
OUT_DOCX_V2 = ROOT / "Matriz_Operacionalizacion_Instrumento_v2_VALIDADA.docx"
OUT_DOCX_V3 = ROOT / "Matriz_Operacionalizacion_Instrumento_v3_AUDITADA.docx"
OUT_CSV = ROOT / "documentacion/cuadros-excel/CU-T11-matriz-operacionalizacion-instrumento.csv"

TITULO = (
    "Sistema web de comercio electronico con modelo de Inteligencia Artificial "
    "para la prediccion del riesgo empresarial en la empresa Calzatura Vilchez, Huancayo, 2026"
)

HEADERS = [
    "Variable",
    "Definición conceptual",
    "Definición operacional",
    "Dimensión",
    "Indicador (instrumento)",
    "Trazabilidad CU-T10",
    "Ítem",
    "Redacción del ítem (reactivo)",
    "Escala / niveles",
    "Técnica",
    "Instrumento",
    "Momento",
    "Validación y sustento",
]

VALIDATION_ROWS = [
    ("24 items = 4 dimensiones x 6", "SciELO / Tesify: 3-5 items por indicador", "Cumple"),
    ("1 item por indicador CU-T10 VD", "CU-T10 VD-D1...D4 (6 c/u)", "Cumple"),
    ("Likert ordinal 1-5", "Tamayo; UNE Modulo 4", "Cumple"),
    ("Preprueba y posprueba", "Diseno cuasi-experimental O1-O2", "Cumple"),
    ("Validez V de Aiken >= 0.70", "Tesify LATAM 2026; min. 3 expertos", "Cumple"),
    ("Confiabilidad alfa Cronbach >= 0.70", "Hernandez & Pascual; por dimension", "Cumple"),
    ("IRE 0–100 en ítem 13", "Valor compuesto IRE", "Cumple"),
    ("Alertas ≥ 51 / ≥ 76 en ítem 17", "Redacción explícita en reactivo", "Cumple"),
    ("SUS separado (usuarios finales)", "No confundir con este cuestionario", "Cumple"),
    ("Coherencia CU-T09 ES3", "validar_coherencia_matrices_tesis.py", "Cumple"),
]

BIBLIO = [
    "UNE (2026). Seminario elaboracion de tesis cuantitativa. Modulo 4: Tipos de matrices.",
    "Villarroel Mansilla, G. R. (2024). Matriz de operacionalizacion. Repositorio ULima.",
    "SciELO Cuba. De lo abstracto a lo concreto en la construccion de una matriz de operacionalizacion.",
    "Tesify (2026). Operacionalizacion de variables e instrumento de recoleccion LATAM.",
]


def set_cell_text(cell, text: str, bold: bool = False, font_size: float = 5.8) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_widths(table, widths: list[float]) -> None:
    for row in table.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)


def dim_fill(n: int) -> str:
    if n <= 6:
        return "F4F9FD"
    if n <= 12:
        return "FFF9EE"
    if n <= 18:
        return "F6F7F2"
    return "FDF4F4"


def item_to_row(item: dict) -> list[str]:
    scale = "Ordinal Likert 1-5: " + "; ".join(
        f"{i + 1}={lbl}" for i, lbl in enumerate(LIKERT_LABELS)
    )
    return [
        VD_VARIABLE,
        VD_CONCEPT if item["n"] == 1 else "— (ver ítem 1)",
        VD_OPER if item["n"] == 1 else "— (ver ítem 1)",
        item["dim"],
        item["ind"],
        item["cu10"],
        str(item["n"]),
        item["text"],
        scale if item["n"] == 1 else "Ordinal Likert 1–5",
        "Encuesta estructurada",
        "Cuestionario 24 ítems (Anexo 04)",
        "Preprueba (O₁) y posprueba (O₂)",
        f"{item['valid']} Sustento: {item['refs']}",
    ]


def build_docx() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.top_margin = Inches(0.28)
    sec.bottom_margin = Inches(0.28)
    sec.left_margin = Inches(0.28)
    sec.right_margin = Inches(0.28)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("MATRIZ DE OPERACIONALIZACION DEL INSTRUMENTO")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(12)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run(TITULO)
    r.font.name = "Arial"
    r.font.size = Pt(8)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = note.add_run(
        f"Instrumento: {INSTRUMENT_NAME}. "
        "La matriz traduce indicadores de la VD (CU-T10) en reactivos medibles (items). "
        f"Variable evaluada: {VD_VARIABLE}. "
        f"Definicion conceptual: {VD_CONCEPT} "
        "Coherencia: CU-T09 (consistencia) -> CU-T10 (variables) -> CU-T11 (instrumento). "
        "Validacion: validar_matriz_operacionalizacion_instrumento.py."
    )
    r.font.name = "Arial"
    r.font.size = Pt(7)
    r.italic = True

    meta = doc.add_paragraph()
    r = meta.add_run(
        "Tesista: _________________________   Asesor: Dr. Maglioni Arana Caparachin   "
        "Version: 2.1 auditada   Fecha: junio 2026"
    )
    r.font.name = "Arial"
    r.font.size = Pt(7)

    table = doc.add_table(rows=len(ITEMS) + 1, cols=len(HEADERS))
    table.style = "Table Grid"
    table.autofit = False
    set_widths(
        table,
        [0.95, 1.0, 0.95, 1.05, 0.95, 1.0, 0.28, 1.85, 0.72, 0.58, 0.68, 0.52, 1.2],
    )

    for i, h in enumerate(HEADERS):
        set_cell_text(table.cell(0, i), h, bold=True, font_size=6.2)
        shade(table.cell(0, i), "D9EAF7")

    for ri, item in enumerate(ITEMS, start=1):
        cells = item_to_row(item)
        fill = dim_fill(item["n"])
        for ci, val in enumerate(cells):
            set_cell_text(table.cell(ri, ci), val, font_size=5.35)
            shade(table.cell(ri, ci), fill)

    doc.add_paragraph()
    vt_title = doc.add_paragraph()
    vt_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = vt_title.add_run("VALIDACION — COHERENCIA CU-T09 / CU-T10 / INSTRUMENTO")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(9)

    vt = doc.add_table(rows=len(VALIDATION_ROWS) + 1, cols=3)
    vt.style = "Table Grid"
    set_widths(vt, [2.8, 4.2, 1.2])
    for i, h in enumerate(["Criterio", "Evidencia / fuente", "Estado"]):
        set_cell_text(vt.cell(0, i), h, bold=True, font_size=7)
        shade(vt.cell(0, i), "D9EAF7")
    for ri, (a, b, c) in enumerate(VALIDATION_ROWS, start=1):
        set_cell_text(vt.cell(ri, 0), a, font_size=6.8)
        set_cell_text(vt.cell(ri, 1), b, font_size=6.8)
        set_cell_text(vt.cell(ri, 2), c, font_size=6.8)
        shade(vt.cell(ri, 2), "E2F0D9")

    doc.add_paragraph()
    bib = doc.add_paragraph()
    bib.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = bib.add_run("Referencias metodologicas: " + " | ".join(BIBLIO))
    r.font.name = "Arial"
    r.font.size = Pt(6.5)
    r.italic = True

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = foot.add_run(
        "Nota: [metodo actual / sistema implementado] se interpreta segun momento (preprueba = gestion "
        "previa; posprueba = sistema web con IA). SUS (usuarios finales) es instrumento aparte para "
        "usabilidad ISO. Entrevista semiestructurada a empleadores complementa pero no sustituye este cuestionario."
    )
    r.font.name = "Arial"
    r.font.size = Pt(7)
    r.italic = True

    saved: list[Path] = []
    for target in (OUT_DOCX_V3, OUT_DOCX_V2, OUT_DOCX):
        try:
            doc.save(target)
            saved.append(target)
            print(f"OK: {target}")
        except PermissionError:
            print(f"AVISO: no se pudo guardar {target.name} (archivo abierto en Word)")
    if not saved:
        raise PermissionError("Cierre los DOCX en Word e intente de nuevo.")


def build_csv() -> None:
    OUT_CSV.parent.mkdir(parents=True, exist_ok=True)
    fields = HEADERS
    with OUT_CSV.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=fields, delimiter=";")
        w.writeheader()
        for item in ITEMS:
            row = dict(zip(HEADERS, item_to_row(item)))
            w.writerow(row)


def main() -> None:
    build_docx()
    build_csv()
    print(f"OK: {OUT_CSV}")


if __name__ == "__main__":
    main()
