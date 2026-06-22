"""
Genera Word con enlaces PDF/DOI de los 43 artículos del estado del arte
y la justificación de inclusión en la tesis Calzatura Vilchez.
"""
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Enlaces_PDFs_Estado_Arte_43_Articulos.docx"
REFS_PATH = ROOT / "documentacion/referencias-estado-arte-43-verificadas.json"
SOURCE = ROOT / "scripts/generar_estado_arte_43_tablas.py"

# PDF directo verificado o acceso abierto conocido (complementa JSON)
PDF_OVERRIDES: dict[str, tuple[str, str]] = {
    "02": (
        "https://www.tandfonline.com/doi/pdf/10.1080/10580530.2020.1814461",
        "PDF abierto (Taylor & Francis)",
    ),
    "03": (
        "https://misq.umn.edu/misq/article-pdf/41/1/223/5892/11_si_introduction.pdf",
        "PDF abierto (MIS Quarterly — introducción)",
    ),
    "10": (
        "https://link.springer.com/content/pdf/10.1007/s11747-019-00696-0.pdf",
        "PDF Springer (puede requerir acceso)",
    ),
    "11": (
        "https://journals.sagepub.com/doi/pdf/10.1177/0008125619862257",
        "PDF Sage (puede requerir acceso)",
    ),
}

EJE_LABELS = {
    1: "Comercio electrónico y transformación comercial",
    2: "Inteligencia artificial y analítica para decisiones",
    3: "Predicción, forecasting y riesgo empresarial",
    4: "Ingeniería, arquitectura y calidad del sistema con IA",
    5: "Metodología, algoritmos y fuentes complementarias",
}


def _load_articles() -> list[dict]:
    text = SOURCE.read_text(encoding="utf-8")
    marker = "# ══════════════════════════════════════════════════════════════════════════════\n#  BUILD DOCUMENT"
    code = text.split(marker)[0]
    ns: dict = {"__file__": str(SOURCE), "__name__": "generar_estado_arte_43_tablas"}
    exec(compile(code, str(SOURCE), "exec"), ns)  # noqa: S102
    return ns["ENRICHED_ARTICLES"]


def _font(run, size: float, bold=False, italic=False, color: str | None = None):
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _add_hyperlink(paragraph, text: str, url: str, size: float = 9):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    r_pr.extend([u, color, sz])
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _resolve_links(num: str, ref: dict) -> tuple[str, str, str, str]:
    """Returns (pdf_url, pdf_label, doi_url, access_note)."""
    doi = ref.get("doi", "")
    doi_url = ref.get("url") or (f"https://doi.org/{doi}" if doi else "")
    local = ref.get("local")

    if ref.get("pdf"):
        pdf_url = ref["pdf"]
        note = "PDF directo verificado"
    elif num in PDF_OVERRIDES:
        pdf_url, note = PDF_OVERRIDES[num]
    elif doi.startswith("10.1371/"):
        pdf_url = (
            f"https://journals.plos.org/plosone/article/file?id={doi}&type=printable"
        )
        note = "PDF abierto (PLOS ONE)"
    elif doi.startswith("10.3390/"):
        m = re.search(r"10\.3390/([^\s]+)", doi)
        suffix = m.group(1) if m else doi.split("/", 1)[-1]
        pdf_url = f"https://www.mdpi.com/{suffix}/pdf"
        note = "PDF abierto (MDPI)"
    elif "springer" in ref.get("pdf", "") or doi.startswith("10.1007/"):
        pdf_url = f"https://link.springer.com/content/pdf/{doi}.pdf"
        note = "PDF Springer (verificar acceso abierto)"
    elif doi.startswith("10.1145/"):
        pdf_url = f"https://dl.acm.org/doi/pdf/{doi}"
        note = "PDF ACM DL (puede requerir acceso institucional)"
    elif doi.startswith("10.1155/"):
        pdf_url = f"https://onlinelibrary.wiley.com/doi/pdf/{doi}"
        note = "PDF Wiley (puede requerir acceso)"
    elif doi.startswith("10.1016/"):
        pdf_url = doi_url
        note = "Elsevier ScienceDirect — descargar PDF desde la página DOI (acceso institucional)"
    elif doi.startswith("10.1108/"):
        pdf_url = doi_url
        note = "Emerald — descargar PDF desde la página DOI (acceso institucional)"
    elif doi.startswith("10.1109/"):
        pdf_url = doi_url
        note = "IEEE Xplore — descargar PDF desde la página DOI (acceso institucional)"
    elif doi.startswith("10.25300/"):
        pdf_url = doi_url
        note = "MIS Quarterly — ver enlace PDF alternativo arriba si aplica"
    elif doi.startswith("10.2307/") or doi.startswith("10.1111/j.1540"):
        pdf_url = doi_url
        note = "Artículo clásico — acceso vía JSTOR/Wiley desde DOI"
    elif doi.startswith("10.1023/"):
        pdf_url = doi_url
        note = "Springer/Kluwer clásico — acceso vía DOI"
    else:
        pdf_url = doi_url
        note = "Acceso vía DOI — descargar PDF desde el editor"

    if local:
        local_path = ROOT / local
        if local_path.exists():
            note += f" | Copia local en repo: {local}"

    return pdf_url, note, doi_url, doi


def _para(doc, text, size=10, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, sa=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(text)
    _font(r, size, bold)
    return p


def build_document(articles: list[dict], refs: dict) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.70)
    sec.bottom_margin = Inches(0.70)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)

    _para(
        doc,
        "ENLACES PDF — ESTADO DEL ARTE (43 ARTÍCULOS)",
        14,
        True,
        WD_ALIGN_PARAGRAPH.CENTER,
        sa=6,
    )
    _para(
        doc,
        "Tesis: Sistema web e-commerce + IA para predicción del riesgo empresarial (IRE) — "
        "Calzatura Vilchez · Piero Vilchez · Universidad Continental · 2026",
        9,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        sa=2,
    )
    _para(
        doc,
        "Este documento lista el enlace de descarga o acceso al PDF de cada fuente, "
        "el DOI verificado y el motivo de su inclusión en el estado del arte de la tesis.",
        9,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        sa=8,
    )

    open_count = sum(
        1
        for a in articles
        if refs.get(a["num"], {}).get("pdf")
        or a["num"] in PDF_OVERRIDES
        or refs.get(a["num"], {}).get("doi", "").startswith(("10.1371/", "10.3390/"))
    )
    _para(
        doc,
        f"Resumen: 43 artículos · {open_count} con PDF directo o acceso abierto conocido · "
        f"1 PDF descargado localmente en el repositorio (artículo 43 — Dai et al. 2024). "
        f"Los demás requieren acceso institucional UCV o descarga manual desde el DOI.",
        9,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        sa=10,
    )

    # Tabla índice rápido
    idx = doc.add_table(rows=1 + len(articles), cols=5)
    idx.style = "Table Grid"
    headers = ["N.°", "Autor / título (abrev.)", "Enlace PDF", "DOI", "Eje"]
    for i, h in enumerate(headers):
        c = idx.cell(0, i)
        c.text = h
        for run in c.paragraphs[0].runs:
            _font(run, 7, bold=True)

    for row_i, art in enumerate(articles, start=1):
        num = art["num"]
        ref = refs.get(num, {})
        pdf_url, _, doi_url, _ = _resolve_links(num, ref)
        author = art["authors"].split(" et al")[0][:35]
        title_short = art["title"][:40] + ("…" if len(art["title"]) > 40 else "")

        idx.cell(row_i, 0).text = num
        idx.cell(row_i, 1).text = f"{author} — {title_short}"
        p_pdf = idx.cell(row_i, 2).paragraphs[0]
        p_pdf.text = ""
        _add_hyperlink(p_pdf, "Abrir PDF / acceso", pdf_url, 7)
        p_doi = idx.cell(row_i, 3).paragraphs[0]
        p_doi.text = ""
        if doi_url:
            _add_hyperlink(p_doi, ref.get("doi", "DOI"), doi_url, 7)
        idx.cell(row_i, 4).text = str(art.get("eje", ""))

    doc.add_page_break()

    current_eje = 0
    for art in articles:
        num = art["num"]
        ref = refs.get(num, {})
        pdf_url, pdf_note, doi_url, doi = _resolve_links(num, ref)
        eje = art.get("eje", 0)

        if eje != current_eje:
            current_eje = eje
            _para(
                doc,
                f"EJE {eje} — {EJE_LABELS.get(eje, '')}",
                11,
                True,
                sa=6,
            )

        _para(doc, f"[{num}] {art['title']}", 10, True, sa=2)
        _para(doc, f"{art['authors']} — {art['journal']}", 9, align=WD_ALIGN_PARAGRAPH.LEFT, sa=4)

        # Enlaces
        pl = doc.add_paragraph()
        pl.paragraph_format.space_after = Pt(2)
        r1 = pl.add_run("Enlace PDF: ")
        _font(r1, 9, bold=True)
        _add_hyperlink(pl, pdf_url if len(pdf_url) < 80 else "Abrir enlace PDF", pdf_url, 9)

        pl2 = doc.add_paragraph()
        pl2.paragraph_format.space_after = Pt(2)
        r2 = pl2.add_run("DOI verificado: ")
        _font(r2, 9, bold=True)
        _add_hyperlink(pl2, doi, doi_url, 9)

        _para(doc, f"Tipo de acceso: {pdf_note}", 8, sa=4)

        if ref.get("local"):
            lp = ROOT / ref["local"]
            if lp.exists():
                _para(
                    doc,
                    f"Archivo local en repositorio: {ref['local']} "
                    f"(ruta: {lp.as_posix()})",
                    8,
                    sa=4,
                )

        if ref.get("note"):
            _para(doc, f"Nota bibliográfica: {ref['note']}", 8, sa=4)

        # Por qué
        pq = doc.add_paragraph()
        pq.paragraph_format.space_before = Pt(2)
        pq.paragraph_format.space_after = Pt(6)
        rl = pq.add_run("Por qué se incluye en la tesis: ")
        _font(rl, 9, bold=True, color="1A5276")
        rt = pq.add_run(art.get("contrib", "—"))
        _font(rt, 9, italic=True)

        if art.get("evidencia"):
            pe = doc.add_paragraph()
            pe.paragraph_format.space_after = Pt(10)
            el = pe.add_run("Trazabilidad al sistema implementado: ")
            _font(el, 8, bold=True, color="117A65")
            et = pe.add_run(art["evidencia"])
            _font(et, 8)

    return doc


def main():
    articles = _load_articles()
    refs = json.loads(REFS_PATH.read_text(encoding="utf-8"))
    assert len(articles) == 43, f"Se esperaban 43 artículos, hay {len(articles)}"

    doc = build_document(articles, refs)
    doc.save(OUT)
    print(f"Documento generado: {OUT}")
    print(f"Artículos: {len(articles)}")


if __name__ == "__main__":
    main()
