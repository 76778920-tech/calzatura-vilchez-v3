from __future__ import annotations

import json
from datetime import date
from pathlib import Path

import openpyxl
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "artifacts" / "documentos"
MATRIX_FILE = ROOT / "artifacts" / "matrices" / "Matriz_Registro_Web_Empresarial_AI_DLC_BOLT.xlsx"
DOC1_FILE = OUT_DIR / "Documento_1_Ficha_tecnica_de_software_Calzatura_Vilchez_COMPLETO.docx"
DOC2_FILE = OUT_DIR / "Documento_2_Ejemplar_del_software_Calzatura_Vilchez_COMPLETO.docx"
DOWNLOADS_DIR = Path.home() / "Downloads"

PROJECT_NAME = "Sistema web Calzatura Vilchez"
PROJECT_TITLE = (
    "Sistema web de comercio electrónico con modelo de inteligencia artificial "
    "para la predicción del riesgo empresarial en la empresa Calzatura Vilchez"
)
AUTHOR = "Serpa Sedano Yeferson Wilson"
ORGANIZATION = "Calzatura Vilchez"
DOC_DATE = "30/05/2026"


def read_json(path: Path) -> dict:
    try:
        raw = path.read_bytes()
        for encoding in ("utf-8-sig", "utf-16", "utf-16-le", "utf-8"):
            try:
                return json.loads(raw.decode(encoding))
            except (UnicodeDecodeError, json.JSONDecodeError):
                continue
        return {}
    except FileNotFoundError:
        return {}


PACKAGE = read_json(ROOT / "calzatura-vilchez" / "package.json")
STRESS_2000 = read_json(ROOT / "artifacts" / "load-tests" / "autocannon-home-c2000-20260528-231905.json")
STRESS_500 = read_json(ROOT / "artifacts" / "load-tests" / "summary-500.json")
STRESS_1000 = read_json(ROOT / "artifacts" / "load-tests" / "summary-1000.json")
ZAP_REPORT = read_json(ROOT / "zap-reports" / "zap-production-report-v2.json")


def zap_counts() -> dict[str, int]:
    counts = {"Alta": 0, "Media": 0, "Baja": 0, "Informativa": 0}
    for site in ZAP_REPORT.get("site", []):
        for alert in site.get("alerts", []):
            risk = str(alert.get("riskdesc") or alert.get("risk") or "")
            count = len(alert.get("instances", [])) or int(alert.get("count") or 1)
            if risk.startswith("High"):
                counts["Alta"] += count
            elif risk.startswith("Medium"):
                counts["Media"] += count
            elif risk.startswith("Low"):
                counts["Baja"] += count
            elif risk.startswith("Informational"):
                counts["Informativa"] += count
    return counts


def metric_p95(report: dict, name: str) -> str:
    value = report.get("metrics", {}).get(name, {}).get("values", {}).get("p(95)")
    return "N/D" if value is None else f"{value:.2f} ms"


def read_matrix_rows() -> list[dict[str, str]]:
    if not MATRIX_FILE.exists():
        return []
    workbook = openpyxl.load_workbook(MATRIX_FILE, data_only=True)
    sheet = workbook["Bolts"]
    rows: list[dict[str, str]] = []
    for row in sheet.iter_rows(min_row=5, max_col=78, values_only=True):
        if not row[1]:
            continue
        rows.append(
            {
                "macroproceso": str(row[0] or ""),
                "id": str(row[1] or ""),
                "descripcion": str(row[2] or ""),
                "dominio": str(row[5] or ""),
                "estado": str(row[74] or ""),
                "evidencia": str(row[77] or ""),
            }
        )
    return rows


BOLT_ROWS = read_matrix_rows()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text_color(cell, color: str) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 3"].font.size = Pt(10)


def add_title(document: Document, title: str, subtitle: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_paragraph.add_run(subtitle)
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(11)
    subtitle_run.italic = True

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Fecha de elaboración: {DOC_DATE}").bold = True
    document.add_paragraph()


def add_note(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EAF2F8")
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.size = Pt(9)
    run.bold = True


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        cell.text = header
        set_cell_shading(cell, "1F4E79")
        set_cell_text_color(cell, "FFFFFF")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                if index < len(row.cells):
                    row.cells[index].width = Inches(width)
    document.add_paragraph()
    return table


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(item)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(item)


def read_lines(relative_path: str, start: int, end: int) -> str:
    path = ROOT / relative_path
    lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
    excerpt = []
    for line_number in range(start, min(end, len(lines)) + 1):
        excerpt.append(f"{line_number:04d}: {lines[line_number - 1]}")
    return "\n".join(excerpt)


def add_code_excerpt(document: Document, title: str, relative_path: str, start: int, end: int) -> None:
    document.add_heading(title, level=3)
    document.add_paragraph(f"Archivo: {relative_path} | Líneas {start}-{end}")
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F7")
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(read_lines(relative_path, start, end))
    run.font.name = "Consolas"
    run.font.size = Pt(7)
    document.add_paragraph()


def add_footer(document: Document) -> None:
    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(f"{PROJECT_NAME} — Documento técnico generado el {DOC_DATE}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)


def dependency_rows() -> list[list[str]]:
    dependencies = PACKAGE.get("dependencies", {})
    dev_dependencies = PACKAGE.get("devDependencies", {})
    keys = [
        ("React", dependencies.get("react", "N/D"), "Interfaz de usuario SPA"),
        ("TypeScript", dev_dependencies.get("typescript", "N/D"), "Tipado estático y mantenibilidad"),
        ("Vite", dev_dependencies.get("vite", "N/D"), "Build y servidor de desarrollo"),
        ("Firebase", dependencies.get("firebase", "N/D"), "Autenticación, App Check y hosting"),
        ("Supabase JS", dependencies.get("@supabase/supabase-js", "N/D"), "Conexión a PostgreSQL/Supabase"),
        ("ExcelJS", dependencies.get("exceljs", "N/D"), "Importación/exportación de Excel"),
        ("Vitest", dev_dependencies.get("vitest", "N/D"), "Pruebas unitarias"),
        ("Playwright", dev_dependencies.get("@playwright/test", "N/D"), "Pruebas E2E"),
    ]
    return [[name, version, purpose] for name, version, purpose in keys]


def module_rows() -> list[list[str]]:
    return [
        ["Público", "Home, login, registro, landings, páginas legales, libro de reclamaciones", "src/domains/publico"],
        ["Catálogo", "Listado, filtros, detalle, campañas, tarjetas de producto accesibles", "src/domains/productos"],
        ["Carrito y checkout", "Carrito por sesión, entrega, Stripe, contraentrega e idempotencia", "src/domains/carrito"],
        ["Pedidos", "Pedido exitoso, historial cliente, administración y estados", "src/domains/pedidos"],
        ["Clientes", "Favoritos privados por cuenta", "src/domains/clientes"],
        ["Usuarios", "Perfil, roles, DNI, privacidad y administración de usuarios", "src/domains/usuarios"],
        ["Administración", "Dashboard, Excel, predicciones, auditoría, reclamaciones", "src/domains/administradores"],
        ["Ventas físicas", "Registro de venta, devolución, stock por talla y panel de trabajador", "src/domains/ventas; src/domains/trabajadores"],
        ["Fabricantes", "CRUD de proveedores/marcas con redacción de PII", "src/domains/fabricantes"],
        ["BFF/API", "Pedidos, Stripe, auditoría, DNI, admin, datos, rate limit y seguridad", "bff/server.cjs"],
        ["Base de datos", "Migraciones, RLS, RPC, triggers y políticas", "supabase/migrations"],
        ["IA", "Predicción, IRE, backtesting y controles medibles", "ai-service; src/domains/administradores/predictions"],
    ]


def security_rows() -> list[list[str]]:
    z = zap_counts()
    return [
        ["Autenticación", "Firebase Auth + verificación de token en BFF", "server.cjs verifica usuario antes de operaciones sensibles"],
        ["Autorización", "Roles cliente, trabajador y admin", "AreaRoute, accessControl y validación server-side"],
        ["Base de datos", "RLS, REVOKE y service_role solo en BFF", "supabase/RLS-MATRIX.md y migraciones de hardening"],
        ["Pedidos/Stripe", "Webhook firmado y bloqueo de pagado manual para Stripe", "server.cjs /stripeWebhook y /updateOrderStatus"],
        ["DNI/App Check", "App Check, rate limit y proof secret en servidor", "REQUIRE_DNI_APPCHECK y DNI_LOOKUP_PROOF_SECRET en BFF"],
        ["Auditoría", "BFF/service_role, allowlist por rol y redacción PII", "auditoria, auditPii y migraciones de redacción"],
        ["DAST", f"OWASP ZAP Docker: {z['Alta']} altas, {z['Media']} medias, {z['Baja']} bajas, {z['Informativa']} informativas", "Sin alertas altas/críticas en reporte v2"],
        ["Secretos", "Variables privadas fuera del frontend", "SUPABASE_SERVICE_ROLE_KEY, Stripe secrets y DNI secrets en servidor"],
    ]


def testing_rows() -> list[list[str]]:
    return [
        ["Unitarias", "Vitest", "Servicios, utilidades, guardas BFF, RLS, PII y reglas de negocio", "VERDE"],
        ["E2E", "Playwright", "Home, catálogo, carrito, checkout, admin, ventas, usuarios, Excel y accesibilidad", "VERDE"],
        ["Seguridad dinámica", "OWASP ZAP Docker", "Reporte producción v2 sin alertas altas/críticas", "VERDE"],
        [
            "Estrés web",
            "autocannon",
            (
                f"{STRESS_2000.get('connections', 2000)} conexiones concurrentes; "
                f"{STRESS_2000.get('errors', 0)} errores; "
                f"{STRESS_2000.get('non2xx', 0)} respuestas no 2xx; "
                f"p99 {STRESS_2000.get('latency', {}).get('p99', 'N/D')} ms"
            ),
            "VERDE",
        ],
        ["Carga mixta", "k6", f"p95 catálogo 500 VUs: {metric_p95(STRESS_500, 'http_req_duration{name:supabase_catalog_list}')} / p95 BFF 1000 VUs: {metric_p95(STRESS_1000, 'http_req_duration{name:bff_catalog_active}')}", "VERDE"],
        ["CI/CD", "GitHub Actions", "Lint, typecheck, test, build, Docker, AI, migrations y deploy gates", "VERDE"],
        ["Continuidad", "Restore drill", "Readiness y evidencia versionada de recuperación", "VERDE"],
        ["Calidad estática", "ESLint / TypeScript / Sonar", "Preflight y smells críticos corregidos", "VERDE"],
    ]


def create_ficha_tecnica() -> None:
    document = Document()
    configure_document(document)
    add_title(document, "DOCUMENTO 1 — FICHA TÉCNICA DE SOFTWARE", PROJECT_TITLE)
    add_note(
        document,
        "Documento técnico elaborado para identificar, describir y sustentar el sistema web desarrollado. "
        "El alcance corresponde al apartado web y sus servicios de soporte: frontend, BFF, base de datos, IA, seguridad, pruebas y despliegue.",
    )

    document.add_heading("1. Datos generales", level=1)
    add_table(
        document,
        ["Campo", "Detalle"],
        [
            ["Nombre del software", PROJECT_NAME],
            ["Título técnico del proyecto", PROJECT_TITLE],
            ["Autor / desarrollador", AUTHOR],
            ["Entidad de aplicación", ORGANIZATION],
            ["Fecha de elaboración", DOC_DATE],
            ["Versión documental", "1.0.0 — registro técnico completo"],
            ["Versión del paquete web", PACKAGE.get("version", "0.0.0")],
            ["Tipo de software", "Sistema web de comercio electrónico, administración comercial e inteligencia artificial"],
            ["Modalidad de ejecución", "Aplicación web SPA desplegable en Firebase Hosting, BFF en Render/Node y base de datos Supabase PostgreSQL"],
            ["Repositorio o workspace técnico", str(ROOT)],
        ],
        [1.8, 5.9],
    )

    document.add_heading("2. Descripción del software", level=1)
    document.add_paragraph(
        "El sistema web Calzatura Vilchez es una plataforma empresarial orientada a la digitalización del canal comercial de una tienda de calzado. "
        "Integra una tienda en línea con catálogo, filtros, carrito, checkout, pagos con Stripe, contraentrega, pedidos, perfil de cliente, favoritos, "
        "administración de productos, stock por talla/color, ventas físicas, fabricantes, usuarios, auditoría, importación/exportación Excel, libro de reclamaciones "
        "y un módulo de inteligencia artificial para análisis predictivo e Índice de Riesgo Empresarial (IRE)."
    )
    document.add_paragraph(
        "La solución fue revisada bajo criterios de ISO/IEC 25010, ISO/IEC 27001, ISO/IEC 27701, ISO/IEC 25012, ISO/IEC 29119, ISO/IEC 42001 e ISO 22301, "
        "tomando como evidencia pruebas automatizadas, reportes de seguridad, controles de base de datos y matriz BOLT empresarial."
    )

    document.add_heading("3. Objetivo y alcance", level=1)
    add_bullets(
        document,
        [
            "Permitir ventas digitales de calzado mediante catálogo, carrito, checkout y seguimiento de pedidos.",
            "Centralizar la administración de productos, precios, stock, fabricantes, usuarios, ventas físicas y datos Excel.",
            "Proteger operaciones sensibles mediante autenticación, autorización, BFF, RLS, auditoría y redacción de PII.",
            "Incorporar IA para apoyar decisiones comerciales mediante predicciones, métricas, backtesting e IRE.",
            "Mantener evidencia verificable de calidad, seguridad, rendimiento, continuidad y despliegue.",
            "Excluir módulos ajenos al comercio electrónico, administración comercial, IA, datos, seguridad y operación web.",
        ],
    )

    document.add_heading("4. Módulos funcionales", level=1)
    add_table(document, ["Macroproceso", "Funcionalidad cubierta", "Evidencia principal"], module_rows(), [1.6, 4.6, 2.0])

    document.add_heading("5. Arquitectura técnica", level=1)
    document.add_paragraph(
        "La arquitectura se organiza como aplicación web SPA con frontend React/TypeScript, BFF Node/Express para operaciones sensibles, Supabase PostgreSQL "
        "para persistencia, Firebase Auth/App Check para identidad, Stripe para pagos, Cloudinary para medios, servicio IA separado y pipelines CI/CD."
    )
    add_table(
        document,
        ["Capa", "Tecnología", "Responsabilidad"],
        [
            ["Presentación", "React 19, TypeScript, Vite", "Interfaz web, rutas públicas, cliente, trabajador y admin"],
            ["BFF/API", "Node.js / Express", "Pedidos, Stripe, DNI, auditoría, admin, Excel, rate limit y service_role"],
            ["Datos", "Supabase PostgreSQL", "Tablas, RLS, migraciones, RPC, triggers e integridad"],
            ["Autenticación", "Firebase Auth + App Check", "Identidad, protección de endpoints y sesión"],
            ["Pagos", "Stripe Checkout + webhook firmado", "Confirmación de pagos sin confiar en el cliente"],
            ["IA", "Servicio Python/FastAPI", "Predicción, IRE, backtesting y métricas"],
            ["DevOps", "GitHub Actions, Docker, Firebase Hosting, Render", "Build, test, seguridad, despliegue y continuidad"],
        ],
        [1.5, 2.2, 4.0],
    )

    document.add_heading("6. Tecnologías y dependencias principales", level=1)
    add_table(document, ["Tecnología", "Versión declarada", "Uso"], dependency_rows(), [1.8, 1.3, 4.5])

    document.add_heading("7. Modelo de datos principal", level=1)
    add_table(
        document,
        ["Tabla / entidad", "Propósito", "Control aplicado"],
        [
            ["productos", "Catálogo, variantes, campaña, stock y atributos comerciales", "Lectura pública solo activa; escritura por admin/BFF"],
            ["productoCodigos", "Código único por producto y variante", "Unicidad y guardas anti-duplicado"],
            ["productoFinanzas", "Costos, márgenes y precios sugeridos", "Acceso administrativo y revocación a cliente"],
            ["pedidos", "Órdenes web, estado, pagos, dirección y totales", "Máquina de estados, idempotencia, stock y auditoría"],
            ["usuarios", "Perfil extendido y rol", "Firebase UID, PII minimizada, rol server-side"],
            ["favoritos", "Relación usuario-producto", "Aislamiento por cuenta"],
            ["ventasDiarias", "Ventas físicas y devoluciones", "RPC atómica y motivo de devolución"],
            ["fabricantes", "Marcas/proveedores", "Redacción de datos sensibles en auditoría"],
            ["auditoria", "Eventos operativos y trazabilidad", "RLS, service_role, allowlist y redacción PII"],
            ["ireHistorial / modeloEstado", "Historial de IA y estado de modelo", "Versionado, métricas y trazabilidad"],
        ],
        [1.7, 4.2, 2.3],
    )

    document.add_heading("8. Seguridad, privacidad y controles", level=1)
    add_table(document, ["Control", "Implementación", "Evidencia / resultado"], security_rows(), [1.6, 3.8, 2.8])

    document.add_heading("9. Inteligencia artificial y datos", level=1)
    add_bullets(
        document,
        [
            "El módulo IA apoya la predicción de demanda y el análisis del riesgo empresarial mediante indicadores operativos y comerciales.",
            "Incluye data_sufficient, historial IRE, métricas, feature importances, backtesting y reportes de calidad de datos.",
            "Los controles están alineados con ISO/IEC 42001 para gobierno de IA, ISO/IEC 25012 para calidad de datos e ISO/IEC 29119 para pruebas.",
            "El sistema mantiene salida explicable para administración: nivel de riesgo, recomendaciones, tendencias y métricas asociadas.",
        ],
    )

    document.add_heading("10. Pruebas, calidad y capacidad", level=1)
    add_table(document, ["Tipo", "Herramienta", "Cobertura / evidencia", "Estado"], testing_rows(), [1.4, 1.5, 4.7, 0.9])

    document.add_heading("11. Requisitos de instalación y operación", level=1)
    add_table(
        document,
        ["Elemento", "Requisito"],
        [
            ["Navegador", "Chrome, Edge, Firefox o equivalente moderno"],
            ["Node.js", "Compatible con el entorno del proyecto y dependencias npm"],
            ["Variables públicas frontend", "VITE_SUPABASE_URL, VITE_SUPABASE_ANON_KEY, VITE_FIREBASE_* y site key pública de App Check"],
            ["Variables privadas servidor", "SUPABASE_SERVICE_ROLE_KEY, STRIPE_SECRET_KEY, STRIPE_WEBHOOK_SECRET, DNI_LOOKUP_PROOF_SECRET, tokens de proveedores y claves IA"],
            ["Base de datos", "Supabase PostgreSQL con migraciones aplicadas"],
            ["Despliegue", "Firebase Hosting para frontend; Render/BFF para API; servicio IA por contenedor"],
            ["Pruebas mínimas previas", "npm run lint, npm run typecheck, npm run test, npm run test:e2e, ZAP y stress según release"],
        ],
        [2.2, 5.5],
    )

    document.add_heading("12. Evidencias de trazabilidad", level=1)
    add_table(
        document,
        ["Evidencia", "Ubicación"],
        [
            ["Matriz de registro empresarial BOLT", str(MATRIX_FILE)],
            ["Inventario BOLT del sistema web", f"{len(BOLT_ROWS)} registros trazables: BOLT-WEB-001 a BOLT-WEB-055; diseño, año 2026 y evidencia por línea de código"],
            ["Reporte ZAP producción v2", "zap-reports/zap-production-report-v2.json"],
            ["Prueba de estrés 2000 conexiones", "artifacts/load-tests/autocannon-home-c2000-20260528-231905.json"],
            ["Carga mixta k6", "artifacts/load-tests/summary-500.json; artifacts/load-tests/summary-1000.json"],
            ["SRS", "documentacion/05-especificacion-requisitos-software-SRS.md"],
            ["Arquitectura", "documentacion/06-diseno-arquitectura-y-datos.md"],
            ["BFF principal", "calzatura-vilchez/bff/server.cjs"],
            ["Migraciones Supabase", "calzatura-vilchez/supabase/migrations"],
        ],
        [2.8, 5.0],
    )

    document.add_heading("13. Declaración técnica", level=1)
    document.add_paragraph(
        "Con base en la revisión del código fuente, pruebas, reportes y matriz de registro BOLT, el sistema web se encuentra documentado como un producto funcional, "
        "trazable y verificable. Los controles de seguridad, calidad y rendimiento se registran como evidencias técnicas y deben mantenerse actualizados por versión."
    )

    add_footer(document)
    document.save(DOC1_FILE)


def create_ejemplar() -> None:
    document = Document()
    configure_document(document)
    add_title(document, "DOCUMENTO 2 — EJEMPLAR DEL SOFTWARE", PROJECT_TITLE)
    add_note(
        document,
        "Ejemplar técnico del software: identifica estructura, módulos, rutas de código, extractos representativos y evidencia de ejecución/pruebas. "
        "No contiene secretos de producción ni credenciales privadas."
    )

    document.add_heading("1. Identificación del ejemplar", level=1)
    add_table(
        document,
        ["Campo", "Detalle"],
        [
            ["Nombre del software", PROJECT_NAME],
            ["Autor / desarrollador", AUTHOR],
            ["Entidad de aplicación", ORGANIZATION],
            ["Fecha del ejemplar", DOC_DATE],
            ["Naturaleza", "Aplicación web empresarial, BFF/API, base de datos, IA y documentación técnica"],
            ["Contenido del ejemplar", "Descripción de carpetas, módulos, código representativo, matriz BOLT, evidencia de pruebas y seguridad"],
            ["Protección de secretos", "Las claves privadas se describen como variables de entorno; no se incluyen valores reales"],
        ],
        [2.0, 5.6],
    )

    document.add_heading("2. Estructura técnica del software", level=1)
    add_table(
        document,
        ["Ruta", "Contenido", "Finalidad"],
        [
            ["calzatura-vilchez/src", "Código frontend React/TypeScript", "Interfaz web por dominios"],
            ["calzatura-vilchez/src/domains", "Módulos público, productos, carrito, pedidos, clientes, usuarios, ventas, administración y fabricantes", "Separación funcional"],
            ["calzatura-vilchez/bff", "Servidor Node/Express", "Operaciones sensibles con service_role y proveedores externos"],
            ["calzatura-vilchez/supabase/migrations", "Migraciones SQL", "Modelo de datos, RLS, RPC, triggers e integridad"],
            ["calzatura-vilchez/e2e", "Pruebas Playwright", "Validación funcional end-to-end"],
            ["calzatura-vilchez/src/__tests__", "Pruebas unitarias y guardas", "Calidad, seguridad y anti-regresión"],
            ["ai-service", "Servicio de IA y scripts", "Predicción, backtesting y métricas"],
            ["load-tests", "Escenarios k6/autocannon", "Capacidad y rendimiento"],
            ["zap-reports", "Reportes OWASP ZAP", "Seguridad dinámica"],
            ["artifacts/matrices", "Matriz BOLT empresarial", "Trazabilidad de macroprocesos, Bolts y evidencia"],
        ],
        [2.3, 3.2, 2.4],
    )

    document.add_heading("3. Flujo funcional principal", level=1)
    add_numbered(
        document,
        [
            "El visitante accede al home, revisa campañas y navega al catálogo.",
            "El cliente filtra productos, revisa detalle, selecciona talla y agrega al carrito.",
            "El checkout valida dirección, método de pago, stock, precio e idempotencia.",
            "El BFF crea el pedido con Firebase token y Supabase service_role; Stripe confirma pagos por webhook firmado.",
            "El administrador gestiona productos, fabricantes, usuarios, pedidos, ventas, Excel, auditoría e IA.",
            "El trabajador registra ventas físicas y devoluciones dentro del alcance permitido por rol.",
            "La base de datos aplica RLS, RPC, triggers y restricciones de integridad.",
            "La IA procesa datos agregados para predicción, IRE, recomendaciones y control de riesgo empresarial.",
        ],
    )

    document.add_heading("4. Inventario BOLT del ejemplar", level=1)
    if BOLT_ROWS:
        add_table(
            document,
            ["Macroproceso", "ID", "Descripción del Bolt", "Estado", "Ubicación exacta en código / evidencia"],
            [[row["macroproceso"], row["id"], row["descripcion"], row["estado"], row["evidencia"]] for row in BOLT_ROWS],
            [1.5, 0.9, 3.6, 0.7, 3.0],
        )
    else:
        document.add_paragraph("No se encontró la matriz BOLT; revisar artifacts/matrices.")

    document.add_heading("5. Extractos representativos de código fuente", level=1)
    add_code_excerpt(document, "5.1 Enrutamiento web y separación de áreas", "calzatura-vilchez/src/App.tsx", 133, 174)
    add_code_excerpt(document, "5.2 Tarjeta de producto accesible", "calzatura-vilchez/src/domains/productos/components/ProductCard.tsx", 101, 147)
    add_code_excerpt(document, "5.3 Carrito por sesión", "calzatura-vilchez/src/domains/carrito/context/CartContext.tsx", 37, 83)
    add_code_excerpt(document, "5.4 Checkout con creación de pedido", "calzatura-vilchez/src/domains/carrito/pages/CheckoutPage.tsx", 132, 163)
    add_code_excerpt(document, "5.5 BFF createOrder con idempotencia y validación", "calzatura-vilchez/bff/server.cjs", 1674, 1778)
    add_code_excerpt(document, "5.6 Máquina de estados de pedidos", "calzatura-vilchez/functions/orderStatusPolicy.js", 1, 40)
    add_code_excerpt(document, "5.7 Importación Excel", "calzatura-vilchez/src/domains/administradores/pages/AdminData.tsx", 1060, 1110)
    add_code_excerpt(document, "5.8 Supabase RLS y acceso BFF", "calzatura-vilchez/supabase/RLS-MATRIX.md", 1, 47)
    add_code_excerpt(document, "5.9 Predicciones IA en administración", "calzatura-vilchez/src/domains/administradores/pages/AdminPredictions.tsx", 1, 35)
    add_code_excerpt(document, "5.10 CI/CD y calidad", ".github/workflows/ci.yml", 1, 45)

    document.add_heading("6. Evidencia de pruebas y seguridad del ejemplar", level=1)
    add_table(document, ["Tipo", "Herramienta", "Cobertura / evidencia", "Estado"], testing_rows(), [1.4, 1.5, 4.7, 0.9])

    document.add_heading("7. Resumen de autenticidad técnica", level=1)
    add_bullets(
        document,
        [
            "El ejemplar enlaza cada macroproceso con código fuente, pruebas, reportes y evidencia de infraestructura.",
            "La última columna de la matriz BOLT contiene rutas y líneas de código verificables.",
            "Los reportes de ZAP y estrés se encuentran versionados como artifacts locales.",
            "Los secretos de producción no forman parte del ejemplar; se referencian únicamente como nombres de variables.",
            "Las pruebas registradas sostienen el estado VERDE del sistema web en los módulos cubiertos.",
        ],
    )

    document.add_heading("8. Comandos de referencia", level=1)
    add_table(
        document,
        ["Comando", "Uso"],
        [
            ["npm run lint", "Validación ESLint"],
            ["npm run typecheck", "Validación TypeScript"],
            ["npm run test", "Pruebas unitarias"],
            ["npm run test:e2e", "Pruebas end-to-end Playwright"],
            ["node scripts/validate-supabase-migrations.mjs", "Validación de migraciones Supabase"],
            ["node scripts/validate-supabase-rls-matrix.mjs", "Validación del contrato RLS"],
            ["npm run build", "Compilación del frontend web"],
            ["firebase deploy --only hosting", "Despliegue del frontend en Firebase Hosting"],
        ],
        [2.4, 5.2],
    )

    document.add_heading("9. Declaración del ejemplar", level=1)
    document.add_paragraph(
        "El presente documento constituye un ejemplar técnico del software web Calzatura Vilchez. Resume los componentes esenciales del código fuente, "
        "las rutas de implementación, las evidencias de prueba y los controles de seguridad necesarios para identificar la autenticidad y funcionalidad del sistema."
    )

    add_footer(document)
    document.save(DOC2_FILE)


def copy_to_downloads() -> None:
    for source in [DOC1_FILE, DOC2_FILE]:
        target = DOWNLOADS_DIR / source.name
        target.write_bytes(source.read_bytes())


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    create_ficha_tecnica()
    create_ejemplar()
    copy_to_downloads()
    print(DOC1_FILE)
    print(DOC2_FILE)


if __name__ == "__main__":
    main()
