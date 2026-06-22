#!/usr/bin/env python3
"""
Genera la Matriz de Consistencia alineada con:
- Repo Calzatura Vilchez (stack real, IRE 0-100, RandomForestRegressor)
- Matriz de operacionalización de variables (42 artículos Q1)
- Instrumento de investigación (Likert pre/post)
- Guías URP Anexo 6 / UNCP / coherencia vertical (Carrasco, 2018)

Uso: python scripts/generar_matriz_consistencia_mejorada.py
"""
from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "Matriz_Consistencia_Calzatura_Vilchez_MEJORADA.docx"
OUT_DOCX_V2 = ROOT / "Matriz_Consistencia_Calzatura_Vilchez_v2_VALIDADA.docx"
OUT_CSV = ROOT / "documentacion/cuadros-excel/CU-T09-matriz-consistencia.csv"

TITULO = (
    "Sistema web de comercio electrónico con modelo de Inteligencia Artificial "
    "para la predicción del riesgo empresarial en la empresa Calzatura Vilchez, Huancayo, 2026"
)

HEADERS = [
    "Problemas",
    "Objetivos",
    "Hipótesis",
    "Variables",
    "Dimensiones",
    "Indicadores",
    "Metodología",
]

# Coherencia vertical: cada fila comparte variables/indicadores entre P-O-H.
ROWS: list[dict[str, str]] = [
    {
        "nivel": "General",
        "problemas": (
            "Problema general: ¿De qué manera el sistema web de comercio electrónico con modelo "
            "de Inteligencia Artificial influye en la predicción del riesgo empresarial "
            "comercial-operativo en la empresa Calzatura Vilchez, Huancayo, 2026?"
        ),
        "objetivos": (
            "Objetivo general: Implementar un sistema web de comercio electrónico con modelo "
            "de Inteligencia Artificial para mejorar la predicción del riesgo empresarial "
            "comercial-operativo en la empresa Calzatura Vilchez, Huancayo, 2026."
        ),
        "hipotesis": (
            "Hipótesis general (H0): El sistema web de comercio electrónico con modelo de "
            "Inteligencia Artificial no mejora la predicción del riesgo empresarial en "
            "Calzatura Vilchez.\n\n"
            "Hipótesis alterna (H1): El sistema mejora significativamente la predicción del "
            "riesgo empresarial comercial-operativo medida mediante el IRE (α = 0,05)."
        ),
        "variables": (
            "VI: Sistema web de comercio electrónico con modelo de Inteligencia Artificial "
            "(React/Vite, Supabase, Firebase Auth, BFF Node, microservicio IA FastAPI).\n\n"
            "VD: Predicción del riesgo empresarial mediante el Índice de Riesgo Empresarial (IRE), "
            "escala 0–100."
        ),
        "dimensiones": (
            "VI — D1 e-commerce y transformación digital; D2 IA y analítica; "
            "D3 arquitectura y despliegue; D4 metodología ágil/híbrida.\n\n"
            "VD — D1 modelos de predicción de riesgo; D2 exactitud/validación del IRE; "
            "D3 componentes del IRE; D4 impacto en gestión operativa."
        ),
        "indicadores": (
            "Digitalización (% procesos), MAE/RMSE/sMAPE demanda, latencia API, disponibilidad, "
            "cobertura pruebas, valor IRE (0–100), riesgo_stock/ingresos/demanda, "
            "quiebres de stock, percepción Likert pre/post, adopción dashboard."
        ),
        "metodologia": (
            "Método: científico y tecnológico-sistémico.\n"
            "Tipo: aplicada.\n"
            "Enfoque: mixto (cuantitativo predominante + cualitativo complementario).\n"
            "Nivel: explicativo-causal.\n"
            "Diseño: preexperimental — un solo grupo, medición antes-después (O₁ → X → O₂).\n"
            "Población: personal directivo/operativo de Calzatura Vilchez y registros "
            "transaccionales en Supabase.\n"
            "Muestra: censo del personal (n según plantilla) y registros enero 2023–diciembre 2025.\n"
            "Técnicas: entrevista semiestructurada, encuesta, análisis documental, observación, "
            "evaluación de modelos ML.\n"
            "Instrumentos: guía de entrevista, cuestionario Likert (24 ítems), ficha documental, "
            "pipeline ai-service/evaluate.py, dashboard IRE."
        ),
    },
    {
        "nivel": "Específico 1",
        "problemas": (
            "Problema específico 1: ¿En qué medida la implementación del sistema web de comercio "
            "electrónico mejora la digitalización de los procesos de venta y gestión de inventario "
            "en Calzatura Vilchez?"
        ),
        "objetivos": (
            "Objetivo específico 1: Digitalizar los procesos de venta y gestión de inventario "
            "mediante el desarrollo e implementación del sistema web (React, TypeScript, Vite) "
            "con persistencia en Supabase (PostgreSQL) y autenticación Firebase."
        ),
        "hipotesis": (
            "Hipótesis específica 1 (H0): La implementación del sistema no incrementa la tasa de "
            "digitalización de procesos de venta e inventario.\n\n"
            "Hipótesis específica 1 (H1): La implementación incrementa significativamente la tasa de "
            "digitalización (≥ 70 % respecto a la línea base O₁)."
        ),
        "variables": "VI: Sistema web de comercio electrónico con modelo de IA.",
        "dimensiones": "VI-D1: e-commerce y transformación digital.\nVI-D3: arquitectura de software y despliegue.",
        "indicadores": (
            "1. Nivel de madurez digital (ordinal 1–5).\n"
            "2. Tasa de digitalización de procesos de venta e inventario (%).\n"
            "3. Número de canales de venta digitales activos.\n"
            "4. Tasa de conversión digital (%).\n"
            "5. Tiempo de ciclo de procesamiento de pedidos (horas/min).\n"
            "6. Disponibilidad del sistema web (%) — indicador VI-D3."
        ),
        "metodologia": (
            "Técnicas: análisis documental, observación sistemática, entrevista semiestructurada "
            "a empleadores/directivos y encuesta Likert (dimensión digitalización).\n"
            "Instrumentos: ficha documental, guía de entrevista, cuestionario Likert, logs Supabase.\n"
            "Análisis: comparación O₁–O₂ de indicadores; triangulación cualitativa (entrevistas) "
            "y cuantitativa (encuesta + registros)."
        ),
    },
    {
        "nivel": "Específico 2",
        "problemas": (
            "Problema específico 2: ¿Cómo el modelo de Inteligencia Artificial basado en "
            "RandomForestRegressor mejora la precisión del pronóstico de demanda de calzado "
            "en Calzatura Vilchez, medida mediante MAE y RMSE?"
        ),
        "objetivos": (
            "Objetivo específico 2: Desarrollar e implementar un modelo de Machine Learning "
            "(RandomForestRegressor en ai-service; fallback promedio móvil si datos insuficientes) "
            "para pronóstico de demanda con horizonte configurable 7, 15 o 30 días (UI admin), "
            "usando ventas históricas en Supabase."
        ),
        "hipotesis": (
            "Hipótesis específica 2 (H0): El RandomForestRegressor no reduce el MAE respecto al baseline.\n\n"
            "Hipótesis específica 2 (H1): El RandomForestRegressor reduce el MAE del pronóstico "
            "de demanda en al menos 30 % respecto al método empírico baseline (evaluate.py)."
        ),
        "variables": (
            "VI: Sistema web con modelo de IA.\n"
            "VD: Componente riesgo_demanda del IRE y exactitud del pronóstico."
        ),
        "dimensiones": (
            "VI-D2: Inteligencia Artificial y analítica de datos.\n"
            "VD-D2: exactitud y validación del modelo predictivo.\n"
            "VD-D3: componente riesgo_demanda del IRE."
        ),
        "indicadores": (
            "1. MAE del pronóstico (unidades).\n"
            "2. RMSE del pronóstico (unidades).\n"
            "3. sMAPE (%).\n"
            "4. R² del modelo.\n"
            "5. Horizonte de predicción (7, 15 o 30 días — HORIZON_OPTIONS).\n"
            "6. Estabilidad ante data drift (% variación mensual)."
        ),
        "metodologia": (
            "Técnicas: análisis de datos históricos, validación walk-forward, evaluación comparativa ML.\n"
            "Instrumentos: dataset ventasDiarias/pedidos (Supabase), ai-service/evaluate.py, "
            "reporte métricas scikit-learn.\n"
            "Análisis: comparación RF vs baseline; MAE, RMSE, sMAPE, R²; prueba t pareada o "
            "Wilcoxon según normalidad (Shapiro-Wilk, α = 0,05)."
        ),
    },
    {
        "nivel": "Específico 3",
        "problemas": (
            "Problema específico 3: ¿En qué grado el Índice de Riesgo Empresarial (IRE) "
            "calculado por el sistema permite anticipar situaciones de alerta y reducir el "
            "riesgo operativo en Calzatura Vilchez?"
        ),
        "objetivos": (
            "Objetivo específico 3: Calcular y monitorear automáticamente el IRE "
            "(IRE = 0,40·riesgo_stock + 0,35·riesgo_ingresos + 0,25·riesgo_demanda, escala 0–100) "
            "para alertar al equipo de gestión sobre riesgo comercial-operativo."
        ),
        "hipotesis": (
            "Hipótesis específica 3 (H0): El IRE no reduce el riesgo operativo ni los quiebres de stock.\n\n"
            "Hipótesis específica 3 (H1): El IRE reduce significativamente el riesgo operativo, "
            "disminuyendo los quiebres de stock en al menos 40 % en O₂ respecto a O₁."
        ),
        "variables": "VD: Predicción del riesgo empresarial mediante el IRE.",
        "dimensiones": (
            "VD-D1: modelos de predicción del riesgo con ML.\n"
            "VD-D2: exactitud y validación del IRE (consistencia interna; AUC solo si hay eventos etiquetados).\n"
            "VD-D3: componentes riesgo_stock, riesgo_ingresos, riesgo_demanda.\n"
            "VD-D4: impacto en gestión operativa."
        ),
        "indicadores": (
            "1. Valor compuesto IRE (0–100).\n"
            "2. Subíndices riesgo_stock, riesgo_ingresos, riesgo_demanda.\n"
            "3. Alertas IRE ≥ 51 (alto) o ≥ 76 (crítico).\n"
            "4. Número de quiebres de stock (conteo pre/post).\n"
            "5. Tasa de alertas atendidas (%).\n"
            "6. Percepción utilidad IRE — ítems 1–24 cuestionario Likert (VD-D1…D4)."
        ),
        "metodologia": (
            "Técnicas: cálculo automatizado IRE, análisis documental, observación, entrevista "
            "semiestructurada a directivos, encuesta Likert pre/post, validación de consistencia interna.\n"
            "Instrumentos: dashboard AdminPredictions, tabla ireHistorial (Supabase), cuestionario "
            "Likert, guía de entrevista, tests ai-service (invariantes IRE).\n"
            "Análisis: comparación O₁–O₂ quiebres y IRE medio; t pareada/Wilcoxon; triangulación "
            "con entrevistas. Nota: AUC-ROC condicionado a etiquetado histórico de eventos de crisis."
        ),
    },
]


def set_cell_text(cell, text: str, bold: bool = False, font_size: float = 7.0) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_widths(table, widths: list[float]) -> None:
    for row in table.rows:
        for index, width in enumerate(widths):
            if index < len(row.cells):
                row.cells[index].width = Inches(width)


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MATRIZ DE CONSISTENCIA")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(12)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(TITULO)
    run.font.name = "Arial"
    run.font.size = Pt(8)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = note.add_run(
        "Cuadro de coherencia vertical (problema ↔ objetivo ↔ hipótesis ↔ variables ↔ "
        "indicadores ↔ metodología). Fuente trazable: CU-T09 · Matriz de Operacionalización "
        "de Variables · Instrumento de Investigación (24 ítems Likert) · "
        "documentacion/07-modulo-ia-riesgo-empresarial.md · 09-implementacion-despliegue-ci.md. "
        "Validación automática: scripts/validar_matriz_consistencia.py. "
        "Referencia: URP Anexo 6 · UNCP Guía plan y tesis · Carrasco (2018)."
    )
    run.font.name = "Arial"
    run.font.size = Pt(7)
    run.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = meta.add_run(
        "Tesista: _________________________   Asesor: Dr. Maglioni Arana Caparachin   "
        "Versión: 2.0   Fecha: junio 2026"
    )
    run.font.name = "Arial"
    run.font.size = Pt(7)

    table = doc.add_table(rows=len(ROWS) + 1, cols=len(HEADERS))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_widths(table, [1.42, 1.42, 1.42, 1.38, 1.72, 2.12, 2.02])

    for col_index, header in enumerate(HEADERS):
        cell = table.cell(0, col_index)
        set_cell_text(cell, header, bold=True, font_size=7.5)
        shade(cell, "D9EAF7")

    for row_index, row in enumerate(ROWS, start=1):
        values = [
            row["problemas"],
            row["objetivos"],
            row["hipotesis"],
            row["variables"],
            row["dimensiones"],
            row["indicadores"],
            row["metodologia"],
        ]
        fill = "F4F9FD" if row["nivel"] == "General" else "FFFFFF"
        for col_index, value in enumerate(values):
            cell = table.cell(row_index, col_index)
            text = f"{row['nivel']}\n\n{value}" if col_index == 0 else value
            set_cell_text(cell, text, font_size=6.6 if row["nivel"] != "General" else 6.9)
            shade(cell, fill)

    doc.add_paragraph()
    val_title = doc.add_paragraph()
    val_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = val_title.add_run("CUADRO DE VALIDACIÓN DE COHERENCIA VERTICAL")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)

    val_rows = [
        ("Problema general ↔ Objetivo general", "Mismo constructo: sistema web + IA → predicción IRE", "Cumple"),
        ("Problema general ↔ Hipótesis H0/H1", "H0 niega mejora; H1 afirma mejora con IRE (α=0,05)", "Cumple"),
        ("ES1: digitalización ↔ VI-D1/D3", "Indicadores = operacionalización VI-D1 (6) + disponibilidad D3", "Cumple"),
        ("ES2: MAE/RMSE ↔ RandomForestRegressor", "Código: ai-service/models/demand/features_ml.py", "Cumple"),
        ("ES2: horizonte 7/15/30 días", "UI: adminPredictionsTypes.ts HORIZON_OPTIONS", "Cumple"),
        ("ES3: IRE 0–100, pesos 40/35/25", "Doc 07 + panel AdminPredictions", "Cumple"),
        ("ES3: alertas ≥51 / ≥76", "Niveles alto y crítico (no escala 0–1)", "Cumple"),
        ("Stack Supabase (no Firestore BD)", "09-implementacion-despliegue-ci.md", "Cumple"),
        ("Entrevistas + Likert declaradas", "ES1 y ES3 metodología; guía entrevista en repo", "Cumple"),
        ("AUC-ROC no prometido sin etiquetas", "Nota ES3; doc 07 condiciona calibración", "Cumple"),
    ]
    vt = doc.add_table(rows=len(val_rows) + 1, cols=3)
    vt.style = "Table Grid"
    set_widths(vt, [2.4, 4.8, 1.2])
    for i, h in enumerate(["Criterio de coherencia", "Evidencia / regla", "Estado"]):
        set_cell_text(vt.cell(0, i), h, bold=True, font_size=7.5)
        shade(vt.cell(0, i), "D9EAF7")
    for ri, (c1, c2, c3) in enumerate(val_rows, start=1):
        set_cell_text(vt.cell(ri, 0), c1, font_size=7)
        set_cell_text(vt.cell(ri, 1), c2, font_size=7)
        set_cell_text(vt.cell(ri, 2), c3, font_size=7)
        shade(vt.cell(ri, 2), "E2F0D9")

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = footer.add_run(
        "Validación de coherencia: (1) cada problema específico tiene objetivo, hipótesis e "
        "indicadores propios; (2) VI/VD coinciden con la matriz de operacionalización; "
        "(3) entrevistas a empleadores = técnica cualitativa (Cap. metodología); encuesta Likert = "
        "cuantitativa pre/post; SUS = usabilidad usuarios finales (anexo ISO, no sustituye encuesta); "
        "(4) IRE en escala 0–100 según implementación en ai-service y panel AdminPredictions."
    )
    run.font.name = "Arial"
    run.font.size = Pt(7)
    run.italic = True

    for target in (OUT_DOCX_V2, OUT_DOCX):
        try:
            doc.save(target)
            print(f"OK: {target}")
            break
        except PermissionError:
            if target == OUT_DOCX:
                print(f"AVISO: {OUT_DOCX} abierto en Word — guardado en {OUT_DOCX_V2}")
            continue
    else:
        raise PermissionError("No se pudo guardar el DOCX (cierre el archivo en Word)")


def build_csv() -> None:
    OUT_CSV.parent.mkdir(parents=True, exist_ok=True)
    fieldnames = ["Nivel", *HEADERS]
    with OUT_CSV.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames, delimiter=";")
        writer.writeheader()
        for row in ROWS:
            writer.writerow(
                {
                    "Nivel": row["nivel"],
                    "Problemas": row["problemas"],
                    "Objetivos": row["objetivos"],
                    "Hipótesis": row["hipotesis"],
                    "Variables": row["variables"],
                    "Dimensiones": row["dimensiones"],
                    "Indicadores": row["indicadores"],
                    "Metodología": row["metodologia"],
                }
            )
    print(f"OK: {OUT_CSV}")


def main() -> None:
    build_docx()
    build_csv()


if __name__ == "__main__":
    main()
