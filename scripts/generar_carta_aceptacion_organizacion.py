from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT = Path("c:/Cazatura Vilchez V3/Carta_Aceptacion_Organizacion_Calzatura_Vilchez.docx")


def set_font(run, size: float = 11, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def add_paragraph(doc: Document, text: str = "", size: float = 11, bold: bool = False, align=None):
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_font(run, size, bold)
    return paragraph


def add_justified(doc: Document, text: str):
    paragraph = add_paragraph(doc, text, 11)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(8)
    return paragraph


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

add_paragraph(doc, "ANEXO 05. CARTA DE ACEPTACIÓN DE LA ORGANIZACIÓN", 12, True, WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc)

add_paragraph(doc, "Huancayo, 6 de junio de 2026", 11, False, WD_ALIGN_PARAGRAPH.RIGHT)
add_paragraph(doc)

add_paragraph(doc, "Señores:", 11, True)
add_paragraph(doc, "UNIVERSIDAD CONTINENTAL", 11, True)
add_paragraph(doc, "Facultad de Ingeniería", 11)
add_paragraph(doc, "Escuela Académico Profesional de Ingeniería de Sistemas e Informática", 11)
add_paragraph(doc, "Presente.-", 11)
add_paragraph(doc)

add_paragraph(doc, "Asunto: Carta de aceptación para el desarrollo de proyecto de investigación", 11, True)
add_paragraph(doc)

add_justified(
    doc,
    "Yo, [NOMBRE DEL REPRESENTANTE LEGAL], identificado(a) con DNI N.° [DNI], en mi calidad de [CARGO] "
    "de la empresa CALZATURA VILCHEZ, con RUC N.° [RUC] y domicilio en [DIRECCIÓN DE LA EMPRESA], "
    "por medio de la presente manifiesto la aceptación de la organización para que el Bach. Piero Vilchez "
    "desarrolle su proyecto de investigación titulado: “Sistema web de comercio electrónico con modelo de "
    "Inteligencia Artificial para la predicción del riesgo empresarial en la empresa Calzatura Vilchez”."
)

add_justified(
    doc,
    "La organización autoriza al investigador a realizar el levantamiento de información necesario para el "
    "desarrollo del estudio durante el periodo académico 2026, incluyendo la revisión de registros históricos "
    "de ventas, inventario, pedidos e ingresos disponibles; la observación de los procesos de venta y gestión "
    "de inventario; y la aplicación de los instrumentos de investigación al personal directivo y operativo "
    "que participe voluntariamente."
)

add_justified(
    doc,
    "Asimismo, se autoriza el desarrollo, implementación y validación del sistema web de comercio electrónico "
    "con modelo de Inteligencia Artificial, así como el cálculo y monitoreo del Índice de Riesgo Empresarial "
    "(IRE), con la finalidad de evaluar su contribución a la predicción del riesgo empresarial, la gestión de "
    "inventarios y la toma de decisiones operativas de Calzatura Vilchez."
)

add_justified(
    doc,
    "La información proporcionada por la empresa será utilizada exclusivamente con fines académicos y de "
    "investigación. El investigador se compromete a mantener la confidencialidad de los datos sensibles, "
    "resguardar la identidad de los participantes cuando corresponda y presentar los resultados de manera "
    "agregada, sin afectar los intereses comerciales ni la operación normal de la organización."
)

add_justified(
    doc,
    "La presente carta se emite a solicitud del interesado para los fines de aprobación, ejecución y "
    "sustentación del proyecto de investigación ante la Universidad Continental."
)

add_paragraph(doc)
add_paragraph(doc, "Atentamente,", 11)
add_paragraph(doc)
add_paragraph(doc)
add_paragraph(doc)

signature = doc.add_paragraph()
signature.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = signature.add_run("________________________________________")
set_font(run, 11)

for text in [
    "[NOMBRE DEL REPRESENTANTE LEGAL]",
    "[CARGO]",
    "CALZATURA VILCHEZ",
    "DNI N.° [DNI]",
    "RUC N.° [RUC]",
    "Firma y sello de la organización",
]:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_font(run, 10, text in ["[NOMBRE DEL REPRESENTANTE LEGAL]", "CALZATURA VILCHEZ"])

doc.save(OUT)
print(OUT)
