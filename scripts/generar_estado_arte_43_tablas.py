from pathlib import Path
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path("c:/Cazatura Vilchez V3/Estado_del_Arte_43_Tablas_CORREGIDO.docx")

# ── helpers ──────────────────────────────────────────────────────────────────

def _font(run, size, bold=False, color=None, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def _shade(cell, fill):
    tc = cell._tc.get_or_add_tcPr()
    shd = tc.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")

def _cell(cell, text, bold=False, size=7.5, align=WD_ALIGN_PARAGRAPH.LEFT,
          color=None, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    _font(r, size, bold, color, italic)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def _cell2(cell, label, name, size=7.5):
    """Two-run cell: bold label + normal name."""
    cell.text = ""
    p = cell.paragraphs[0]
    r1 = p.add_run(label)
    _font(r1, size, bold=True)
    r2 = p.add_run(name)
    _font(r2, size, bold=False)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def _widths(table, widths):
    for row in table.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)

def _para(doc, text, size=10, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
          color=None, italic=False, sb=0, sa=2):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(text)
    _font(r, size, bold, color, italic)
    return p

# ── article renderer ──────────────────────────────────────────────────────────

def add_article(doc, art):
    """
    art keys:
      num, title, authors, journal, eje (optional),
      note (optional correction string),
      vi_name, vi_dims [(dim_name, [ind,...]),...], vi_instr [str,...],
      vd_name, vd_dims [...], vd_instr [str,...]
    """
    # heading
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(f"Artículo [{art['num']}] — ")
    _font(r1, 9, bold=True)
    r2 = p.add_run(art['title'])
    _font(r2, 9, bold=True, italic=True)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    r3 = p2.add_run(f"{art['authors']} — {art['journal']}")
    _font(r3, 8, italic=True)

    if art.get('note'):
        pn = doc.add_paragraph()
        pn.paragraph_format.space_before = Pt(0)
        pn.paragraph_format.space_after = Pt(3)
        rn = pn.add_run(f"CORRECCIÓN: {art['note']}")
        _font(rn, 8, bold=True, color="C00000")

    vi_rows = sum(len(inds) for _, inds in art['vi_dims'])
    vd_rows = sum(len(inds) for _, inds in art['vd_dims'])
    total = 1 + vi_rows + vd_rows

    tbl = doc.add_table(rows=total, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _widths(tbl, [1.52, 1.52, 2.76, 1.45])

    # header
    for i, h in enumerate(["Variables", "Dimensiones", "Indicadores", "Instrumentos"]):
        c = tbl.cell(0, i)
        _cell(c, h, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, color="FFFFFF")
        _shade(c, "1F4E79")

    # ── VI rows ──
    row = 1
    vi_start = row
    for dim_name, inds in art['vi_dims']:
        dim_start = row
        for ind in inds:
            _cell(tbl.cell(row, 2), ind, size=7.5)
            _shade(tbl.cell(row, 2), "FFFFFF")
            row += 1
        if len(inds) > 1:
            tbl.cell(dim_start, 1).merge(tbl.cell(row - 1, 1))
        _cell(tbl.cell(dim_start, 1), dim_name, bold=True, size=7.5,
              align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade(tbl.cell(dim_start, 1), "D6E4F0")

    if vi_rows > 1:
        tbl.cell(vi_start, 0).merge(tbl.cell(vi_start + vi_rows - 1, 0))
    _cell2(tbl.cell(vi_start, 0), "Variable\nindependiente:\n", art['vi_name'])
    _shade(tbl.cell(vi_start, 0), "D9EAF7")

    if vi_rows > 1:
        tbl.cell(vi_start, 3).merge(tbl.cell(vi_start + vi_rows - 1, 3))
    _cell(tbl.cell(vi_start, 3), "\n".join(art['vi_instr']), size=7.5)
    _shade(tbl.cell(vi_start, 3), "EBF5FB")

    # ── VD rows ──
    vd_start = row
    for dim_name, inds in art['vd_dims']:
        dim_start = row
        for ind in inds:
            _cell(tbl.cell(row, 2), ind, size=7.5)
            _shade(tbl.cell(row, 2), "FFFEF8")
            row += 1
        if len(inds) > 1:
            tbl.cell(dim_start, 1).merge(tbl.cell(row - 1, 1))
        _cell(tbl.cell(dim_start, 1), dim_name, bold=True, size=7.5,
              align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade(tbl.cell(dim_start, 1), "FFF5CC")

    if vd_rows > 1:
        tbl.cell(vd_start, 0).merge(tbl.cell(vd_start + vd_rows - 1, 0))
    _cell2(tbl.cell(vd_start, 0), "Variable\ndependiente:\n", art['vd_name'])
    _shade(tbl.cell(vd_start, 0), "FFF2CC")

    if vd_rows > 1:
        tbl.cell(vd_start, 3).merge(tbl.cell(vd_start + vd_rows - 1, 3))
    _cell(tbl.cell(vd_start, 3), "\n".join(art['vd_instr']), size=7.5)
    _shade(tbl.cell(vd_start, 3), "FFFDF0")

    doc.add_paragraph().paragraph_format.space_after = Pt(1)

    if art.get("contrib"):
        pc = doc.add_paragraph()
        pc.paragraph_format.space_before = Pt(0)
        pc.paragraph_format.space_after  = Pt(4)
        pc.paragraph_format.left_indent  = Inches(0)
        pPr = pc._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "EBF5F0")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        pPr.append(shd)
        r_label = pc.add_run("Contribución a la tesis: ")
        _font(r_label, 7.5, bold=True, color="1A5276")
        r_text = pc.add_run(art["contrib"])
        _font(r_text, 7.5, italic=True)
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    if art.get("evidencia"):
        pe = doc.add_paragraph()
        pe.paragraph_format.space_before = Pt(0)
        pe.paragraph_format.space_after  = Pt(6)
        pPr = pe._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "E8F8F5")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        pPr.append(shd)
        r_label = pe.add_run("Evidencia funcional en Calzatura Vilchez: ")
        _font(r_label, 7.5, bold=True, color="117A65")
        r_text = pe.add_run(art["evidencia"])
        _font(r_text, 7.5, italic=False)


def add_eje_header(doc, eje_num, label):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"EJE {eje_num} — {label}")
    _font(r, 11, bold=True, color="FFFFFF")
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "2E75B6")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    pPr.append(shd)

# ══════════════════════════════════════════════════════════════════════════════
#  ARTICLE DATA
# ══════════════════════════════════════════════════════════════════════════════

ARTICLES = [

# ─── EJE 1 ───────────────────────────────────────────────────────────────────
{"eje": 1,
 "num": "01", "title": "Digital transformation: A multidisciplinary reflection and research agenda",
 "authors": "Verhoef, P.C. et al. (2021)",
 "journal": "Journal of Business Research · Q1 · JIF 10.97",
 "vi_name": "Transformación Digital\nEmpresarial",
 "vi_dims": [
   ("Estrategia digital", ["Nivel de madurez digital (escala 1–5)",
                            "Inversión en TI (% del presupuesto total)",
                            "N.° de plataformas digitales adoptadas"]),
   ("Modelo de negocio digital", ["Canales digitales activos (n.°)",
                                   "Automatización de procesos (%)",
                                   "Integración de ecosistemas digitales (escala)"]),
   ("Orientación al cliente digital", ["NPS digital (Net Promoter Score)",
                                        "Tasa de personalización de ofertas (%)",
                                        "Interacciones omnicanal (n.°/mes)"]),
 ],
 "vi_instr": ["Revisión sistemática de literatura",
              "Análisis bibliométrico (Scopus/WoS)",
              "Encuesta Likert 5 puntos a directivos",
              "Análisis de regresión múltiple"],
 "vd_name": "Desempeño\nOrganizacional",
 "vd_dims": [
   ("Eficiencia operativa", ["Reducción de costos operativos (%)",
                              "Tiempo de ciclo de pedido (días)",
                              "ROI tecnológico (%)"]),
   ("Ventaja competitiva", ["Cuota de mercado digital (%)",
                             "N.° de innovaciones lanzadas/año",
                             "CSAT — Satisfacción del cliente (escala 1–10)"]),
 ],
 "vd_instr": ["Revisión sistemática de literatura",
              "Análisis bibliométrico (Scopus/WoS)",
              "Encuesta Likert 5 puntos a directivos",
              "Análisis de regresión múltiple"],
 "contrib": ("Aporta el marco teórico de la transformación digital empresarial que justifica la propuesta del sistema "
             "web en Calzatura Vilchez. Los indicadores de madurez digital, automatización de procesos y ROI "
             "tecnológico son métricas comparables con los resultados esperados del e-commerce propuesto (Variable "
             "Independiente — Digitalización comercial)."),
},

{"eje": 1,
 "num": "02", "title": "COVID-19 Pandemic: Shifting Digital Transformation to a High-Speed Gear in SMEs",
 "authors": "Soto-Acosta, P. (2020)",
 "journal": "Information Systems Management · Q1 · JIF 7.15",
 "vi_name": "Transformación Digital\nAcelerada por COVID-19",
 "vi_dims": [
   ("Adopción tecnológica de emergencia", ["Tiempo de adopción digital (días)",
                                            "N.° de herramientas digitales adoptadas",
                                            "% de procesos digitalizados durante pandemia"]),
   ("Reconfiguración del modelo de negocio", ["Ventas en línea / ventas totales (%)",
                                               "Reducción de operaciones presenciales (%)",
                                               "N.° de nuevos canales digitales habilitados"]),
 ],
 "vi_instr": ["Análisis de caso longitudinal",
              "Entrevistas semiestructuradas a gerentes de PYMEs",
              "Revisión de literatura (Scopus, WoS)",
              "Análisis estadístico descriptivo"],
 "vd_name": "Resiliencia y Desempeño\nde PYMEs",
 "vd_dims": [
   ("Continuidad operativa", ["Tasa de continuidad del negocio (%)",
                               "Ingresos mantenidos vs. período pre-crisis (%)"]),
   ("Adaptación digital sostenida", ["Crecimiento de ventas digitales post-pandemia (%)",
                                      "Tasa de retención de clientes digitales (%)",
                                      "NPS post-pandemia"]),
 ],
 "vd_instr": ["Análisis de caso longitudinal",
              "Entrevistas semiestructuradas a gerentes de PYMEs",
              "Revisión de literatura (Scopus, WoS)",
              "Análisis estadístico descriptivo"],
 "contrib": ("Sustenta la urgencia de digitalización acelerada en PYMEs, análoga a la situación de Calzatura Vilchez. "
             "Los indicadores de tasa de continuidad del negocio y crecimiento de ventas digitales son benchmarks "
             "directamente aplicables para medir el impacto del módulo de e-commerce propuesto (Variable "
             "Independiente — Digitalización de ventas en PYME)."),
},

{"eje": 1,
 "num": "03", "title": "Digital Innovation Management: Reinventing Innovation Management Research in a Digital World",
 "authors": "Nambisan, S. et al. (2017)",
 "journal": "MIS Quarterly · Q1 · JIF 12.50",
 "vi_name": "Gestión de Innovación\nDigital",
 "vi_dims": [
   ("Límites de la innovación digital", ["N.° de actores en el ecosistema digital",
                                          "Grado de apertura de la plataforma (escala 1–5)",
                                          "Tasa de co-creación de valor (%)"]),
   ("Procesos de innovación digital", ["Velocidad de iteración (sprints/mes)",
                                        "N.° de prototipos generados",
                                        "Tiempo de lanzamiento al mercado (días)"]),
   ("Resultados de la innovación", ["N.° de nuevas funcionalidades entregadas",
                                     "Adopción de innovación por usuarios (%)",
                                     "Valor generado por innovación (S/.)"]),
 ],
 "vi_instr": ["Revisión conceptual y marco integrador",
              "Análisis de casos de empresas digitales globales",
              "Revisión Scopus / IEEE Xplore / AIS"],
 "vd_name": "Desempeño Empresarial\npor Innovación",
 "vd_dims": [
   ("Competitividad", ["Posición relativa en el mercado (ranking sectorial)",
                        "N.° de innovaciones lanzadas/año"]),
   ("Sostenibilidad de la innovación", ["Tasa de retención de innovaciones activas (%)",
                                         "Inversión en I+D como % de ingresos"]),
 ],
 "vd_instr": ["Revisión conceptual y marco integrador",
              "Análisis de casos de empresas digitales globales",
              "Revisión Scopus / IEEE Xplore / AIS"],
 "contrib": ("Provee el marco conceptual de la innovación digital como proceso iterativo, respaldando la metodología "
             "ágil adoptada en el desarrollo del sistema. La velocidad de iteración y la tasa de adopción de "
             "innovación son indicadores de gestión del proyecto directamente aplicables al proceso de "
             "implementación propuesto (Variable Independiente — Marco metodológico del sistema)."),
},

{"eje": 1,
 "num": "04", "title": "The digital transformation of business models in the creative industries",
 "authors": "Li, F. (2020)",
 "journal": "British Journal of Management · Q1 · JIF 5.57",
 "vi_name": "Transformación Digital del\nModelo de Negocio",
 "vi_dims": [
   ("Digitalización de productos y servicios", ["% de oferta convertida a formato digital",
                                                 "N.° de servicios digitalizados",
                                                 "Nuevos canales digitales creados"]),
   ("Plataformización del negocio", ["N.° de socios integrados a la plataforma",
                                      "Ingresos generados por plataforma (% del total)",
                                      "Tasa de crecimiento de usuarios activos (%)"]),
 ],
 "vi_instr": ["Estudio de casos múltiples (caso cruzado)",
              "Entrevistas a directivos (n=24)",
              "Marcos teóricos comparativos (Value Chain)",
              "Análisis temático inductivo"],
 "vd_name": "Valor Empresarial\nCreado y Capturado",
 "vd_dims": [
   ("Creación de valor", ["Ingresos incrementales por digitalización (%)",
                           "ROI del modelo de negocio digital"]),
   ("Captura de valor", ["Margen neto del negocio digital (%)",
                          "Cuota de mercado ganada (%)"]),
 ],
 "vd_instr": ["Estudio de casos múltiples (caso cruzado)",
              "Entrevistas a directivos (n=24)",
              "Marcos teóricos comparativos (Value Chain)",
              "Análisis temático inductivo"],
 "contrib": ("Justifica la transformación del modelo de negocio presencial de Calzatura Vilchez hacia un canal "
             "digital de e-commerce. Los indicadores de ingresos incrementales por digitalización y ROI del canal "
             "digital son métricas de impacto económico esperadas tras la implementación del sistema propuesto "
             "(Variable Independiente — Digitalización del modelo comercial)."),
},

{"eje": 1,
 "num": "05", "title": "Why do small and medium enterprises use social media marketing and what is the impact",
 "authors": "Chatterjee, S. y Kumar Kar, A. (2020)",
 "journal": "Journal of Business Research · Q1 · JIF 10.97",
 "vi_name": "Adopción de Marketing\nDigital en PYMEs",
 "vi_dims": [
   ("Factores de adopción tecnológica", ["Percepción de utilidad (Likert 1–5)",
                                          "Facilidad de uso percibida (Likert 1–5)",
                                          "Presión competitiva y social percibida"]),
   ("Intensidad de uso del canal digital", ["Frecuencia de publicaciones (n.°/semana)",
                                             "Presupuesto digital (% de ventas totales)",
                                             "N.° de plataformas sociales activas"]),
 ],
 "vi_instr": ["Encuesta cuantitativa (n=302, Likert 5)",
              "Modelado PLS-SEM (SmartPLS 3)",
              "Muestreo aleatorio estratificado",
              "Análisis de confiabilidad (α Cronbach)"],
 "vd_name": "Desempeño Empresarial\nde PYMEs",
 "vd_dims": [
   ("Desempeño comercial", ["Crecimiento de ventas (%)",
                             "N.° de nuevos clientes captados",
                             "Tasa de conversión digital (%)"]),
   ("Reconocimiento y posicionamiento de marca", ["N.° de seguidores/fans activos",
                                                   "Engagement rate (%)",
                                                   "Alcance mensual de publicaciones"]),
 ],
 "vd_instr": ["Encuesta cuantitativa (n=302, Likert 5)",
              "Modelado PLS-SEM (SmartPLS 3)",
              "Muestreo aleatorio estratificado",
              "Análisis de confiabilidad (α Cronbach)"],
 "contrib": ("Valida empíricamente que la adopción de canales digitales en PYMEs genera crecimiento de ventas y "
             "captación de nuevos clientes, justificando el módulo de ventas en línea del sistema. Los indicadores "
             "de tasa de conversión digital y crecimiento de ventas son KPIs comparables con los resultados "
             "esperados del e-commerce de Calzatura Vilchez (Variable Independiente — Canal digital)."),
},

{"eje": 1,
 "num": "06", "title": "Digital marketing: A framework, review and research agenda",
 "authors": "Kannan, P.K. (2017)",
 "journal": "International Journal of Research in Marketing · Q1 · JIF 7.20",
 "vi_name": "Estrategia de Marketing\nDigital",
 "vi_dims": [
   ("Canales y medios digitales", ["N.° de canales digitales activos",
                                    "Costo por lead digital — CPL (S/.)",
                                    "Alcance mensual (impresiones)"]),
   ("Personalización y segmentación", ["CTR — Click-through rate (%)",
                                        "Tasa de apertura de emails (%)",
                                        "Tasa de personalización de mensajes (%)"]),
   ("Analítica de marketing", ["ROI de campañas digitales (%)",
                                "CAC — Costo de adquisición de cliente (S/.)",
                                "LTV — Valor de vida del cliente (S/.)"]),
 ],
 "vi_instr": ["Revisión sistemática (meta-análisis)",
              "Marco conceptual integrador",
              "Encuesta a consumidores online",
              "Google Analytics / CRM (datos secundarios)"],
 "vd_name": "Comportamiento del\nConsumidor Digital",
 "vd_dims": [
   ("Intención de compra", ["Tasa de intención medida en encuesta (Likert)",
                             "Tasa de conversión real (%)"]),
   ("Fidelización del cliente", ["Tasa de retención (%)",
                                  "NPS",
                                  "Frecuencia de recompra (veces/mes)"]),
 ],
 "vd_instr": ["Revisión sistemática (meta-análisis)",
              "Marco conceptual integrador",
              "Encuesta a consumidores online",
              "Google Analytics / CRM (datos secundarios)"],
 "contrib": ("Fundamenta la estrategia de marketing digital que acompañará al módulo de e-commerce del sistema. "
             "El CTR, la tasa de conversión y el LTV son métricas de evaluación del canal digital que pueden "
             "medirse en la plataforma propuesta; sirven como indicadores secundarios del desempeño comercial "
             "del sistema (Variable Independiente — Módulo comercial digital)."),
},

{"eje": 1,
 "num": "07", "title": "The impact of digital transformation on the retailing value chain",
 "authors": "Reinartz, W.; Wiegand, N. y Imschloss, M. (2019)",
 "journal": "International Journal of Research in Marketing · Q1 · JIF 7.20",
 "vi_name": "Transformación Digital\nen Retail",
 "vi_dims": [
   ("Digitalización de la cadena de valor", ["% de procesos clave digitalizados",
                                              "N.° de integraciones vía API activas",
                                              "Tiempo de procesamiento de pedidos (horas)"]),
   ("Integración omnicanal", ["N.° de canales integrados",
                               "Consistencia de experiencia entre canales (%)",
                               "Tasa de compra omnicanal (%)"]),
 ],
 "vi_instr": ["Análisis de cadena de valor (Porter adaptado)",
              "Revisión sistemática de literatura",
              "Entrevistas a directivos de retail (n=18)",
              "Datos de POS y CRM (análisis cuantitativo)"],
 "vd_name": "Cadena de Valor\nMinorista",
 "vd_dims": [
   ("Eficiencia operacional", ["Rotación de inventario (veces/año)",
                                "Costo de fulfillment por pedido (S/.)",
                                "Tiempo de entrega al cliente (días)"]),
   ("Satisfacción y lealtad del cliente", ["CSAT (escala 1–10)",
                                            "Tasa de devoluciones (%)",
                                            "NPS del canal digital"]),
 ],
 "vd_instr": ["Análisis de cadena de valor (Porter adaptado)",
              "Revisión sistemática de literatura",
              "Entrevistas a directivos de retail (n=18)",
              "Datos de POS y CRM (análisis cuantitativo)"],
 "contrib": ("ARTÍCULO NÚCLEO. Fundamenta directamente la digitalización de la cadena de valor minorista con "
             "énfasis en la integración omnicanal, gestión de pedidos y stock, análoga a la operación de "
             "Calzatura Vilchez. Los indicadores de rotación de inventario, tiempo de entrega y NPS del canal "
             "digital son KPIs del módulo e-commerce adaptables al IRE (Variable Independiente — Sistema web)."),
},

{"eje": 1,
 "num": "08", "title": "The digitalization of retailing: An exploratory framework",
 "authors": "Hagberg, J.; Sundstrom, M. y Egels-Zandén, N. (2016)",
 "journal": "International Journal of Retail & Distribution Management · Q1 · JIF 4.50",
 "vi_name": "Digitalización del Retail",
 "vi_dims": [
   ("Adopción de tecnologías en el punto de venta", ["N.° de tecnologías digitales adoptadas",
                                                      "Inversión en tecnología retail (% de ventas)",
                                                      "Nivel de madurez digital del punto de venta (escala)"]),
   ("Cambio en procesos y oferta retail", ["% de procesos de caja automatizados",
                                            "N.° de SKUs disponibles en canal online",
                                            "N.° de servicios digitales nuevos ofrecidos"]),
 ],
 "vi_instr": ["Marco exploratorio conceptual",
              "Análisis de tendencias tecnológicas (revisión de literatura)",
              "Estudio de casos de retail europeo",
              "Entrevistas semiestructuradas a gerentes"],
 "vd_name": "Modelo de Negocio\nMinorista Transformado",
 "vd_dims": [
   ("Propuesta de valor digital", ["Tasa de adopción de servicios digitales por clientes (%)",
                                    "Diferenciación percibida respecto a competidores (Likert)"]),
   ("Experiencia del cliente en tienda", ["Tiempo promedio de compra (minutos)",
                                           "Satisfacción en punto de venta (CSAT)",
                                           "Tasa de conversión en tienda física (%)"]),
 ],
 "vd_instr": ["Marco exploratorio conceptual",
              "Análisis de tendencias tecnológicas (revisión de literatura)",
              "Estudio de casos de retail europeo",
              "Entrevistas semiestructuradas a gerentes"],
 "contrib": ("ARTÍCULO NÚCLEO. Provee el modelo de digitalización del punto de venta minorista directamente "
             "aplicable al contexto de una empresa de calzado. El N.° de SKUs en canal online y la tasa de "
             "adopción de servicios digitales por clientes son indicadores del módulo de catálogo y ventas del "
             "sistema propuesto (Variable Independiente — E-commerce en retail de calzado)."),
},

# ─── EJE 2 ───────────────────────────────────────────────────────────────────
{"eje": 2,
 "num": "09", "title": "Artificial intelligence for decision making in the era of Big Data",
 "authors": "Yuan, Y.; Edwards, J.S. y Dwivedi, Y.K. (2021)",
 "journal": "International Journal of Information Systems · Q1 · JIF 9.14",
 "vi_name": "IA para Toma de\nDecisiones con Big Data",
 "vi_dims": [
   ("Capacidades de procesamiento de datos", ["Volumen procesado (GB/día)",
                                               "Velocidad de análisis (registros/segundo)",
                                               "Variedad de fuentes de datos integradas (n.°)"]),
   ("Algoritmos de IA aplicados", ["Accuracy del modelo (%)",
                                    "Recall (%)",
                                    "F1-score"]),
   ("Integración empresarial de IA", ["N.° de procesos asistidos por IA",
                                       "Tasa de adopción interna (%)",
                                       "ROI de implementación de IA (%)"]),
 ],
 "vi_instr": ["Revisión sistemática (PRISMA)",
              "Encuesta a tomadores de decisión",
              "Benchmarking de algoritmos (Python/R)",
              "Análisis de tendencias (2010–2021)"],
 "vd_name": "Calidad de las Decisiones\nEmpresariales",
 "vd_dims": [
   ("Velocidad de decisión", ["Tiempo de decisión (horas)",
                               "Ciclo análisis-acción reducido (%)"]),
   ("Precisión decisional", ["Tasa de error decisional (%)",
                              "Reducción de riesgo operativo (%)",
                              "Ahorro por mejores decisiones (S/.)"]),
 ],
 "vd_instr": ["Revisión sistemática (PRISMA)",
              "Encuesta a tomadores de decisión",
              "Benchmarking de algoritmos (Python/R)",
              "Análisis de tendencias (2010–2021)"],
 "contrib": ("ARTÍCULO NÚCLEO. Establece el fundamento teórico del componente de IA para toma de decisiones en "
             "el sistema propuesto. Accuracy, Recall y F1-score son las métricas estándar de evaluación del "
             "módulo predictivo del IRE. El ROI de implementación de IA y la reducción del riesgo operativo son "
             "los KPIs de impacto gerencial del sistema en Calzatura Vilchez (Variable Dependiente — IRE)."),
},

{"eje": 2,
 "num": "10", "title": "How artificial intelligence will change the future of marketing",
 "authors": "Davenport, T. et al. (2020)",
 "journal": "Journal of the Academy of Marketing Science · Q1 · JIF 18.53",
 "vi_name": "Inteligencia Artificial\nAplicada al Marketing",
 "vi_dims": [
   ("Personalización automática con IA", ["Tasa de personalización de contenidos (%)",
                                           "Relevancia de recomendaciones (CTR %)",
                                           "Tiempo de generación de recomendación (ms)"]),
   ("Automatización de procesos de marketing", ["N.° de campañas automatizadas",
                                                  "Reducción de tiempo operativo de marketing (%)",
                                                  "ROI de campañas basadas en IA (%)"]),
 ],
 "vi_instr": ["Revisión de literatura prospectiva",
              "Análisis Delphi (rondas de expertos)",
              "Entrevistas a directores de marketing (n=12)",
              "Análisis de plataformas IA comerciales"],
 "vd_name": "Efectividad del\nMarketing Digital",
 "vd_dims": [
   ("Desempeño de campañas", ["Tasa de conversión (%)",
                               "CAC — Costo de adquisición (S/.)",
                               "ROAS — Retorno en gasto publicitario"]),
   ("Experiencia y retención del cliente", ["NPS",
                                             "Satisfacción con recomendaciones (Likert)",
                                             "Tasa de retención (%)"]),
 ],
 "vd_instr": ["Revisión de literatura prospectiva",
              "Análisis Delphi (rondas de expertos)",
              "Entrevistas a directores de marketing (n=12)",
              "Análisis de plataformas IA comerciales"],
 "contrib": ("Justifica el uso de IA para personalización y automatización en el sistema de e-commerce propuesto. "
             "Las métricas de ROI de campañas con IA y tasa de conversión evalúan el impacto comercial del módulo "
             "IA como soporte a decisiones de ventas; complementa el análisis del subíndice de ingresos del IRE "
             "(Variable Independiente — Módulo IA comercial)."),
},

{"eje": 2,
 "num": "11", "title": "Organizational decision-making structures in the age of artificial intelligence",
 "authors": "Shrestha, Y.R.; Ben-Menahem, S.M. y von Krogh, G. (2019)",
 "journal": "Academy of Management Perspectives (antes AME) · Q1 · JIF 9.87",
 "vi_name": "Integración de IA en\nEstructuras de Decisión\nOrganizacional",
 "vi_dims": [
   ("Nivel de automatización decisional", ["% de decisiones automatizadas en procesos clave",
                                            "Grado de autonomía de agentes IA (escala 1–5)",
                                            "N.° de áreas funcionales con IA integrada"]),
   ("Rediseño organizacional para IA", ["N.° de roles modificados por IA",
                                         "Reducción de capas jerárquicas de aprobación",
                                         "Tiempo promedio de aprobación reducido (%)"]),
 ],
 "vi_instr": ["Revisión conceptual y teórica",
              "Análisis de casos de empresas con IA (n=9)",
              "Teoría fundamentada en datos",
              "Entrevistas semiestructuradas a altos directivos"],
 "vd_name": "Desempeño\nOrganizacional con IA",
 "vd_dims": [
   ("Agilidad organizacional", ["Tiempo de respuesta a cambios del mercado (días)",
                                 "Tasa de adaptación a nuevas condiciones (encuesta)"]),
   ("Calidad de las decisiones", ["Tasa de decisiones correctas evaluadas ex-post (%)",
                                   "Reducción de sesgos cognitivos (escala Likert)"]),
 ],
 "vd_instr": ["Revisión conceptual y teórica",
              "Análisis de casos de empresas con IA (n=9)",
              "Teoría fundamentada en datos",
              "Entrevistas semiestructuradas a altos directivos"],
 "contrib": ("ARTÍCULO NÚCLEO. Fundamenta la integración de IA como apoyo a las decisiones gerenciales operativas "
             "de Calzatura Vilchez. La reducción de riesgo operativo y el ahorro por mejores decisiones son KPIs "
             "directamente vinculados al impacto del IRE sobre la gestión empresarial. NOTA: La revista fue "
             "renombrada de 'Academy of Management Executive' a 'Academy of Management Perspectives' en 2014 "
             "(Variable Dependiente — Impacto del IRE en decisiones gerenciales)."),
},

{"eje": 2,
 "num": "12", "title": "Big data analytics and firm performance: Effects of dynamic capabilities",
 "authors": "Amba, S.F. et al. (2017)",
 "journal": "International Journal of Production Economics · Q1 · JIF 9.83",
 "vi_name": "Analítica de Big Data",
 "vi_dims": [
   ("Infraestructura de datos", ["Capacidad de almacenamiento (TB)",
                                  "Velocidad de procesamiento (Mbps)",
                                  "N.° de fuentes de datos integradas"]),
   ("Capacidades analíticas dinámicas", ["N.° de modelos predictivos en producción",
                                          "Cobertura analítica (% de procesos cubiertos)",
                                          "Exactitud promedio de modelos (%)"]),
 ],
 "vi_instr": ["Encuesta cuantitativa (n=302)",
              "PLS-SEM (SmartPLS 3)",
              "SPSS 24 para análisis factorial",
              "Prueba de confiabilidad (α de Cronbach > 0.70)"],
 "vd_name": "Desempeño de la Empresa",
 "vd_dims": [
   ("Desempeño operativo", ["Eficiencia de procesos (%)",
                             "Reducción de costos operativos (%)",
                             "Tiempo de ciclo reducido (%)"]),
   ("Desempeño estratégico", ["Crecimiento de ingresos (%)",
                               "Cuota de mercado (%)",
                               "ROI de inversión en tecnología (%)"]),
 ],
 "vd_instr": ["Encuesta cuantitativa (n=302)",
              "PLS-SEM (SmartPLS 3)",
              "SPSS 24 para análisis factorial",
              "Prueba de confiabilidad (α de Cronbach > 0.70)"],
 "contrib": ("ARTÍCULO NÚCLEO. Valida empíricamente la relación entre analítica de Big Data y desempeño "
             "empresarial. La exactitud promedio de modelos y la reducción de costos operativos son métricas "
             "de calidad y valor del sistema propuesto. Justifica que la inversión en analítica de datos de "
             "Calzatura Vilchez generará mejoras medibles en su desempeño operativo y estratégico "
             "(Variable Dependiente — Desempeño empresarial con IA)."),
},

{"eje": 2,
 "num": "13", "title": "Statistical and Machine Learning forecasting methods: Concerns and ways forward",
 "authors": "Makridakis, S.; Spiliotis, E. y Assimakopoulos, V. (2018)",
 "journal": "PLOS ONE · Q1 · JIF 3.75",
 "vi_name": "Método de Pronóstico\n(Estadístico vs. Machine\nLearning)",
 "vi_dims": [
   ("Modelos estadísticos clásicos", ["MAE — Error Absoluto Medio",
                                       "MASE — Error Absoluto Escalado Medio",
                                       "sMAPE — Error Porcentual Absoluto Simétrico"]),
   ("Modelos de Machine Learning evaluados", ["RMSE — Raíz del Error Cuadrático Medio",
                                               "Tiempo de entrenamiento (segundos)",
                                               "Complejidad computacional (notación Big-O)"]),
   ("Configuración y ajuste del modelo", ["N.° de hiperparámetros optimizados",
                                           "Tamaño del conjunto de entrenamiento (registros)",
                                           "Horizonte de predicción (períodos futuros)"]),
 ],
 "vi_instr": ["Competencia M4 (100.000 series temporales heterogéneas)",
              "Python (statsmodels, scikit-learn)",
              "R (forecast, M4comp2018)",
              "Benchmarking estadístico comparativo"],
 "vd_name": "Precisión del Pronóstico\nde Demanda",
 "vd_dims": [
   ("Exactitud del pronóstico", ["sMAPE (%)",
                                  "MASE",
                                  "OWA — Overall Weighted Average"]),
   ("Estabilidad y robustez del modelo", ["Varianza del error entre horizontes",
                                           "Sesgo del modelo (bias)",
                                           "Consistencia de resultados entre series"]),
 ],
 "vd_instr": ["Competencia M4 (100.000 series temporales heterogéneas)",
              "Python (statsmodels, scikit-learn)",
              "R (forecast, M4comp2018)",
              "Benchmarking estadístico comparativo"],
 "contrib": ("ARTÍCULO NÚCLEO. Establece MAE, MASE, RMSE y sMAPE como las métricas estándar para evaluar la "
             "precisión del módulo predictivo del IRE. Demuestra que los modelos ML superan a los métodos "
             "estadísticos clásicos en series temporales complejas, justificando la elección de Random Forest "
             "frente a ARIMA en el sistema propuesto (Variable Dependiente — Precisión del IRE)."),
},

{"eje": 2,
 "num": "14", "title": "Retail forecasting: Research and practice",
 "authors": "Fildes, R.; Ma, S. y Kolassa, S. (2019)",
 "journal": "International Journal of Forecasting · Q1 · JIF 7.22",
 "vi_name": "Modelos de Predicción de\nDemanda en Retail",
 "vi_dims": [
   ("Factores que determinan la demanda retail", ["Precio por unidad (S/.)",
                                                   "Efecto de promociones (% de incremento de demanda)",
                                                   "Estacionalidad y efecto de calendario"]),
   ("Configuración del modelo predictivo", ["Horizonte de pronóstico (semanas)",
                                             "Granularidad del modelo (SKU-tienda)",
                                             "Datos históricos mínimos requeridos (días)"]),
 ],
 "vi_instr": ["Revisión sistemática de literatura (Scopus/WoS)",
              "Análisis empírico de datos de punto de venta (POS)",
              "Benchmarking de modelos (ARIMA, RF, XGBoost)",
              "Python / R para implementación"],
 "vd_name": "Precisión del Pronóstico\nde Ventas Retail",
 "vd_dims": [
   ("Error de pronóstico", ["MASE",
                             "RMSE",
                             "Sesgo del pronóstico (%)"]),
   ("Impacto en gestión de inventario", ["Nivel de quiebre de stock evitado (%)",
                                          "Exceso de inventario reducido (%)",
                                          "Costo del error de pronóstico (S/./período)"]),
 ],
 "vd_instr": ["Revisión sistemática de literatura (Scopus/WoS)",
              "Análisis empírico de datos de punto de venta (POS)",
              "Benchmarking de modelos (ARIMA, RF, XGBoost)",
              "Python / R para implementación"],
 "contrib": ("ARTÍCULO NÚCLEO. Fundamenta el pronóstico de demanda en retail como componente central del "
             "subíndice de riesgo de demanda (25% del IRE). El nivel de quiebre de stock evitado y el exceso de "
             "inventario reducido son los KPIs operativos directamente medibles en Calzatura Vilchez. El contexto "
             "de datos de punto de venta es equivalente al historial de ventas disponible en el sistema "
             "(Variable Dependiente — Subíndice de riesgo de demanda/stock del IRE)."),
},

{"eje": 2,
 "num": "15", "title": "Waiting for a sales renaissance in the fourth industrial revolution: Machine learning and selling",
 "authors": "Syam, N. y Sharma, A. (2018)",
 "journal": "Industrial Marketing Management · Q1 · JIF 8.49",
 "vi_name": "IA y Automatización\nAplicadas a Ventas",
 "vi_dims": [
   ("Automatización del proceso de ventas", ["% de tareas de ventas automatizadas",
                                              "Tiempo ahorrado por vendedor (horas/semana)",
                                              "N.° de interacciones con clientes automatizadas"]),
   ("Augmentación del vendedor con IA", ["Tasa de adopción de herramientas IA por fuerza de ventas (%)",
                                          "Mejora en tasa de cierre de ventas (%)",
                                          "N.° de insights accionables generados por el sistema"]),
 ],
 "vi_instr": ["Revisión teórica-prospectiva de literatura",
              "Análisis de tendencias industriales",
              "Entrevistas a directores comerciales (n=15)",
              "Análisis de plataformas CRM con IA (Salesforce, HubSpot)"],
 "vd_name": "Desempeño de la Fuerza\nde Ventas",
 "vd_dims": [
   ("Productividad individual del vendedor", ["Ingresos por vendedor (S/./mes)",
                                               "N.° de cuentas gestionadas activamente",
                                               "Tasa de retención de clientes (%)"]),
   ("Eficiencia del proceso de venta", ["Tiempo de ciclo de venta (días)",
                                         "CAC — Costo de adquisición de cliente (S/.)",
                                         "Tasa de conversión de oportunidades (%)"]),
 ],
 "vd_instr": ["Revisión teórica-prospectiva de literatura",
              "Análisis de tendencias industriales",
              "Entrevistas a directores comerciales (n=15)",
              "Análisis de plataformas CRM con IA (Salesforce, HubSpot)"],
 "contrib": ("Sustenta la automatización del proceso de ventas con IA como mejora de productividad comercial en "
             "Calzatura Vilchez. La mejora en tasa de cierre de ventas y la reducción del tiempo de ciclo "
             "respaldan los beneficios esperados del módulo de gestión de ventas del sistema; son indicadores "
             "comparables con el subíndice de ingresos del IRE (Variable Independiente — Módulo de ventas IA)."),
},

{"eje": 2,
 "num": "16", "title": "Artificial intelligence and innovation management: A review, framework, and research agenda",
 "authors": "Haefner, N. et al. (2021)",
 "journal": "R&D Management · Q1 · JIF 7.01",
 "vi_name": "IA como Motor de\nInnovación Empresarial",
 "vi_dims": [
   ("Exploración de ideas asistida por IA", ["N.° de ideas generadas por IA por período",
                                              "Tasa de aceptación de ideas propuestas por IA (%)",
                                              "Diversidad temática de ideas (índice de entropía)"]),
   ("Aceleración del proceso de innovación", ["Reducción del tiempo de desarrollo (%)",
                                               "Reducción del costo de I+D por IA (%)",
                                               "N.° de ciclos de iteración reducidos"]),
 ],
 "vi_instr": ["Revisión sistemática (PRISMA)",
              "Análisis bibliométrico (VOSviewer)",
              "Marco conceptual integrador multinivel",
              "Encuesta a gerentes de I+D+i"],
 "vd_name": "Resultados de Innovación\nEmpresarial",
 "vd_dims": [
   ("Output de innovación", ["N.° de nuevos productos lanzados/año",
                              "Tiempo al mercado — TTM (días)",
                              "Tasa de éxito del lanzamiento (%)"]),
   ("Capacidad de innovación organizacional", ["Nivel de madurez de capacidades IA (escala 1–5)",
                                                "Inversión en I+D (% de ingresos)",
                                                "N.° de patentes registradas/año"]),
 ],
 "vd_instr": ["Revisión sistemática (PRISMA)",
              "Análisis bibliométrico (VOSviewer)",
              "Marco conceptual integrador multinivel",
              "Encuesta a gerentes de I+D+i"],
 "contrib": ("Justifica la IA como driver estratégico de innovación que eleva la competitividad de Calzatura "
             "Vilchez frente a competidores con canales solo presenciales. El nivel de madurez de capacidades IA "
             "fundamenta la inversión tecnológica propuesta; el número de nuevos productos/servicios lanzados es "
             "un indicador de impacto estratégico del sistema (Marco conceptual de justificación de la propuesta)."),
},

{"eje": 2,
 "num": "17", "title": "Artificial intelligence (AI) and its implications for market knowledge in B2B marketing",
 "authors": "Paschen, J.; Kietzmann, J. y Kietzmann, T.C. (2019)",
 "journal": "Industrial Marketing Management · Q1 · JIF 8.49",
 "vi_name": "IA Aplicada al\nConocimiento de Mercado\nB2B",
 "vi_dims": [
   ("Generación de conocimiento de mercado con IA", ["N.° de fuentes de datos externas integradas",
                                                      "Velocidad de actualización del conocimiento (horas)",
                                                      "Profundidad del insight generado (escala)"]),
   ("Capacidades predictivas del sistema IA", ["Precisión de predicción del comportamiento del cliente B2B (%)",
                                                "Tiempo de respuesta del sistema de insights (ms)",
                                                "N.° de variables analizadas simultáneamente"]),
 ],
 "vi_instr": ["Revisión conceptual de literatura",
              "Análisis de casos B2B con IA (n=6)",
              "Entrevistas semiestructuradas a expertos de marketing",
              "Framework inductivo de análisis temático"],
 "vd_name": "Ventaja Competitiva en\nMercados B2B",
 "vd_dims": [
   ("Diferenciación del servicio", ["Tasa de retención de clientes B2B (%)",
                                     "CES — Customer Effort Score"]),
   ("Eficiencia comercial B2B", ["Tiempo de ciclo de ventas B2B (días)",
                                  "Tasa de conversión de prospectos a clientes (%)",
                                  "CAC B2B (S/.)"]),
 ],
 "vd_instr": ["Revisión conceptual de literatura",
              "Análisis de casos B2B con IA (n=6)",
              "Entrevistas semiestructuradas a expertos de marketing",
              "Framework inductivo de análisis temático"],
 "contrib": ("Sustenta el conocimiento de mercado como input clave del modelo predictivo del IRE. La profundidad "
             "del insight generado y el número de variables analizadas simultáneamente son características "
             "deseables del módulo IA de Calzatura Vilchez; los datos de comportamiento de clientes son "
             "insumos directos para el subíndice de riesgo de demanda (Variable Dependiente — Riesgo de demanda)."),
},

# ─── EJE 3 ───────────────────────────────────────────────────────────────────
{"eje": 3,
 "num": "18", "title": "Deep learning for financial applications: A survey",
 "authors": "Ozbayoglu, A.M.; Gudelek, M.U. y Sezer, O.B. (2020)",
 "journal": "Applied Soft Computing · Q1 · JIF 8.26",
 "vi_name": "Modelos de Deep Learning\npara Predicción Financiera",
 "vi_dims": [
   ("Arquitectura de la red neuronal profunda", ["N.° de capas ocultas del modelo",
                                                  "N.° de neuronas por capa",
                                                  "Función de activación (ReLU, sigmoid, tanh)"]),
   ("Datos financieros de entrada", ["N.° de variables financieras utilizadas",
                                      "Horizonte temporal de datos históricos (años)",
                                      "Frecuencia de actualización de datos"]),
   ("Optimización del modelo", ["Tasa de aprendizaje (learning rate)",
                                 "N.° de épocas de entrenamiento",
                                 "Técnica de regularización aplicada (dropout, L2)"]),
 ],
 "vi_instr": ["TensorFlow / Keras (Python)",
              "Conjuntos de datos financieros públicos (Compustat, WRDS)",
              "Validación cruzada k-fold (k=10)",
              "SHAP para interpretabilidad de predicciones"],
 "vd_name": "Precisión en Predicción\nde Riesgo Financiero",
 "vd_dims": [
   ("Métricas de exactitud del modelo", ["Accuracy (%)",
                                          "AUC-ROC",
                                          "F1-score",
                                          "Precision y Recall"]),
   ("Estabilidad temporal del modelo", ["Variación del error por horizonte (δ accuracy)",
                                         "Sesgo del modelo en datos fuera de muestra"]),
 ],
 "vd_instr": ["TensorFlow / Keras (Python)",
              "Conjuntos de datos financieros públicos (Compustat, WRDS)",
              "Validación cruzada k-fold (k=10)",
              "SHAP para interpretabilidad de predicciones"],
 "contrib": ("Antecedente metodológico de deep learning para predicción financiera. Establece AUC-ROC, F1-score "
             "y SHAP como métricas estándar de comparación de modelos predictivos. Justifica metodológicamente "
             "por qué Random Forest es preferible a DL para el IRE dado el tamaño limitado del dataset de "
             "una PYME. NOTA: el riesgo modelado es insolvencia financiera, diferente al IRE comercial-operativo."),
},

{"eje": 3,
 "num": "19", "title": "Machine learning models and bankruptcy prediction",
 "authors": "Barboza, F.; Kimura, H. y Altman, E. (2017)",
 "journal": "Expert Systems with Applications · Q1 · JIF 8.67",
 "vi_name": "Modelos de Machine\nLearning para Predicción\nde Quiebra",
 "vi_dims": [
   ("Variables financieras de entrada", ["Ratios de liquidez (Current Ratio, Quick Ratio)",
                                          "Ratios de endeudamiento (Debt/Equity)",
                                          "Rentabilidad (ROA, ROE, Margen neto)"]),
   ("Algoritmos de clasificación evaluados", ["Random Forest (n_estimators, max_depth)",
                                               "SVM (kernel, C, gamma)",
                                               "Boosting (XGBoost / AdaBoost)"]),
   ("Esquema de validación del modelo", ["Tamaño de muestra de entrenamiento (n)",
                                          "Técnica de validación (CV k-fold vs. holdout)",
                                          "Balance de clases (SMOTE aplicado)"]),
 ],
 "vi_instr": ["Python (scikit-learn)",
              "Base de datos COMPUSTAT (EE.UU., 1985–2013)",
              "Matriz de confusión y curva ROC",
              "XGBoost, Random Forest, SVM"],
 "vd_name": "Predicción de Quiebra\nEmpresarial",
 "vd_dims": [
   ("Exactitud del clasificador", ["Accuracy (%)",
                                    "AUC-ROC",
                                    "Especificidad (%)",
                                    "Sensibilidad (recall) (%)"]),
   ("Ventaja sobre modelos clásicos", ["Mejora de AUC vs. Altman Z-score",
                                        "Reducción del error tipo II (falsos negativos %)",
                                        "Mejora vs. regresión logística simple"]),
 ],
 "vd_instr": ["Python (scikit-learn)",
              "Base de datos COMPUSTAT (EE.UU., 1985–2013)",
              "Matriz de confusión y curva ROC",
              "XGBoost, Random Forest, SVM"],
 "contrib": ("ANTECEDENTE DIRECTO. Demuestra empíricamente que Random Forest supera a modelos clásicos (Altman "
             "Z-score, regresión logística) en predicción de riesgo empresarial, validando la elección "
             "algorítmica del IRE. La mejora de AUC vs. Altman y la reducción del error tipo II son métricas "
             "de comparación aplicables. NOTA: Diferenciar explícitamente: riesgo = quiebra financiera "
             "corporativa; en la tesis = IRE comercial-operativo de PYME."),
},

{"eje": 3,
 "num": "20", "title": "Bankruptcy prediction using imaged financial ratios and convolutional neural networks",
 "authors": "Osaka, T. (2019)",
 "journal": "Expert Systems with Applications · Q1 · JIF 8.67",
 "vi_name": "Ratios Financieros como\nImágenes para Redes CNN",
 "vi_dims": [
   ("Representación visual de datos financieros", ["N.° de ratios financieros convertidos a imagen",
                                                    "Resolución de la imagen generada (píxeles × píxeles)",
                                                    "Dimensión temporal representada (años de datos)"]),
   ("Arquitectura de la CNN", ["N.° de capas convolucionales",
                                "Tamaño del kernel (filtro)",
                                "N.° de filtros por capa convolucional"]),
 ],
 "vi_instr": ["TensorFlow / Keras (Python)",
              "Pillow / OpenCV para generación de imágenes",
              "Base de datos EDGAR (SEC, EE.UU.)",
              "Validación cruzada k-fold estratificada"],
 "vd_name": "Predicción de Quiebra\nEmpresarial por Imagen",
 "vd_dims": [
   ("Exactitud del modelo CNN", ["Accuracy (%)",
                                  "AUC-ROC",
                                  "F1-score"]),
   ("Comparativa con modelos tabulares clásicos", ["Mejora de exactitud vs. Altman Z-score (%)",
                                                    "Mejora de AUC vs. regresión logística"]),
 ],
 "vd_instr": ["TensorFlow / Keras (Python)",
              "Pillow / OpenCV para generación de imágenes",
              "Base de datos EDGAR (SEC, EE.UU.)",
              "Validación cruzada k-fold estratificada"],
 "contrib": ("Antecedente de CNN para predicción de riesgo con datos financieros visualizados. Justifica "
             "metodológicamente la elección de Random Forest sobre CNN en el IRE: datos de Calzatura Vilchez "
             "son tabulares estructurados, no imágenes; el modelo tabular es más interpretable y eficiente "
             "con datasets pequeños. Sirve como comparativa de enfoques en la sección de selección del modelo."),
},

{"eje": 3,
 "num": "21", "title": "Corporate default predictions using machine learning: Literature review",
 "authors": "Kim, M.J.; Cho, S.W. y Ryu, I. (2021)",
 "journal": "Sustainability · Q1 · JIF 3.89",
 "vi_name": "Métodos de ML para\nPredicción de Default\nEmpresarial",
 "vi_dims": [
   ("Variables predictoras utilizadas", ["Ratios financieros (liquidez, rentabilidad, apalancamiento)",
                                          "Variables macroeconómicas (inflación, tasa de interés)",
                                          "Datos alternativos (sentimiento de mercado, noticias)"]),
   ("Tipos de modelos de ML comparados", ["Regresión logística (baseline)",
                                           "Random Forest y Gradient Boosting",
                                           "Redes neuronales artificiales y LSTM"]),
 ],
 "vi_instr": ["Revisión sistemática (PRISMA, 2010–2020)",
              "Python (scikit-learn, XGBoost, LightGBM)",
              "SHAP para interpretabilidad",
              "Análisis bibliométrico VOSviewer"],
 "vd_name": "Precisión en Predicción\nde Incumplimiento",
 "vd_dims": [
   ("Métricas de desempeño del modelo", ["AUC-ROC",
                                          "Accuracy (%)",
                                          "Recall / sensibilidad (%)",
                                          "Precision (%)"]),
   ("Importancia de características", ["Feature importance (MDI)",
                                        "Correlación de Pearson y Spearman",
                                        "SHAP values para interpretación"]),
 ],
 "vd_instr": ["Revisión sistemática (PRISMA, 2010–2020)",
              "Python (scikit-learn, XGBoost, LightGBM)",
              "SHAP para interpretabilidad",
              "Análisis bibliométrico VOSviewer"],
 "contrib": ("ANTECEDENTE DIRECTO. La revisión sistemática valida Random Forest y XGBoost como los mejores "
             "algoritmos para predicción de riesgo empresarial. El análisis SHAP de importancia de "
             "características es replicable en el IRE para interpretar el peso de los subíndices de demanda, "
             "stock e ingresos. NOTA: Diferenciar incumplimiento financiero corporativo de IRE operacional PYME."),
},

{"eje": 3,
 "num": "22", "title": "Dynamics of firm financial evolution and bankruptcy prediction",
 "authors": "Du Jardin, P. (2021)",
 "journal": "Expert Systems with Applications · Q1 · JIF 8.67",
 "vi_name": "Evolución Financiera\nDinámica de la Empresa",
 "vi_dims": [
   ("Indicadores financieros dinámicos (variaciones temporales)",
    ["Variación del ROA año sobre año (Δ%)",
     "Variación del ratio de liquidez (Δ)",
     "Tendencia del margen operativo (pendiente)"]),
   ("Modelos predictivos de aprendizaje secuencial",
    ["LSTM — Long Short-Term Memory",
     "RNN — Redes Neuronales Recurrentes",
     "Modelos de supervivencia (Cox, Kaplan-Meier)"]),
 ],
 "vi_instr": ["Python (TensorFlow, Keras)",
              "Base de datos Diane (empresas francesas, n=15.000)",
              "Análisis de supervivencia (Kaplan-Meier)",
              "Validación walk-forward temporal"],
 "vd_name": "Predicción de Quiebra\ncon Perspectiva Dinámica",
 "vd_dims": [
   ("Exactitud de predicción longitudinal",
    ["AUC-ROC por horizonte (1, 2 y 3 años antes)",
     "Accuracy en cada horizonte temporal"]),
   ("Ventaja sobre modelos estáticos",
    ["Mejora de AUC-ROC vs. modelo de corte transversal (%)",
     "Anticipación de señal de distress financiero (años)"]),
 ],
 "vd_instr": ["Python (TensorFlow, Keras)",
              "Base de datos Diane (empresas francesas, n=15.000)",
              "Análisis de supervivencia (Kaplan-Meier)",
              "Validación walk-forward temporal"],
 "contrib": ("Antecedente de predicción dinámica longitudinal del riesgo con datos temporales. Justifica el uso "
             "de datos históricos secuenciales de múltiples períodos para entrenar el modelo del IRE. La "
             "anticipación de la señal de riesgo (1–3 años antes) fundamenta el horizonte temporal del módulo "
             "predictivo. NOTA: Riesgo modelado = quiebra financiera; en la tesis = IRE comercial-operativo PYME."),
},

{"eje": 3,
 "num": "23",
 "title": "Financial ratios and corporate governance indicators in bankruptcy prediction: A comprehensive study",
 "authors": "Jiang, D. et al. (2018)",
 "journal": "Journal of Accounting and Public Policy · Q1 · JIF 3.60",
 "note": ("Tabla reposicionada al marco teórico. Usa datos de grandes corporaciones EE.UU. "
          "(COMPUSTAT/CRSP 1996–2015), alejados del riesgo comercial-operativo del IRE. "
          "Citar como antecedente histórico de predicción de riesgo, NO como evidencia directa."),
 "vi_name": "Ratios Financieros e\nIndicadores de Gobernanza\nCorporativa",
 "vi_dims": [
   ("Ratios financieros clásicos", ["Liquidez corriente (Current Ratio)",
                                     "Ratio de endeudamiento (Debt/Assets)",
                                     "Margen neto (%)",
                                     "ROA y ROE (%)"]),
   ("Indicadores de gobernanza corporativa", ["Tamaño del directorio (n.° miembros)",
                                               "% de directores independientes",
                                               "Separación de roles CEO-Presidente (dummy)"]),
 ],
 "vi_instr": ["Regresión logística binaria (Logit)",
              "Análisis discriminante múltiple (MDA)",
              "SPSS / Stata",
              "Base de datos COMPUSTAT / CRSP (1996–2015)"],
 "vd_name": "Predicción de Quiebra\nEmpresarial",
 "vd_dims": [
   ("Exactitud del modelo integrado", ["Accuracy del modelo combinado (%)",
                                        "AUC-ROC combinado",
                                        "Mejora vs. modelo solo con ratios financieros"]),
   ("Poder explicativo de las variables", ["Significancia estadística (p < 0.05) por variable",
                                            "R² ajustado del modelo",
                                            "Odds ratio por cada variable significativa"]),
 ],
 "vd_instr": ["Regresión logística binaria (Logit)",
              "Análisis discriminante múltiple (MDA)",
              "SPSS / Stata",
              "Base de datos COMPUSTAT / CRSP (1996–2015)"],
 "contrib": ("REPOSICIONADO AL MARCO TEÓRICO. Antecedente histórico de predicción de quiebra con ratios "
             "financieros en grandes corporaciones EE.UU. (COMPUSTAT/CRSP 1996–2015). Se cita únicamente como "
             "evidencia de la evolución del campo desde MDA hasta ML moderno. NO presentar como definición del "
             "riesgo de Calzatura Vilchez: el IRE mide riesgo comercial-operativo de PYME, no insolvencia contable."),
},

{"eje": 3,
 "num": "24", "title": "Deep learning models for bankruptcy prediction using textual disclosures",
 "authors": "Cai, F. et al. (2019)",
 "journal": "European Accounting Review · Q1 · JIF 4.50",
 "vi_name": "Modelos de Deep Learning\ncon Datos Textuales (NLP)",
 "vi_dims": [
   ("Procesamiento de lenguaje natural", ["Vocabulario del modelo (n.° de tokens únicos)",
                                           "Representación vectorial (word2vec, BERT, TF-IDF)",
                                           "N.° de palabras procesadas por documento"]),
   ("Arquitectura del modelo de DL", ["Capas LSTM o Transformer utilizadas",
                                       "N.° de épocas de entrenamiento",
                                       "Tasa de aprendizaje (learning rate)"]),
 ],
 "vi_instr": ["Python (NLTK, SpaCy, HuggingFace Transformers)",
              "TensorFlow / PyTorch",
              "Informes anuales 10-K de SEC EDGAR",
              "Validación cruzada estratificada"],
 "vd_name": "Precisión en Predicción\nde Quiebra Textual",
 "vd_dims": [
   ("Exactitud del modelo basado en texto", ["Accuracy (%)",
                                              "AUC-ROC",
                                              "F1-score"]),
   ("Valor incremental del texto vs. ratios", ["Mejora de AUC vs. modelo solo con ratios financieros",
                                                "Información incremental aportada por el texto (%)"]),
 ],
 "vd_instr": ["Python (NLTK, SpaCy, HuggingFace Transformers)",
              "TensorFlow / PyTorch",
              "Informes anuales 10-K de SEC EDGAR",
              "Validación cruzada estratificada"],
 "contrib": ("Antecedente de NLP para predicción de riesgo con datos textuales. Justifica la elección de datos "
             "cuantitativos estructurados (ventas, stock, demanda) para el IRE en lugar de datos textuales: "
             "Calzatura Vilchez no genera informes anuales 10-K ni documentos financieros estructurados. "
             "Sirve como comparativa metodológica que respalda el diseño de variables del IRE."),
},

{"eje": 3,
 "num": "25", "title": "Variable selection and corporate bankruptcy forecasts",
 "authors": "Tian, S.; Yu, Y. y Guo, H. (2015)",
 "journal": "Journal of Banking & Finance · Q1 · JIF 3.84",
 "vi_name": "Selección de Variables\npara Pronóstico de Quiebra",
 "vi_dims": [
   ("Variables financieras candidatas", ["Ratios de liquidez (Current Ratio, Quick Ratio)",
                                          "Ratios de apalancamiento (Debt/Equity)",
                                          "Ratios de rentabilidad (ROA, ROE, Margen)"]),
   ("Métodos de selección de variables", ["LASSO — Regularización L1",
                                           "Stepwise selection (forward, backward)",
                                           "Importancia de variables por Random Forest"]),
 ],
 "vi_instr": ["LASSO, Ridge, Elastic Net (glmnet en R)",
              "Python / R para modelado",
              "COMPUSTAT (EE.UU., 1980–2010)",
              "PCA — Análisis de Componentes Principales"],
 "vd_name": "Pronóstico de Quiebra\nEmpresarial",
 "vd_dims": [
   ("Exactitud del modelo con variables seleccionadas",
    ["AUC-ROC con subconjunto de variables",
     "Accuracy (%)",
     "N.° de variables óptimas seleccionadas"]),
   ("Interpretabilidad del modelo", ["N.° de coeficientes significativos (p < 0.05)",
                                      "Variable de mayor importancia predictiva",
                                      "Odds ratio por variable clave"]),
 ],
 "vd_instr": ["LASSO, Ridge, Elastic Net (glmnet en R)",
              "Python / R para modelado",
              "COMPUSTAT (EE.UU., 1980–2010)",
              "PCA — Análisis de Componentes Principales"],
 "contrib": ("ARTÍCULO NÚCLEO METODOLÓGICO. Fundamenta la selección de variables predictivas del IRE usando "
             "LASSO, Stepwise y Random Forest. Los métodos de selección de características son directamente "
             "replicables para determinar el peso relativo de demanda (25%), stock (40%) e ingresos (35%) en "
             "el índice. Los coeficientes significativos identificados equivalen a los pesos del IRE "
             "(Variable Dependiente — Diseño y calibración del modelo IRE)."),
},

{"eje": 3,
 "num": "26",
 "title": "Financial ratios, discriminant analysis and the prediction of corporate bankruptcy",
 "authors": "Altman, E.I. (1968)",
 "journal": "The Journal of Finance · Q1 · JIF 7.75",
 "vi_name": "Ratios Financieros para\nPredicción de Quiebra\n(Modelo Z-score)",
 "vi_dims": [
   ("Ratios del modelo Z-score de Altman",
    ["X1: Capital trabajo / Activo total",
     "X2: Utilidades retenidas / Activo total",
     "X3: EBIT / Activo total",
     "X4: Valor mercado patrimonio / Pasivo total",
     "X5: Ventas / Activo total"]),
   ("Método de clasificación discriminante",
    ["Coeficientes de la función discriminante múltiple",
     "Punto de corte Z = 2.675",
     "Ponderación de cada ratio en la función"]),
 ],
 "vi_instr": ["Análisis discriminante múltiple (MDA)",
              "SPSS / SAS",
              "Base de datos empresas manufactureras EE.UU. (n=33 quebradas + 33 sanas)",
              "Estados financieros auditados"],
 "vd_name": "Clasificación del\nRiesgo de Quiebra",
 "vd_dims": [
   ("Exactitud de clasificación", ["Tasa de acierto en muestra original (%)",
                                    "Tasa de acierto en muestra de validación (%)",
                                    "Error tipo I y tipo II (%)"]),
   ("Zonas de riesgo definidas", ["Zona segura: Z > 2.99",
                                   "Zona gris: 1.81 < Z < 2.99",
                                   "Zona de distress: Z < 1.81"]),
 ],
 "vd_instr": ["Análisis discriminante múltiple (MDA)",
              "SPSS / SAS",
              "Base de datos empresas manufactureras EE.UU. (n=33 quebradas + 33 sanas)",
              "Estados financieros auditados"],
 "contrib": ("REFERENCIA FUNDACIONAL. El modelo Z-score de Altman (1968) establece el concepto de índice "
             "compuesto de riesgo como antecedente histórico directo del IRE. El punto de corte Z = 2.675 es "
             "el análogo conceptual del umbral del IRE. Se cita en el marco teórico como la primera propuesta "
             "cuantitativa de predicción del riesgo empresarial; la tesis evoluciona este concepto hacia datos "
             "operacionales de PYME."),
},

{"eje": 3,
 "num": "27", "title": "Financial ratios as predictors of failure",
 "authors": "Beaver, W.H. (1966)",
 "journal": "Journal of Accounting Research · Q1 · JIF 4.87",
 "vi_name": "Ratios Financieros\nUnivar. como Predictores\nde Quiebra",
 "vi_dims": [
   ("Ratios financieros individuales evaluados",
    ["Cash flow / Total deuda",
     "Utilidad neta / Activo total",
     "Deuda total / Activo total",
     "Capital de trabajo / Activo total",
     "Ratio de liquidez corriente"]),
   ("Horizonte temporal de predicción",
    ["1 año antes de la quiebra",
     "2 años antes",
     "3 años antes",
     "4 años antes",
     "5 años antes"]),
 ],
 "vi_instr": ["Análisis univariado de ratios financieros",
              "Dichotomous classification test",
              "Muestra pareada (n=34 quebradas + 34 sanas, 5 años)",
              "COMPUSTAT / Moody's Industrial Manual"],
 "vd_name": "Discriminación entre\nEmpresas Quebradas y\nSolventes",
 "vd_dims": [
   ("Poder predictivo de cada ratio",
    ["Tasa de clasificación correcta (%) por ratio",
     "Error de predicción en dichotomous classification"]),
   ("Degradación de la señal temporal",
    ["Exactitud (%) del mejor ratio por horizonte",
     "Ratio más estable a largo plazo"]),
 ],
 "vd_instr": ["Análisis univariado de ratios financieros",
              "Dichotomous classification test",
              "Muestra pareada (n=34 quebradas + 34 sanas, 5 años)",
              "COMPUSTAT / Moody's Industrial Manual"],
 "contrib": ("REFERENCIA FUNDACIONAL HISTÓRICA. El análisis univariado de ratios (Beaver, 1966) establece el "
             "origen del campo de predicción cuantitativa del riesgo empresarial. Junto con Altman (1968), "
             "traza la línea histórica desde análisis univariado hasta ML moderno que justifica el módulo IA "
             "del IRE. Se cita en el marco teórico como evidencia de la evolución metodológica, no como "
             "antecedente empírico directo."),
},

# ─── EJE 4 ───────────────────────────────────────────────────────────────────
{"eje": 4,
 "num": "28", "title": "Microservices: The journey so far and challenges ahead",
 "authors": "Jamshidi, P. et al. (2018)",
 "journal": "IEEE Software · Q1 · JIF 4.40",
 "vi_name": "Arquitectura de\nMicroservicios",
 "vi_dims": [
   ("Principios de diseño de microservicios",
    ["Granularidad del servicio (n.° de responsabilidades)",
     "Independencia de despliegue (escala 1–5)",
     "Protocolo de comunicación inter-servicio (REST, gRPC)"]),
   ("Retos de implementación",
    ["Latencia de comunicación inter-servicio (ms)",
     "Complejidad de orquestación (n.° de servicios)",
     "Consistencia eventual de datos (CAP theorem)"]),
 ],
 "vi_instr": ["Revisión sistemática de literatura (IEEE/ACM)",
              "Análisis de 14 casos de sistemas en producción",
              "Docker / Kubernetes",
              "API Gateway (Kong, NGINX, AWS API GW)"],
 "vd_name": "Desempeño del Sistema\nDistribuido",
 "vd_dims": [
   ("Escalabilidad del sistema",
    ["Tiempo de escalado horizontal (segundos)",
     "Recursos adicionales por instancia (CPU %, RAM GB)",
     "N.° de instancias simultáneas soportadas"]),
   ("Mantenibilidad del sistema",
    ["Tiempo de despliegue de actualización (minutos)",
     "Cobertura de pruebas automatizadas (%)",
     "Tasa de fallos en pipeline de CI/CD (%)"]),
 ],
 "vd_instr": ["Revisión sistemática de literatura (IEEE/ACM)",
              "Análisis de 14 casos de sistemas en producción",
              "Docker / Kubernetes",
              "API Gateway (Kong, NGINX, AWS API GW)"],
 "contrib": ("ARTÍCULO NÚCLEO TÉCNICO. Fundamenta la arquitectura de microservicios del sistema propuesto: "
             "cada módulo (e-commerce, IA-IRE, pagos Stripe, pedidos) opera como servicio independiente "
             "comunicado vía REST API, exactamente como en la arquitectura Firebase + Supabase del sistema. "
             "Latencia, escalabilidad y tasa de fallos en CI/CD corresponden a indicadores de la dimensión "
             "Interoperabilidad de ISO/IEC 25000 (Variable Independiente — Arquitectura del sistema)."),
},

{"eje": 4,
 "num": "29",
 "title": "Large-scale machine learning systems in real-world industrial settings",
 "authors": "Swakatare, L.E. et al. (2019)",
 "journal": "ICSE Workshops (IEEE) — Actas de congreso (no revista arbitrada)",
 "note": ("Fuente catalogada erróneamente como Q1. ICSE Workshops es actas de congreso, "
          "NO una revista científica indexada en cuartiles. Citar únicamente como "
          "literatura técnica de referencia. NO presentar como fuente Q1 ante el jurado."),
 "vi_name": "Sistemas de ML a Gran\nEscala en Producción\nIndustrial",
 "vi_dims": [
   ("Escalabilidad del pipeline ML",
    ["Volumen de datos procesados (GB/hora)",
     "Tiempo de reentrenamiento del modelo (horas)",
     "N.° de modelos ML en producción simultánea"]),
   ("Confiabilidad del sistema ML",
    ["Uptime del sistema (%)",
     "Tasa de fallos de predicción en producción (%)",
     "MTTR — Tiempo medio de recuperación ante fallo"]),
 ],
 "vi_instr": ["MLOps (MLflow, Kubeflow)",
              "Monitoreo de datos (Evidently AI, WhyLabs)",
              "Docker / Kubernetes",
              "Entrevistas a ingenieros de ML en producción (n=8)"],
 "vd_name": "Calidad del Sistema ML\nen Producción",
 "vd_dims": [
   ("Precisión en producción",
    ["Delta de accuracy (drift del modelo)",
     "Degradación de F1-score por período",
     "Frecuencia de reentrenamiento necesario (días)"]),
   ("Eficiencia operativa del sistema",
    ["Costo de inferencia (S/. por 1000 predicciones)",
     "Latencia de predicción p95 (ms)",
     "Utilización de recursos (CPU %, RAM %)"]),
 ],
 "vd_instr": ["MLOps (MLflow, Kubeflow)",
              "Monitoreo de datos (Evidently AI, WhyLabs)",
              "Docker / Kubernetes",
              "Entrevistas a ingenieros de ML en producción (n=8)"],
 "contrib": ("CORREGIDO: No es Q1. Literatura técnica de referencia sobre MLOps en producción industrial. "
             "Los indicadores de data drift, uptime y MTTR son conceptos aplicables al despliegue del servicio "
             "de predicción del IRE en Firebase Functions. Se cita como referencia técnica de respaldo sobre "
             "los desafíos de ML en producción, sin declarar cuartil. NO presentar como fuente Q1 al jurado."),
},

{"eje": 4,
 "num": "30",
 "title": "Predicting risk through artificial intelligence based on machine learning algorithms",
 "authors": "Khalid, U. et al. (2022)",
 "journal": "Scientific Programming (Hindawi/Wiley) — JIF 1.89 (verificar cuartil en Scimago antes de sustentar como Q1)",
 "note": ("JIF 1.89 es bajo para Q1. En Scimago aparece como Q2 en varias categorías. "
          "Verificar en scimago.org antes de defender como Q1 ante el jurado. "
          "Si es Q2, citar sin declarar cuartil explícito. El tema encaja directamente "
          "con predicción de riesgo empresarial con ML en PYMEs."),
 "vi_name": "Algoritmos de ML para\nPredicción de Riesgo\nEmpresarial",
 "vi_dims": [
   ("Variables de entrada al modelo",
    ["Indicadores financieros históricos (n.°)",
     "Datos históricos de ventas (períodos mínimos)",
     "Variables macroeconómicas relevantes"]),
   ("Algoritmos de ML evaluados y comparados",
    ["Random Forest Classifier / Regressor",
     "Gradient Boosting (XGBoost, LightGBM)",
     "K-Nearest Neighbors (KNN) y Naive Bayes"]),
 ],
 "vi_instr": ["Python (scikit-learn, XGBoost)",
              "Validación cruzada k-fold (k=10)",
              "SHAP para interpretabilidad de predicciones",
              "Dataset de PYMEs regionales (Perú/Latinoamérica)"],
 "vd_name": "Índice de Riesgo\nEmpresarial Predicho",
 "vd_dims": [
   ("Exactitud de la predicción de riesgo",
    ["Accuracy (%)",
     "AUC-ROC",
     "MAE — Error Absoluto Medio"]),
   ("Comparativa entre algoritmos",
    ["Mejor algoritmo por métrica de evaluación",
     "Tiempo de entrenamiento (segundos)",
     "Robustez ante datos faltantes o ruidosos"]),
 ],
 "vd_instr": ["Python (scikit-learn, XGBoost)",
              "Validación cruzada k-fold (k=10)",
              "SHAP para interpretabilidad de predicciones",
              "Dataset de PYMEs regionales (Perú/Latinoamérica)"],
 "contrib": ("ARTÍCULO DE MAYOR RELEVANCIA TEMÁTICA DIRECTA. Es el antecedente más próximo al IRE: usa ML "
             "para predicción de riesgo empresarial (no quiebra) en PYMEs con datos de Latinoamérica. Random "
             "Forest y XGBoost evaluados son los mismos algoritmos candidatos del sistema propuesto. "
             "VERIFICAR cuartil en scimago.org antes de defender como Q1; si es Q2, citar sin declarar cuartil "
             "(Variable Dependiente — Predicción del IRE en PYME latinoamericana)."),
},

{"eje": 4,
 "num": "31",
 "title": "Deep learning with long short-term memory networks for financial market predictions",
 "authors": "Fischer, T. y Krauss, C. (2018)",
 "journal": "European Journal of Operational Research · Q1 · JIF 6.37",
 "vi_name": "Redes LSTM para\nPredicción de Series\nTemporales Financieras",
 "vi_dims": [
   ("Arquitectura de la red LSTM",
    ["N.° de capas LSTM apiladas",
     "N.° de unidades ocultas por capa",
     "Longitud de la ventana temporal de entrada (días)"]),
   ("Datos históricos de entrenamiento",
    ["N.° de variables de entrada (features)",
     "Período histórico utilizado para entrenamiento (años)",
     "Frecuencia de los datos (diaria, semanal, mensual)"]),
 ],
 "vi_instr": ["Python (TensorFlow / Keras)",
              "S&P 500 (500 acciones, 1992–2015)",
              "Validación walk-forward temporal",
              "Métricas financieras (Sharpe ratio, Drawdown)"],
 "vd_name": "Precisión en Predicción\nde Series Temporales",
 "vd_dims": [
   ("Exactitud predictiva del modelo",
    ["Accuracy de dirección de movimiento (%)",
     "Sharpe ratio del portafolio resultante",
     "RMSE en datos fuera de muestra"]),
   ("Comparativa con benchmarks clásicos",
    ["Mejora vs. Random Walk (% de accuracy)",
     "Mejora vs. regresión logística (ΔAUC)",
     "Retorno acumulado del modelo vs. buy-and-hold (%)"]),
 ],
 "vd_instr": ["Python (TensorFlow / Keras)",
              "S&P 500 (500 acciones, 1992–2015)",
              "Validación walk-forward temporal",
              "Métricas financieras (Sharpe ratio, Drawdown)"],
 "contrib": ("Antecedente de LSTM para series temporales financieras con JIF 6.37 Q1 sólido. Justifica por qué "
             "Random Forest fue seleccionado sobre LSTM para el IRE: el dataset de Calzatura Vilchez es "
             "pequeño y los datos son tabulares estructurados (no secuencias largas), haciendo RF más "
             "interpretable y menos propenso a sobreajuste. La validación walk-forward es replicable para "
             "evaluar el modelo del IRE con datos históricos de ventas."),
},

{"eje": 4,
 "num": "32", "title": "Challenges in deploying machine learning: A survey of case studies",
 "authors": "Paleyes, A.; Urma, R.G. y Lawrence, N.D. (2022)",
 "journal": "ACM Computing Surveys · Q1 · JIF 16.60",
 "vi_name": "Desafíos en el\nDespliegue de Modelos\nML en Producción",
 "vi_dims": [
   ("Preparación y calidad de datos",
    ["% de datos válidos en producción",
     "Tiempo de limpieza y preprocesamiento (horas)",
     "N.° de fuentes de datos integradas al pipeline"]),
   ("Despliegue e integración del modelo",
    ["Tiempo de despliegue del modelo (horas)",
     "N.° de integraciones API requeridas",
     "Compatibilidad con sistemas legados existentes"]),
   ("Monitoreo y mantenimiento en producción",
    ["Frecuencia de reentrenamiento (días)",
     "Delta de accuracy detectado (data drift)",
     "N.° de alertas automatizadas configuradas"]),
 ],
 "vi_instr": ["Revisión sistemática de 209 casos reales (PRISMA)",
              "MLflow / Kubeflow para MLOps",
              "Taxonomía de desafíos (categorización cualitativa)",
              "Docker / Kubernetes en producción"],
 "vd_name": "Éxito del Despliegue\nML en Entornos Reales",
 "vd_dims": [
   ("Fiabilidad del sistema en producción",
    ["Uptime del modelo (%)",
     "Tasa de predicciones fallidas o erróneas (%)",
     "MTTR ante incidentes de producción (horas)"]),
   ("Adopción real por usuarios del negocio",
    ["Tasa de uso efectivo del sistema ML (%)",
     "Satisfacción de usuarios internos (Likert 1–5)",
     "N.° de decisiones asistidas por ML/semana"]),
 ],
 "vd_instr": ["Revisión sistemática de 209 casos reales (PRISMA)",
              "MLflow / Kubeflow para MLOps",
              "Taxonomía de desafíos (categorización cualitativa)",
              "Docker / Kubernetes en producción"],
 "contrib": ("ARTÍCULO NÚCLEO. Identifica los 6 desafíos reales del despliegue ML en producción (datos, modelo, "
             "código, sistema, integración, monitoreo), todos presentes en el servicio de predicción del IRE "
             "vía Firebase Functions. Los indicadores de data drift, MTTR y uptime son KPIs de calidad del "
             "módulo IA del sistema. JIF 16.60 Q1 lo convierte en la fuente técnica más sólida del Eje 4 "
             "(Variable Independiente — Calidad del sistema web con IA)."),
},

{"eje": 4,
 "num": "33", "title": "Software engineering for AI-based systems: A survey",
 "authors": "Martinez-Fernandez, S. et al. (2022)",
 "journal": "ACM Transactions on Software Engineering and Methodology · Q1 · JIF 6.00",
 "vi_name": "Ingeniería de Software\npara Sistemas Basados\nen IA",
 "vi_dims": [
   ("Prácticas de ingeniería de software adaptadas a IA",
    ["N.° de prácticas de SE adaptadas para componentes ML",
     "Cobertura de pruebas de componentes IA (%)",
     "Nivel de documentación de requisitos de IA (escala)"]),
   ("Calidad intrínseca del software con IA",
    ["Deuda técnica de módulos IA (horas estimadas)",
     "N.° de defectos detectados por release",
     "Tiempo de refactorización del código IA (horas)"]),
   ("Proceso de desarrollo iterativo con IA",
    ["Duración del ciclo de experimento (días)",
     "N.° de experimentos por iteración de desarrollo",
     "Tasa de experimentos que avanzan a producción (%)"]),
 ],
 "vi_instr": ["Revisión sistemática SLR (106 estudios, 2014–2021)",
              "Análisis temático cualitativo",
              "SonarQube para métricas de calidad de código",
              "GitHub / GitLab para datos de repositorios"],
 "vd_name": "Calidad del Sistema de\nSoftware con IA",
 "vd_dims": [
   ("Fiabilidad del sistema",
    ["MTBF — Tiempo medio entre fallos (días)",
     "Tasa de fallos en producción (%)",
     "N.° de incidentes críticos por mes"]),
   ("Mantenibilidad y evolución",
    ["Tiempo para incorporar nueva funcionalidad ML (días)",
     "Complejidad ciclomática del código IA (McCabe)",
     "Esfuerzo de mantenimiento mensual (horas)"]),
 ],
 "vd_instr": ["Revisión sistemática SLR (106 estudios, 2014–2021)",
              "Análisis temático cualitativo",
              "SonarQube para métricas de calidad de código",
              "GitHub / GitLab para datos de repositorios"],
 "contrib": ("ARTÍCULO NÚCLEO. Fundamenta las prácticas de ingeniería de software para sistemas con componentes "
             "IA: pruebas, documentación técnica y gestión de la deuda técnica del módulo predictivo. "
             "Cobertura de pruebas del componente IA y MTBF corresponden a la dimensión Fiabilidad de "
             "ISO/IEC 25000. El N.° de defectos por release es un KPI de calidad del módulo IA del sistema "
             "(Variable Independiente — Calidad de software del sistema con IA)."),
},

{"eje": 4,
 "num": "34",
 "title": "Understanding and addressing quality attributes of microservices architecture: A systematic mapping study",
 "authors": "Li, S. et al. (2021)",
 "journal": "Information and Software Technology · Q1 · JIF 3.86",
 "vi_name": "Atributos de Calidad en\nArquitectura de\nMicroservicios",
 "vi_dims": [
   ("Atributos de calidad evaluados (ISO/IEC 25010)",
    ["Escalabilidad: horizontal y vertical",
     "Disponibilidad: uptime (%)",
     "Rendimiento: latencia (ms) y throughput (RPS)"]),
   ("Patrones y estrategias de solución aplicadas",
    ["N.° de patrones aplicados (circuit breaker, saga, CQRS)",
     "N.° de técnicas de monitoreo implementadas",
     "Frecuencia de despliegue (veces/semana — CI/CD)"]),
 ],
 "vi_instr": ["Systematic Mapping Study (SMS) — 60 estudios primarios",
              "Docker / Kubernetes en entornos de prueba",
              "Prometheus + Grafana para monitoreo",
              "Análisis bibliométrico de ACM DL / IEEE Xplore"],
 "vd_name": "Desempeño de la\nArquitectura de\nMicroservicios",
 "vd_dims": [
   ("Eficiencia del sistema",
    ["Latencia de API p95 (ms)",
     "Throughput máximo (requests/segundo)",
     "Tasa de errores de API 5xx (%)"]),
   ("Mantenibilidad del sistema distribuido",
    ["MTTR — Tiempo medio de resolución de incidentes",
     "Cobertura de pruebas de integración (%)",
     "Deuda técnica acumulada (horas estimadas)"]),
 ],
 "vd_instr": ["Systematic Mapping Study (SMS) — 60 estudios primarios",
              "Docker / Kubernetes en entornos de prueba",
              "Prometheus + Grafana para monitoreo",
              "Análisis bibliométrico de ACM DL / IEEE Xplore"],
 "contrib": ("ARTÍCULO NÚCLEO. Fundamenta directamente los atributos de calidad ISO/IEC 25010 del sistema: "
             "escalabilidad, disponibilidad, rendimiento y mantenibilidad. Los patrones circuit breaker y API "
             "Gateway son exactamente los utilizados en la arquitectura Firebase + Supabase del sistema "
             "propuesto. Latencia API p95, throughput y MTTR son métricas de evaluación técnica comparables "
             "con los requisitos no funcionales del sistema (Variable Independiente — ISO/IEC 25000)."),
},

# ─── EJE 5 ───────────────────────────────────────────────────────────────────
{"eje": 5,
 "num": "35", "title": "Random forests",
 "authors": "Breiman, L. (2001)",
 "journal": "Machine Learning · Q1 · JIF 7.40",
 "vi_name": "Algoritmo Random Forest\n(Ensamble de Árboles\nde Decisión)",
 "vi_dims": [
   ("Hiperparámetros del algoritmo",
    ["N.° de árboles del bosque (n_estimators)",
     "Profundidad máxima de árbol (max_depth)",
     "N.° de variables por división (max_features)"]),
   ("Proceso de ensemble por bagging",
    ["Técnica de muestreo con reemplazo (bootstrapping)",
     "Agregación de predicciones (votación / promedio)",
     "Semilla aleatoria (random_state) para reproducibilidad"]),
   ("Medición de importancia de variables",
    ["Mean Decrease Impurity — MDI",
     "Mean Decrease Accuracy — MDA",
     "SHAP values (análisis post-hoc)"]),
 ],
 "vi_instr": ["Python (scikit-learn — RandomForestClassifier/Regressor)",
              "MATLAB (TreeBagger)",
              "Validación cruzada k-fold",
              "Curvas de aprendizaje y de error vs. n_estimators"],
 "vd_name": "Exactitud y Robustez del\nModelo Predictivo",
 "vd_dims": [
   ("Exactitud de clasificación y regresión",
    ["OOB Error — Error fuera de bolsa (%)",
     "Accuracy en clasificación (%)",
     "RMSE en tareas de regresión"]),
   ("Robustez ante sobreajuste (overfitting)",
    ["Generalización en datos no vistos (test set)",
     "Balance bias-variance del modelo",
     "Estabilidad del error vs. n_estimators (curva)"]),
 ],
 "vd_instr": ["Python (scikit-learn — RandomForestClassifier/Regressor)",
              "MATLAB (TreeBagger)",
              "Validación cruzada k-fold",
              "Curvas de aprendizaje y de error vs. n_estimators"],
 "contrib": ("ARTÍCULO FUNDACIONAL ALGORÍTMICO — REFERENCIA OBLIGATORIA. Es la fuente primaria del algoritmo "
             "Random Forest utilizado en el módulo de predicción del IRE. Los hiperparámetros n_estimators, "
             "max_depth y max_features son exactamente los optimizados en el entrenamiento del modelo del IRE. "
             "OOB error y SHAP values son los instrumentos de evaluación e interpretabilidad del índice "
             "(Variable Dependiente — Algoritmo del modelo IRE)."),
},

{"eje": 5,
 "num": "36",
 "title": "A decade of agile methodologies: Towards explaining agile software development",
 "authors": "Dingsøyr, T. et al. (2012)",
 "journal": "Journal of Systems and Software · Q1 · JIF 4.34",
 "vi_name": "Metodologías Ágiles de\nDesarrollo de Software",
 "vi_dims": [
   ("Prácticas ágiles adoptadas",
    ["N.° de ceremonias Scrum implementadas (daily, retro, review)",
     "Duración del sprint (semanas)",
     "Tasa de adopción de TDD — Test-Driven Development (%)"]),
   ("Factores contextuales del equipo",
    ["Tamaño del equipo (n.° de integrantes)",
     "Experiencia en metodologías ágiles (años)",
     "Distribución geográfica del equipo (local/remoto)"]),
 ],
 "vi_instr": ["Revisión sistemática de literatura (2000–2010)",
              "Análisis bibliométrico",
              "Encuesta a equipos ágiles (n=109 proyectos)",
              "Jira / Trello para métricas de sprints"],
 "vd_name": "Desempeño del Proyecto\nde Software",
 "vd_dims": [
   ("Productividad del equipo",
    ["Velocidad del equipo (story points/sprint)",
     "Tasa de completitud de sprints (%)",
     "Déficit de entrega por sprint (%)"]),
   ("Calidad del software entregado",
    ["N.° de defectos detectados por sprint",
     "Cobertura de pruebas unitarias (%)",
     "Deuda técnica acumulada (horas)"]),
 ],
 "vd_instr": ["Revisión sistemática de literatura (2000–2010)",
              "Análisis bibliométrico",
              "Encuesta a equipos ágiles (n=109 proyectos)",
              "Jira / Trello para métricas de sprints"],
 "contrib": ("Fundamenta el uso de metodología Scrum para el desarrollo iterativo del sistema web propuesto. "
             "La velocidad del equipo (story points/sprint) y la cobertura de pruebas son métricas de gestión "
             "del proyecto de desarrollo. Se referencia en la sección de metodología de desarrollo del sistema, "
             "no en el estado del arte de e-commerce/IA/riesgo (Variable Independiente — Proceso de desarrollo)."),
},

{"eje": 5,
 "num": "37",
 "title": "Predicting sprint performance for agile software development using machine learning",
 "authors": "Salgonde, O. y Chari, K. (2021)",
 "journal": "Information Systems Management · Q1 · JIF 7.15",
 "note": ("Esta tabla NO pertenece al estado del arte central. Su tema (predicción de "
          "rendimiento de sprints ágiles) no guarda relación con e-commerce, IA comercial "
          "ni riesgo empresarial. Reubicar en ANEXO de metodología de desarrollo. "
          "NO incluir en el capítulo de Estado del Arte de la tesis."),
 "vi_name": "Variables de Entrada para\nPredicción del Desempeño\nde Sprints",
 "vi_dims": [
   ("Datos históricos del sprint",
    ["Velocidad promedio histórica (story points/sprint)",
     "N.° de items en el product backlog",
     "Capacidad del equipo disponible (horas-persona)"]),
   ("Modelos de ML aplicados al sprint",
    ["Regresión lineal múltiple (baseline)",
     "Random Forest Regressor",
     "XGBoost Regressor (gradient boosting)"]),
 ],
 "vi_instr": ["Python (scikit-learn)",
              "Jira / Azure DevOps (datos históricos de sprints reales)",
              "Validación cruzada k-fold",
              "SMOTE para balanceo de clases"],
 "vd_name": "Desempeño Predicho\ndel Sprint",
 "vd_dims": [
   ("Exactitud de la predicción",
    ["RMSE — Raíz del Error Cuadrático Medio",
     "MAE — Error Absoluto Medio",
     "R² — Coeficiente de determinación"]),
   ("Utilidad práctica del modelo",
    ["Reducción del déficit de entrega del sprint (%)",
     "% de sprints con predicción dentro del margen ±10%",
     "Tiempo ahorrado en planificación (horas/sprint)"]),
 ],
 "vd_instr": ["Python (scikit-learn)",
              "Jira / Azure DevOps (datos históricos de sprints reales)",
              "Validación cruzada k-fold",
              "SMOTE para balanceo de clases"],
 "contrib": ("REUBICADO AL ANEXO DE METODOLOGÍA. Predice el rendimiento de sprints ágiles con ML; no guarda "
             "relación con e-commerce, IA comercial ni riesgo empresarial. Se incluye en este documento solo "
             "como referencia metodológica complementaria para el plan de gestión del proyecto. NO incluir en "
             "el capítulo de Estado del Arte de la tesis; puede citarse en la sección de gestión del proyecto."),
},

{"eje": 5,
 "num": "38",
 "title": "Hybrid software development approaches in practice: A European perspective",
 "authors": "Kuhrmann, M. et al. (2018)",
 "journal": "IEEE Software · Q1 · JIF 4.40",
 "vi_name": "Enfoque Híbrido de\nDesarrollo (Ágil +\nTradicional Planificado)",
 "vi_dims": [
   ("Configuración del proceso híbrido",
    ["% de prácticas ágiles adoptadas en el proceso",
     "% de prácticas tradicionales (waterfall) mantenidas",
     "N.° de fases formales definidas"]),
   ("Contexto organizacional de aplicación",
    ["Tamaño de la organización (n.° de empleados)",
     "Dominio de la aplicación desarrollada (sector)",
     "N.° de años de experiencia con enfoques híbridos"]),
 ],
 "vi_instr": ["Encuesta a organizaciones europeas (n=732 respuestas válidas)",
              "Análisis estadístico descriptivo y factorial",
              "Clustering (k-means) de perfiles de proceso",
              "R / SPSS para análisis multivariado"],
 "vd_name": "Efectividad del Proceso\nde Desarrollo Híbrido",
 "vd_dims": [
   ("Éxito del proyecto",
    ["Entrega en plazo (% de proyectos a tiempo)",
     "Desviación de presupuesto (%)",
     "Satisfacción del cliente con el producto (Likert)"]),
   ("Flexibilidad y control del proceso",
    ["Capacidad de respuesta a cambios (escala 1–5)",
     "Nivel de documentación generada (escala)",
     "Tasa de riesgos identificados y mitigados (%)"]),
 ],
 "vd_instr": ["Encuesta a organizaciones europeas (n=732 respuestas válidas)",
              "Análisis estadístico descriptivo y factorial",
              "Clustering (k-means) de perfiles de proceso",
              "R / SPSS para análisis multivariado"],
 "contrib": ("Justifica el enfoque híbrido ágil-planificado adoptado en el desarrollo del sistema (Scrum + fases "
             "formales de tesis). La entrega en plazo y la satisfacción del cliente son métricas de éxito del "
             "proyecto de desarrollo. Se referencia en la sección de metodología de desarrollo; complementa "
             "al artículo 36 en la justificación del proceso de implementación del sistema propuesto."),
},

{"eje": 5,
 "num": "39", "title": "Prescriptive analytics: Literature review and research challenges",
 "authors": "Lepenioti, K. et al. (2020)",
 "journal": "International Journal of Information Management · Q1 · JIF 18.70",
 "vi_name": "Analítica Prescriptiva\npara Optimización de\nDecisiones",
 "vi_dims": [
   ("Tipos de técnicas de analítica prescriptiva",
    ["Optimización matemática (PL, PNL, algoritmos evolutivos)",
     "Simulación (Monte Carlo, Discrete Event Simulation — DES)",
     "Reglas de negocio y sistemas expertos basados en IA"]),
   ("Nivel de automatización de la prescripción",
    ["Grado de autonomía del sistema (escala 1–5)",
     "N.° de decisiones prescritas automáticamente por período",
     "Tiempo de implementación de la prescripción (ms)"]),
   ("Calidad de los datos de entrada",
    ["Completitud de los datos disponibles (%)",
     "Frescura/lag de los datos (horas de retraso)",
     "N.° de fuentes de datos integradas al modelo"]),
 ],
 "vi_instr": ["Revisión sistemática (PRISMA) — 209 publicaciones (Scopus/WoS)",
              "Taxonomía de técnicas (categorización inductiva)",
              "Python / IBM CPLEX para optimización",
              "Framework de madurez de analítica prescriptiva"],
 "vd_name": "Valor de las Decisiones\nPrescritas por el Sistema",
 "vd_dims": [
   ("Calidad y aceptación de la prescripción",
    ["Tasa de adopción de prescripciones por usuarios (%)",
     "Mejora de KPIs objetivo tras implementar prescripción (%)",
     "ROI de las decisiones tomadas con prescripción (%)"]),
   ("Eficiencia operativa generada",
    ["Reducción de costos por optimización (%)",
     "Reducción del tiempo de ciclo de decisión (horas)",
     "N.° de alternativas evaluadas automáticamente"]),
 ],
 "vd_instr": ["Revisión sistemática (PRISMA) — 209 publicaciones (Scopus/WoS)",
              "Taxonomía de técnicas (categorización inductiva)",
              "Python / IBM CPLEX para optimización",
              "Framework de madurez de analítica prescriptiva"],
 "contrib": ("ARTÍCULO NÚCLEO. Fundamenta la analítica prescriptiva del módulo IA: el sistema no solo predice "
             "el IRE, sino que prescribe acciones correctivas (ajustar stock, modificar pedidos a proveedores, "
             "revisar precios). El ROI de decisiones prescritas y la tasa de adopción de prescripciones son "
             "KPIs de impacto gerencial del módulo de recomendaciones del sistema "
             "(Variable Dependiente — Analítica prescriptiva basada en IRE)."),
},

{"eje": 5,
 "num": "40", "title": "Big data analytics in operations management",
 "authors": "Choi, T.M.; Wallace, S.W. y Wang, Y. (2018)",
 "journal": "Production and Operations Management · Q1 · JIF 4.60",
 "vi_name": "Analítica de Big Data en\nGestión de Operaciones",
 "vi_dims": [
   ("Capacidades de los datos disponibles",
    ["Volumen de datos de operaciones (TB)",
     "Velocidad de actualización de datos (lotes/hora)",
     "Variedad y fuentes de datos integradas (n.°)"]),
   ("Modelos analíticos aplicados a operaciones",
    ["Modelos de optimización de inventario (LP, MIP)",
     "Modelos de pronóstico de demanda (ARIMA, ML)",
     "Modelos de gestión de la cadena de suministro"]),
 ],
 "vi_instr": ["Revisión de literatura operacional (POM, MSOM)",
              "Modelos de optimización (Python: Scipy, PuLP)",
              "Análisis de series temporales (SARIMA, Prophet)",
              "Simulación de cadena de suministro"],
 "vd_name": "Desempeño de la Gestión\nde Operaciones",
 "vd_dims": [
   ("Eficiencia de la cadena de suministro",
    ["Reducción de costos de inventario (%)",
     "Mejora del nivel de servicio (fill rate %)",
     "Reducción del tiempo de ciclo operativo (%)"]),
   ("Capacidad de respuesta al mercado",
    ["Tiempo de respuesta a variaciones de demanda (días)",
     "Tasa de fulfillment de pedidos (%)",
     "N.° de quiebres de stock evitados por período"]),
 ],
 "vd_instr": ["Revisión de literatura operacional (POM, MSOM)",
              "Modelos de optimización (Python: Scipy, PuLP)",
              "Análisis de series temporales (SARIMA, Prophet)",
              "Simulación de cadena de suministro"],
 "contrib": ("ARTÍCULO NÚCLEO. Sustenta Big Data en gestión de operaciones con énfasis directo en inventario y "
             "demanda, los dos subíndices más pesados del IRE (stock 40% + demanda 25% = 65%). Los indicadores "
             "de quiebres de stock evitados, fill rate y tiempo de respuesta a variaciones de demanda son los "
             "KPIs operativos exactos del sistema propuesto, medibles antes y después de la implementación "
             "(Variable Dependiente — Subíndices de riesgo de stock y demanda del IRE)."),
},

{"eje": 5,
 "num": "41", "title": "AI lifecycle models need to be revised: An exploratory study in Fintech",
 "authors": "Haakman, M. et al. (2021)",
 "journal": "Empirical Software Engineering · Q1 · JIF 3.50",
 "vi_name": "Ciclo de Vida del Modelo\nde IA — CRISP-ML(Q)",
 "vi_dims": [
   ("Fases del ciclo de vida ML",
    ["Comprensión del negocio y de los datos",
     "Preparación y transformación de datos",
     "Modelado, entrenamiento y selección",
     "Evaluación y validación del modelo",
     "Despliegue y monitoreo en producción"]),
   ("Métricas de calidad por fase del ciclo",
    ["Completitud de datos mínima requerida (%)",
     "Accuracy mínimo aceptable por el negocio (threshold)",
     "Frecuencia de reentrenamiento necesario (días)"]),
   ("Prácticas de MLOps adoptadas",
    ["N.° de experimentos registrados con trazabilidad",
     "Versionado de modelos (MLflow / DVC)",
     "Automatización del pipeline CI/CD ML (%)"]),
 ],
 "vi_instr": ["CRISP-ML(Q) framework de calidad",
              "MLflow para versionado y registro de experimentos",
              "Entrevistas a equipos de ML en Fintech (n=17 entrevistados)",
              "Análisis temático cualitativo (codificación inductiva)"],
 "vd_name": "Calidad y Madurez del\nSistema de ML en\nProducción",
 "vd_dims": [
   ("Calidad del modelo en producción",
    ["Data drift detectado (delta de accuracy)",
     "F1-score en producción vs. F1-score en evaluación",
     "N.° de reentrenamientos necesarios por mes"]),
   ("Madurez del proceso de desarrollo ML",
    ["Nivel de madurez MLOps alcanzado (escala 1–5)",
     "N.° de fases con criterios de calidad definidos formalmente",
     "Esfuerzo de mantenimiento del sistema (horas/mes)"]),
 ],
 "vd_instr": ["CRISP-ML(Q) framework de calidad",
              "MLflow para versionado y registro de experimentos",
              "Entrevistas a equipos de ML en Fintech (n=17 entrevistados)",
              "Análisis temático cualitativo (codificación inductiva)"],
 "contrib": ("ARTÍCULO NÚCLEO. Establece CRISP-ML(Q) como el ciclo de vida del módulo IA del sistema propuesto. "
             "Las 5 fases (comprensión del negocio, datos, modelado, evaluación, despliegue) son exactamente "
             "las etapas del desarrollo del servicio de predicción del IRE. El nivel de madurez MLOps alcanzado "
             "es la métrica de calidad del proceso de desarrollo del módulo IA, evaluable en la tesis "
             "(Variable Dependiente — Ciclo de vida y calidad del modelo IRE)."),
},

{"eje": 5,
 "num": "42",
 "title": "Information technology and organizational performance: An integrative model of IT business value",
 "authors": "Melville, N.; Kraemer, K. y Gurbaxani, V. (2004)",
 "journal": "MIS Quarterly · Q1 · JIF 12.50",
 "vi_name": "Inversión y Capacidades\nen Tecnología de\nInformación (TI)",
 "vi_dims": [
   ("Recursos tecnológicos de TI",
    ["Inversión en TI (S/. por año)",
     "N.° de sistemas de información implementados",
     "Capacidad de infraestructura tecnológica (escala)"]),
   ("Capacidades humanas en TI",
    ["Nivel de competencias digitales del personal (escala 1–5)",
     "N.° de personal dedicado a TI",
     "Horas de capacitación en TI por empleado/año"]),
   ("Recursos organizacionales complementarios",
    ["Madurez de los procesos de negocio (escala)",
     "Cultura organizacional orientada a datos (escala)",
     "N.° de socios tecnológicos estratégicos"]),
 ],
 "vi_instr": ["Modelo conceptual integrador (meta-análisis de literatura)",
              "Revisión de 200+ estudios (1985–2003)",
              "Regresión múltiple jerárquica",
              "Encuesta a organizaciones Fortune 500 (EE.UU.)"],
 "vd_name": "Desempeño\nOrganizacional\nImpulsado por TI",
 "vd_dims": [
   ("Desempeño de los procesos internos",
    ["Eficiencia de procesos mejorada (%)",
     "Reducción de tiempos de ciclo (%)",
     "Reducción de costos operativos (%)"]),
   ("Desempeño competitivo en el mercado",
    ["Crecimiento de ingresos (%)",
     "Cuota de mercado ganada (%)",
     "ROI de la inversión en TI (%)"]),
 ],
 "vd_instr": ["Modelo conceptual integrador (meta-análisis de literatura)",
              "Revisión de 200+ estudios (1985–2003)",
              "Regresión múltiple jerárquica",
              "Encuesta a organizaciones Fortune 500 (EE.UU.)"],
 "contrib": ("Justifica el valor económico de la inversión en TI para Calzatura Vilchez como PYME. Los "
             "indicadores de ROI de TI, reducción de costos operativos y crecimiento de ingresos son las "
             "métricas de impacto económico esperadas de la propuesta. El modelo IT-Business Value es el marco "
             "conceptual que respalda la rentabilidad de implementar el sistema web con IA "
             "(Marco teórico — Justificación de la inversión tecnológica propuesta)."),
},

# ─── TABLA 43 — NUEVA (obligatoria) ──────────────────────────────────────────
{"eje": 5,
 "num": "43",
 "title": "Revolutionizing cross-border e-commerce: A deep dive into AI and big data-driven innovations for the straw hat industry",
 "authors": "Dai, J.; Mao, X.; Wu, P.; Zhou, H. y Cao, L. (2024)",
 "journal": "PLOS ONE · Q1 · JIF 3.75 · DOI: 10.1371/journal.pone.0305639",
 "note": ("TABLA NUEVA — Agregada como fuente obligatoria. Es la única fuente del conjunto "
          "que une simultáneamente: e-commerce + IA + Big Data + pequeña empresa + demanda "
          "+ decisiones comerciales. Citar en el Eje 2 o como fuente transversal."),
 "vi_name": "IA y Big Data Aplicados\nal E-commerce de\nPequeñas Empresas",
 "vi_dims": [
   ("Innovaciones tecnológicas en e-commerce",
    ["N.° de tecnologías IA implementadas en la plataforma",
     "Volumen de datos procesados por el sistema (GB/día)",
     "Tasa de adopción de herramientas de Big Data (%)"]),
   ("Modelos de predicción de demanda con IA",
    ["Accuracy del modelo predictivo (%)",
     "MAE del pronóstico de demanda",
     "Tiempo de respuesta del sistema IA (ms)"]),
   ("Personalización y decisiones comerciales",
    ["Tasa de personalización de recomendaciones (%)",
     "Precisión de la segmentación de clientes (%)",
     "N.° de decisiones comerciales asistidas por IA/semana"]),
 ],
 "vi_instr": ["Revisión sistemática de literatura (Scopus, WoS)",
              "Modelos de ML y Big Data (Python, scikit-learn)",
              "Análisis de datos de plataformas e-commerce reales",
              "PLOS ONE — Validación empírica en pequeña empresa"],
 "vd_name": "Desempeño Comercial y\nCompetitivo del\nE-commerce",
 "vd_dims": [
   ("Eficiencia comercial",
    ["Crecimiento de ventas en línea (%)",
     "Reducción del tiempo de ciclo de pedido (días)",
     "Tasa de fulfillment de pedidos (%)"]),
   ("Capacidad de respuesta al mercado",
    ["Tiempo de respuesta a variaciones de demanda (días)",
     "N.° de quiebres de stock evitados por período",
     "ROI de la implementación IA en e-commerce (%)"]),
 ],
 "vd_instr": ["Revisión sistemática de literatura (Scopus, WoS)",
              "Modelos de ML y Big Data (Python, scikit-learn)",
              "Análisis de datos de plataformas e-commerce reales",
              "PLOS ONE — Validación empírica en pequeña empresa"],
 "contrib": ("ARTÍCULO TRANSVERSAL OBLIGATORIO — EL MÁS ALINEADO AL TÍTULO DE LA TESIS. Une simultáneamente: "
             "e-commerce + IA + Big Data + pequeña empresa artesanal + demanda + decisiones comerciales, "
             "exactamente el alcance de la propuesta para Calzatura Vilchez. Accuracy del modelo, crecimiento "
             "de ventas en línea y quiebres de stock evitados son KPIs directamente comparables con los del "
             "sistema propuesto. PLOS ONE Q1 JIF 3.75 lo valida como fuente revisada por pares "
             "(Transversal VI-VD — Evidencia empírica más próxima al objeto de estudio)."),
},

]  # end ARTICLES

# ══════════════════════════════════════════════════════════════════════════════
#  EVIDENCIA FUNCIONAL — TRAZABILIDAD INDICADOR → SISTEMA IMPLEMENTADO
# ══════════════════════════════════════════════════════════════════════════════

SYSTEM_EVIDENCE = {
    "01": "Portal web React 19 (calzatura-vilchez) + app Flutter (calzatura-vilchez-mobile) con catálogo, carrito, checkout, panel admin/trabajador y BFF (bff/server.cjs) sobre Supabase PostgreSQL.",
    "02": "Canal digital operativo: pedidos online (Stripe/contraentrega), ventas presenciales en ventasDiarias y continuidad comercial documentada en SRS-Calzatura-Vilchez.md.",
    "03": "Desarrollo iterativo con CI (GitHub Actions), despliegue en Render (BFF + ai-service) y releases móviles APK; adopción medible por módulos activos en producción.",
    "04": "Transformación del modelo presencial a omnicanal: catálogo público (domains/productos), tiendas físicas (domains/público) y ventas integradas vía BFF.",
    "05": "Módulo de ventas web y móvil (AdminSalesPage/StaffSalesPage) con registro diario, documentos y devoluciones; conversión medible en pedidos y ventasDiarias.",
    "06": "Catálogo con filtros, favoritos (tabla favoritos vía BFF /favorites), carrito en sessionStorage y checkout; métricas de conversión extraíbles de pedidos pagados.",
    "07": "Cadena de valor digitalizada: productos → stock por talla/color → pedidos → fulfillment; integración API REST entre frontend, BFF, Supabase y Stripe.",
    "08": "SKUs publicados en canal online (tabla productos + productoCodigos); punto de venta digital con panel de ventas admin/staff y búsqueda por marca/código.",
    "09": "Servicio IA FastAPI (ai-service/main.py) con endpoints /api/predict/demand y /api/predict/combined; métricas MAE/accuracy en modelo_estado (Supabase).",
    "10": "Panel ejecutivo admin consume predicciones IA; recomendaciones de stock y demanda visibles en dashboard web (domains/administradores).",
    "11": "IRE (compute_ire en models/risk.py) asiste decisiones gerenciales: alertas de stock, riesgo por dimensión y nivel Bajo/Moderado/Alto/Crítico.",
    "12": "Pipeline analítico unifica ventasDiarias + pedidos pagados (services/supabase_client.py); modelo RandomForestRegressor en models/demand.py.",
    "13": "Evaluación con MAE/RMSE en predicción de demanda; comparativa RF vs. métodos clásicos documentada en docs/06-ia/modelo-ia.md.",
    "14": "Forecast por producto con horizonte configurable; gráfico semanal de demanda en panel admin; datos históricos desde Supabase.",
    "15": "Módulo ventas con búsqueda DNI (BFF lookupDni.cjs), rangos de precio, historial y documentos PDF; paridad web-móvil vía panel_bff_api.dart.",
    "16": "Innovación IA integrada al e-commerce: predicción + IRE + alertas sin módulo aislado; ciclo CRISP-ML aplicado en ai-service.",
    "17": "Variables de mercado (categoría, marca, destacado, estacionalidad) incorporadas como features del RandomForest en build_daily_sales.",
    "18": "Decisión técnica documentada: RF tabular (scikit-learn 1.5) en lugar de DL/LSTM por tamaño limitado del dataset PYME.",
    "19": "Antecedente algorítmico: RandomForestRegressor seleccionado; métricas AUC/accuracy no aplican a quiebra sino a riesgo operativo IRE.",
    "20": "Comparativa metodológica: datos tabulares estructurados (no imágenes de ratios); justifica RF sobre CNN para Calzatura Vilchez.",
    "21": "Revisión de algoritmos ML respalda elección de RF/XGBoost; IRE mide riesgo comercial-operativo, no default financiero corporativo.",
    "22": "Series temporales de ventas por producto/día construidas en ai-service; horizonte de predicción alineado a gestión de stock.",
    "23": "Marco teórico ISO/IEC 25010 en SRS; no evidencia directa IRE — referencia metodológica de calidad de software.",
    "24": "Datos no textuales: ventas, stock y pedidos estructurados en Supabase; NLP no requerido en el pipeline actual.",
    "25": "Variables IRE calibradas: riesgo_stock (40%), riesgo_ingresos (35%), riesgo_demanda (25%) en models/risk_metadata.py.",
    "26": "Altman Z-Score como antecedente histórico; umbral conceptual análogo al nivel Crítico del IRE (score ≥ 75).",
    "27": "Evolución metodológica univariado → compuesto → ML; IRE es índice compuesto ponderado implementado en producción.",
    "28": "Arquitectura microservicios: frontend React, BFF Node (Render), ai-service FastAPI (Render), Supabase, Firebase Auth, Stripe Functions.",
    "29": "ai-service desplegado en Docker/Render con caché TTL, persistencia modelo_estado y reentrenamiento bajo demanda; referencia técnica MLOps.",
    "30": "IRE implementado con Random Forest y validación sobre datos reales de PYME peruana (Huancayo); tabla ire_historial persiste series.",
    "31": "LSTM descartado por dataset pequeño; validación walk-forward replicable con ventas históricas de ventasDiarias y pedidos.",
    "32": "Despliegue ML en ai-service Render: auth Bearer/Firebase admin, caché, rate limiting, persistencia post-reinicio en Supabase.",
    "33": "Pruebas pytest (ai-service), Vitest (frontend), E2E Playwright, SonarQube en CI; documentación en docs/03-sdd y docs/06-ia.",
    "34": "Atributos ISO/IEC 25010 verificables: latencia BFF, uptime Render, CI/CD GitHub Actions, cobertura de tests automatizados.",
    "35": "Algoritmo Breiman (2001) implementado: RandomForestRegressor con n_estimators, max_depth y max_features en models/demand.py.",
    "36": "Metodología ágil/TDD en repositorio; pruebas unitarias y de integración en cada módulo (web, BFF, ai-service, móvil).",
    "37": "Anexo metodológico: proceso de desarrollo documentado en docs/; no evidencia directa del estado del arte central.",
    "38": "Analítica descriptiva en panel admin: KPIs ventas, ganancia (admin), unidades (staff), gráficos de demanda semanal.",
    "39": "Alertas de stock prescriptivas: productos con fecha de quiebre estimada; acciones correctivas sugeridas en panel ejecutivo.",
    "40": "Big Data operacional: fusión ventas presenciales + online, movimientos de stock y catálogo; procesamiento en pandas/ai-service.",
    "41": "Ciclo CRISP-ML(Q): entrenamiento → despliegue → monitoreo drift → reentrenamiento; metadatos en modelo_estado e ire_historial.",
    "42": "ROI medible: reducción de quiebres de stock, trazabilidad de ventas, automatización de pedidos; valor de TI justificado en SRS.",
    "43": "Alineación total: e-commerce React/Flutter + IA FastAPI + Big Data Supabase + PYME calzado; Dai et al. (2024) como evidencia transversal.",
}

SYSTEM_INSTRUMENT = {
    "01": "↳ Sistema: calzatura-vilchez (React) + calzatura-vilchez-mobile (Flutter) + bff/server.cjs",
    "02": "↳ Sistema: pedidos (domains/pedidos) + ventasDiarias (domains/ventas) + Stripe",
    "03": "↳ Sistema: GitHub Actions CI + Render deploy + docker-compose.yml",
    "04": "↳ Sistema: domains/productos + domains/público + checkout omnicanal",
    "05": "↳ Sistema: panel ventas web/móvil (AdminSalesPage, StaffSalesPage)",
    "06": "↳ Sistema: carrito (sessionStorage) + favoritos BFF + Google Analytics (opcional)",
    "07": "↳ Sistema: BFF REST + Supabase RLS + Stripe Cloud Functions",
    "08": "↳ Sistema: catálogo productos + productoCodigos + panel ventas",
    "09": "↳ Sistema: ai-service/main.py — GET /api/predict/demand, /api/predict/combined",
    "10": "↳ Sistema: panel ejecutivo admin + predicciones IA consumidas vía BFF",
    "11": "↳ Sistema: models/risk.py — compute_ire(), niveles y dimensiones",
    "12": "↳ Sistema: services/supabase_client.py + models/demand.py (RF)",
    "13": "↳ Sistema: métricas MAE/RMSE en entrenamiento; docs/06-ia/modelo-ia.md",
    "14": "↳ Sistema: predict_demand() + gráfico semanal en panel admin",
    "15": "↳ Sistema: BFF lookupDni.cjs + panel_sales_page.dart + sales_register_logic.dart",
    "16": "↳ Sistema: ai-service integrado al e-commerce (no módulo aislado)",
    "17": "↳ Sistema: features categoría/marca/destacado en pipeline demanda",
    "18": "↳ Sistema: scikit-learn RF; sin TensorFlow/Keras en producción",
    "19": "↳ Sistema: RandomForestRegressor — riesgo operativo IRE, no quiebra",
    "20": "↳ Sistema: datos tabulares Supabase; sin CNN en pipeline",
    "21": "↳ Sistema: RF seleccionado; comparativa documentada en modelo-ia.md",
    "22": "↳ Sistema: build_daily_sales_by_product() — series temporales",
    "23": "↳ Sistema: ISO/IEC 25010 en SRS — marco teórico, no IRE directo",
    "24": "↳ Sistema: tablas ventasDiarias, pedidos, productos (no NLP)",
    "25": "↳ Sistema: models/risk_metadata.py — pesos y variables IRE",
    "26": "↳ Sistema: umbral IRE Crítico (≥75) análogo conceptual a Z-Score",
    "27": "↳ Sistema: IRE compuesto ponderado en compute_ire()",
    "28": "↳ Sistema: React + BFF Node + ai-service FastAPI + Supabase",
    "29": "↳ Sistema: ai-service Docker/Render + caché + modelo_estado",
    "30": "↳ Sistema: IRE + RF sobre datos PYME; ire_historial en Supabase",
    "31": "↳ Sistema: validación walk-forward con ventas históricas reales",
    "32": "↳ Sistema: ai-service producción Render — auth, caché, persistencia",
    "33": "↳ Sistema: pytest + Vitest + Playwright E2E + SonarQube CI",
    "34": "↳ Sistema: latencia BFF, uptime Render, GitHub Actions pipeline",
    "35": "↳ Sistema: sklearn.ensemble.RandomForestRegressor en demand.py",
    "36": "↳ Sistema: tests automatizados en web, BFF, ai-service y móvil",
    "37": "↳ Sistema: docs/ metodología — anexo, no estado del arte central",
    "38": "↳ Sistema: KPIs panel admin (ventas, ganancia, demanda semanal)",
    "39": "↳ Sistema: alertas stock con fecha quiebre en panel ejecutivo",
    "40": "↳ Sistema: pandas fusiona ventasDiarias + pedidos en ai-service",
    "41": "↳ Sistema: CRISP-ML — modelo_estado + ire_historial + reentrenamiento",
    "42": "↳ Sistema: trazabilidad ventas/pedidos + automatización checkout",
    "43": "↳ Sistema: e-commerce + IA + Big Data — stack completo operativo",
}

CONTRIB_FIXES = {
    "Firebase + Supabase": "React + BFF Node + Supabase PostgreSQL + ai-service FastAPI",
    "Firebase Functions": "ai-service FastAPI desplegado en Render (Docker)",
    "vía Firebase Functions": "vía ai-service FastAPI (Render)",
}


def enrich_article(art):
    num = art["num"]
    out = dict(art)
    if num in SYSTEM_EVIDENCE:
        out["evidencia"] = SYSTEM_EVIDENCE[num]
    if num in SYSTEM_INSTRUMENT:
        out["vi_instr"] = list(out["vi_instr"]) + [SYSTEM_INSTRUMENT[num]]
    if "contrib" in out:
        for old, new in CONTRIB_FIXES.items():
            out["contrib"] = out["contrib"].replace(old, new)
    return out


ENRICHED_ARTICLES = [enrich_article(a) for a in ARTICLES]

# ══════════════════════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()
sec = doc.sections[0]
sec.top_margin    = Inches(0.60)
sec.bottom_margin = Inches(0.60)
sec.left_margin   = Inches(0.65)
sec.right_margin  = Inches(0.65)

# Cover
_para(doc, "TABLAS DEL ESTADO DEL ARTE — 43 ARTÍCULOS CIENTÍFICOS",
      14, True, WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=4)
_para(doc,
      "Tesis: Propuesta de implementación de un sistema web de comercio electrónico con "
      "inteligencia artificial para la predicción del riesgo empresarial — Calzatura Vilchez",
      9, False, WD_ALIGN_PARAGRAPH.CENTER, italic=True, sb=0, sa=2)
_para(doc,
      "Autor: Piero Vilchez  |  Universidad Continental  |  "
      "Variables · Dimensiones · Indicadores · Instrumentos · Evidencia funcional",
      9, False, WD_ALIGN_PARAGRAPH.CENTER, italic=True, sb=0, sa=6)

# Correction summary box
_para(doc, "CORRECCIONES APLICADAS EN ESTA VERSIÓN", 9, True,
      WD_ALIGN_PARAGRAPH.LEFT, "C00000", sb=4, sa=2)
corrections = [
    "Artículo 29: Eliminada etiqueta Q1. ICSE Workshops es actas de congreso, no revista arbitrada.",
    "Artículo 30: Advertencia de cuartil. Scientific Programming JIF 1.89 requiere verificación en Scimago antes de sustentar como Q1.",
    "Artículo 23: Advertencia de reposicionamiento. Tabla reubicada al marco teórico; no presentar como evidencia directa del IRE.",
    "Artículo 37: Advertencia de reubicación. No pertenece al estado del arte central; mover a Anexo de metodología de desarrollo.",
    "Artículo 43: TABLA NUEVA obligatoria. Dai et al. (2024) — PLOS ONE Q1. Une e-commerce + IA + Big Data + pequeña empresa.",
    "Todas las tablas incluyen evidencia funcional trazada al sistema implementado (web React, BFF, Supabase, ai-service, app Flutter).",
    "Referencias arquitectónicas actualizadas: React + BFF Node + Supabase + ai-service FastAPI (Render), sin Firestore.",
]
for c in corrections:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(c)
    _font(r, 8)

doc.add_paragraph()

# Articles grouped by eje
eje_labels = {
    1: "COMERCIO ELECTRÓNICO, RETAIL DIGITAL Y TRANSFORMACIÓN COMERCIAL",
    2: "INTELIGENCIA ARTIFICIAL Y ANALÍTICA PARA TOMA DE DECISIONES",
    3: "PREDICCIÓN, FORECASTING Y RIESGO EMPRESARIAL COMERCIAL-OPERATIVO",
    4: "INGENIERÍA, ARQUITECTURA, CALIDAD E INTEROPERABILIDAD DEL SISTEMA CON IA",
    5: "METODOLOGÍA, ALGORITMOS Y FUENTES COMPLEMENTARIAS",
}

current_eje = 0
for art in ENRICHED_ARTICLES:
    if art.get("eje", 0) != current_eje:
        current_eje = art["eje"]
        add_eje_header(doc, current_eje, eje_labels[current_eje])
    add_article(doc, art)

# ── Matriz de trazabilidad funcional ──
doc.add_page_break()
_para(doc, "ANEXO — MATRIZ DE TRAZABILIDAD FUNCIONAL (43 ARTÍCULOS)",
      12, True, WD_ALIGN_PARAGRAPH.CENTER, sb=8, sa=4)
_para(doc,
      "Relación entre cada artículo del estado del arte y el componente operativo "
      "del sistema Calzatura Vilchez implementado y desplegado.",
      9, False, WD_ALIGN_PARAGRAPH.CENTER, italic=True, sb=0, sa=6)

trace_tbl = doc.add_table(rows=1 + len(ENRICHED_ARTICLES), cols=4)
trace_tbl.style = "Table Grid"
trace_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
_widths(trace_tbl, [0.45, 2.20, 1.10, 3.50])
for i, h in enumerate(["N.°", "Artículo", "Eje", "Evidencia funcional en sistema"]):
    c = trace_tbl.cell(0, i)
    _cell(c, h, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, color="FFFFFF")
    _shade(c, "1F4E79")

for row_idx, art in enumerate(ENRICHED_ARTICLES, start=1):
    _cell(trace_tbl.cell(row_idx, 0), art["num"], size=7.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell(trace_tbl.cell(row_idx, 1), art["title"][:72] + ("…" if len(art["title"]) > 72 else ""),
          size=7, align=WD_ALIGN_PARAGRAPH.LEFT)
    _cell(trace_tbl.cell(row_idx, 2), str(art.get("eje", "")), size=7.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell(trace_tbl.cell(row_idx, 3), art.get("evidencia", "—"), size=7)
    if row_idx % 2 == 0:
        for col in range(4):
            _shade(trace_tbl.cell(row_idx, col), "F8F9F9")

# ── Validación ──
empty_cells = 0
missing_evidence = []
for art in ENRICHED_ARTICLES:
    if not art.get("evidencia"):
        missing_evidence.append(art["num"])
    for dim_name, inds in art["vi_dims"] + art["vd_dims"]:
        if not dim_name or not inds:
            empty_cells += 1
        for ind in inds:
            if not ind or not ind.strip():
                empty_cells += 1
    if not art.get("vi_instr") or not art.get("vd_instr") or not art.get("contrib"):
        empty_cells += 1

doc.save(OUT)
print(f"Documento generado: {OUT}")
print(f"Artículos: {len(ENRICHED_ARTICLES)}")
print(f"Celdas vacías detectadas: {empty_cells}")
print(f"Artículos sin evidencia: {missing_evidence or 'ninguno'}")
assert len(ENRICHED_ARTICLES) == 43, "Deben ser exactamente 43 artículos"
assert empty_cells == 0, f"Hay {empty_cells} celdas/campos vacíos"
assert not missing_evidence, f"Faltan evidencias en artículos: {missing_evidence}"
