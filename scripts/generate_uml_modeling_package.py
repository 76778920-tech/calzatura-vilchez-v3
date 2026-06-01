from __future__ import annotations

import shutil
import textwrap
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "artifacts" / "diagramas"
PNG_DIR = OUT_DIR / "png"
MMD_DIR = OUT_DIR / "mermaid"
DOC_DIR = ROOT / "artifacts" / "documentos"
DOWNLOADS = Path.home() / "Downloads"

DOCX_NAME = "Documento_4_Diagramas_UML_Modelado_Calzatura_Vilchez_COMPLETO.docx"
ZIP_NAME = "Diagramas_UML_Modelado_Calzatura_Vilchez.zip"


COLORS = {
    "bg": "#F8FAFC",
    "ink": "#0F172A",
    "muted": "#475569",
    "line": "#334155",
    "blue": "#DBEAFE",
    "blue_line": "#2563EB",
    "green": "#DCFCE7",
    "green_line": "#16A34A",
    "amber": "#FEF3C7",
    "amber_line": "#D97706",
    "red": "#FEE2E2",
    "red_line": "#DC2626",
    "violet": "#EDE9FE",
    "violet_line": "#7C3AED",
    "gray": "#E2E8F0",
    "white": "#FFFFFF",
}


@dataclass(frozen=True)
class Diagram:
    key: str
    title: str
    purpose: str
    elements: str
    evidence: str
    image: Path
    mermaid: Path


def ensure_dirs() -> None:
    PNG_DIR.mkdir(parents=True, exist_ok=True)
    MMD_DIR.mkdir(parents=True, exist_ok=True)
    DOC_DIR.mkdir(parents=True, exist_ok=True)


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


F_TITLE = font(42, True)
F_H2 = font(28, True)
F_H3 = font(23, True)
F_TEXT = font(21)
F_SMALL = font(17)
F_SMALL_BOLD = font(17, True)


def text_size(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont) -> tuple[int, int]:
    if not text:
        return 0, 0
    bbox = draw.textbbox((0, 0), text, font=fnt)
    return bbox[2] - bbox[0], bbox[3] - bbox[1]


def wrap_text(text: str, max_chars: int) -> list[str]:
    lines: list[str] = []
    for part in str(text).splitlines() or [""]:
        wrapped = textwrap.wrap(part, width=max_chars, break_long_words=False, replace_whitespace=False)
        lines.extend(wrapped or [""])
    return lines


def draw_multiline(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    fnt: ImageFont.ImageFont,
    fill: str = COLORS["ink"],
    max_chars: int = 32,
    line_gap: int = 7,
    center: bool = False,
) -> int:
    x, y = xy
    lines = wrap_text(text, max_chars)
    line_h = text_size(draw, "Ág", fnt)[1] + line_gap
    for line in lines:
        tx = x
        if center:
            w, _ = text_size(draw, line, fnt)
            tx = x - w // 2
        draw.text((tx, y), line, font=fnt, fill=fill)
        y += line_h
    return y


def center_text(
    draw: ImageDraw.ImageDraw,
    box_xy: tuple[int, int, int, int],
    text: str,
    fnt: ImageFont.ImageFont,
    fill: str = COLORS["ink"],
    max_chars: int = 24,
) -> None:
    x1, y1, x2, y2 = box_xy
    lines = wrap_text(text, max_chars)
    line_h = text_size(draw, "Ág", fnt)[1] + 7
    total_h = line_h * len(lines)
    y = y1 + ((y2 - y1) - total_h) // 2
    for line in lines:
        w, _ = text_size(draw, line, fnt)
        draw.text((x1 + ((x2 - x1) - w) // 2, y), line, font=fnt, fill=fill)
        y += line_h


def rounded_box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    title: str,
    lines: list[str] | None = None,
    fill: str = COLORS["white"],
    outline: str = COLORS["line"],
    radius: int = 18,
    title_fill: str = COLORS["ink"],
) -> None:
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = xy
    draw.text((x1 + 18, y1 + 14), title, font=F_H3, fill=title_fill)
    if lines:
        y = y1 + 56
        for line in lines:
            draw.text((x1 + 18, y), line, font=F_SMALL, fill=COLORS["muted"])
            y += 25
            if y > y2 - 24:
                break


def pill(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    text: str,
    fill: str,
    outline: str,
    max_chars: int = 22,
) -> None:
    draw.rounded_rectangle(xy, radius=60, fill=fill, outline=outline, width=3)
    center_text(draw, xy, text, F_TEXT, max_chars=max_chars)


def arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    fill: str = COLORS["line"],
    width: int = 4,
    dashed: bool = False,
) -> None:
    x1, y1 = start
    x2, y2 = end
    if dashed:
        segments = 18
        for i in range(segments):
            if i % 2 == 0:
                sx = x1 + (x2 - x1) * i / segments
                sy = y1 + (y2 - y1) * i / segments
                ex = x1 + (x2 - x1) * (i + 1) / segments
                ey = y1 + (y2 - y1) * (i + 1) / segments
                draw.line((sx, sy, ex, ey), fill=fill, width=width)
    else:
        draw.line((x1, y1, x2, y2), fill=fill, width=width)
    import math

    angle = math.atan2(y2 - y1, x2 - x1)
    size = 16
    p1 = (x2 - size * math.cos(angle - math.pi / 6), y2 - size * math.sin(angle - math.pi / 6))
    p2 = (x2 - size * math.cos(angle + math.pi / 6), y2 - size * math.sin(angle + math.pi / 6))
    draw.polygon([(x2, y2), p1, p2], fill=fill)


def actor(draw: ImageDraw.ImageDraw, x: int, y: int, label: str) -> None:
    draw.ellipse((x - 20, y - 72, x + 20, y - 32), outline=COLORS["line"], width=4, fill=COLORS["white"])
    draw.line((x, y - 30, x, y + 34), fill=COLORS["line"], width=4)
    draw.line((x - 42, y - 5, x + 42, y - 5), fill=COLORS["line"], width=4)
    draw.line((x, y + 34, x - 38, y + 86), fill=COLORS["line"], width=4)
    draw.line((x, y + 34, x + 38, y + 86), fill=COLORS["line"], width=4)
    draw_multiline(draw, (x - 95, y + 98), label, F_SMALL_BOLD, max_chars=18, center=True)


def canvas(title: str, subtitle: str = "", size: tuple[int, int] = (2200, 1350)) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", size, COLORS["bg"])
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, size[0], 112), fill="#0F172A")
    draw.text((56, 30), title, font=F_TITLE, fill=COLORS["white"])
    if subtitle:
        draw.text((56, 82), subtitle, font=F_SMALL, fill="#CBD5E1")
    return img, draw


def save(img: Image.Image, path: Path) -> Path:
    img.save(path, "PNG", optimize=True)
    return path


def draw_use_case(path: Path) -> None:
    img, draw = canvas(
        "Diagrama de casos de uso general",
        "Actores principales y funciones del sistema web de Calzatura Vilchez",
    )
    boundary = (455, 165, 1745, 1195)
    draw.rounded_rectangle(boundary, radius=34, fill="#EFF6FF", outline=COLORS["blue_line"], width=4)
    draw.text((505, 190), "Sistema web Calzatura Vilchez", font=F_H2, fill=COLORS["blue_line"])

    actor(draw, 190, 360, "Visitante / Cliente")
    actor(draw, 1950, 370, "Administrador")
    actor(draw, 1980, 860, "Trabajador")
    actor(draw, 190, 875, "Stripe")
    actor(draw, 190, 1110, "Servicio IA")

    use_cases = [
        ((585, 285, 910, 385), "Navegar catálogo"),
        ((965, 285, 1290, 385), "Registrarse e iniciar sesión"),
        ((1345, 285, 1670, 385), "Gestionar carrito y checkout"),
        ((610, 485, 960, 590), "Consultar pedidos"),
        ((1040, 490, 1400, 595), "Gestionar productos y stock"),
        ((1480, 490, 1695, 595), "Auditar operaciones"),
        ((610, 700, 985, 810), "Registrar ventas físicas"),
        ((1065, 705, 1425, 815), "Importar y exportar Excel"),
        ((680, 920, 1100, 1035), "Ejecutar predicción IRE"),
        ((1190, 920, 1625, 1035), "Administrar usuarios y fabricantes"),
    ]
    for xy, label in use_cases:
        pill(draw, xy, label, COLORS["white"], COLORS["blue_line"], max_chars=24)

    for target in [(585, 335), (965, 335), (1345, 335), (610, 540)]:
        arrow(draw, (290, 395), target, COLORS["line"], 3)
    arrow(draw, (300, 885), (1345, 340), COLORS["amber_line"], 3, dashed=True)
    arrow(draw, (300, 1105), (680, 980), COLORS["violet_line"], 3, dashed=True)
    for target in [(1040, 540), (1480, 540), (1065, 760), (680, 980), (1190, 980)]:
        arrow(draw, (1860, 400), target, COLORS["line"], 3)
    arrow(draw, (1880, 870), (610, 755), COLORS["line"], 3)

    draw.text((495, 1225), "Alcance excluido: no se modelan procesos fuera del sistema web comercial.", font=F_SMALL, fill=COLORS["muted"])
    save(img, path)


def draw_class(path: Path) -> None:
    img, draw = canvas(
        "Diagrama de clases",
        "Entidades de dominio, servicios principales y relaciones operativas",
        (2300, 1450),
    )
    boxes = {
        "Usuario": (65, 165, 420, 360, COLORS["blue"], COLORS["blue_line"], [
            "+ uid: string",
            "+ email: string",
            "+ rol: cliente|trabajador|admin",
            "+ actualizarPerfil()",
        ]),
        "Producto": (535, 165, 900, 385, COLORS["green"], COLORS["green_line"], [
            "+ id: uuid",
            "+ nombre, categoría, marca",
            "+ precio, stock, tallaStock",
            "+ validarStock()",
        ]),
        "ProductoCódigo": (985, 165, 1325, 340, COLORS["green"], COLORS["green_line"], [
            "+ productoId: uuid",
            "+ código: string",
            "+ únicoPorVariante()",
        ]),
        "ProductoFinanza": (1410, 165, 1790, 360, COLORS["green"], COLORS["green_line"], [
            "+ productoId: uuid",
            "+ costo, margen",
            "+ precioSugerido",
            "+ calcularMargen()",
        ]),
        "Fabricante": (1870, 165, 2225, 340, COLORS["green"], COLORS["green_line"], [
            "+ id: uuid",
            "+ marca: string",
            "+ dni_hash / dni_masked",
        ]),
        "Pedido": (105, 540, 485, 765, COLORS["amber"], COLORS["amber_line"], [
            "+ id: uuid",
            "+ userId: string",
            "+ estado: EstadoPedido",
            "+ métodoPago",
            "+ cambiarEstado()",
        ]),
        "PedidoItem": (575, 550, 930, 745, COLORS["amber"], COLORS["amber_line"], [
            "+ productoId: uuid",
            "+ talla, cantidad",
            "+ precioUnitario",
            "+ subtotal()",
        ]),
        "VentaDiaria": (1035, 545, 1415, 755, COLORS["amber"], COLORS["amber_line"], [
            "+ productoId: uuid",
            "+ cantidad, total",
            "+ canal: tienda",
            "+ registrar()/devolver()",
        ]),
        "Favorito": (1535, 555, 1845, 730, COLORS["blue"], COLORS["blue_line"], [
            "+ userId",
            "+ productoId",
            "+ alternar()",
        ]),
        "Auditoría": (1930, 545, 2245, 755, COLORS["red"], COLORS["red_line"], [
            "+ acción",
            "+ entidad",
            "+ detalle redacted",
            "+ registrarEvento()",
        ]),
        "PredicciónIRE": (360, 930, 785, 1165, COLORS["violet"], COLORS["violet_line"], [
            "+ score: number",
            "+ nivel: bajo|medio|alto",
            "+ data_hash",
            "+ calcularIRE()",
        ]),
        "ModeloEstado": (890, 940, 1275, 1145, COLORS["violet"], COLORS["violet_line"], [
            "+ modeloVersión",
            "+ MAE/MAPE",
            "+ backtesting()",
        ]),
        "BFFService": (1410, 930, 1845, 1165, COLORS["gray"], COLORS["line"], [
            "+ verificarFirebase()",
            "+ validarRol()",
            "+ invocarRPC()",
            "+ redactarPII()",
        ]),
    }
    centers: dict[str, tuple[int, int]] = {}
    for name, (x1, y1, x2, y2, fill, outline, lines) in boxes.items():
        rounded_box(draw, (x1, y1, x2, y2), name, lines, fill, outline)
        centers[name] = ((x1 + x2) // 2, (y1 + y2) // 2)

    relations = [
        ("Usuario", "Pedido", "1..* pedidos"),
        ("Usuario", "Favorito", "1..* favoritos"),
        ("Pedido", "PedidoItem", "1..* ítems"),
        ("PedidoItem", "Producto", "*..1 producto"),
        ("Producto", "ProductoCódigo", "1..1"),
        ("Producto", "ProductoFinanza", "1..1"),
        ("Producto", "Fabricante", "*..1 marca"),
        ("Producto", "VentaDiaria", "1..* ventas"),
        ("BFFService", "Auditoría", "registra"),
        ("BFFService", "Pedido", "controla estados"),
        ("BFFService", "VentaDiaria", "RPC atómica"),
        ("BFFService", "PredicciónIRE", "consulta IA"),
        ("PredicciónIRE", "ModeloEstado", "métricas"),
    ]
    for a, b, label in relations:
        ax, ay = centers[a]
        bx, by = centers[b]
        arrow(draw, (ax, ay), (bx, by), COLORS["line"], 3)
        lx, ly = (ax + bx) // 2, (ay + by) // 2
        draw.rounded_rectangle((lx - 85, ly - 18, lx + 85, ly + 18), radius=10, fill=COLORS["white"], outline="#CBD5E1")
        center_text(draw, (lx - 85, ly - 18, lx + 85, ly + 18), label, F_SMALL, max_chars=18)

    draw.text((65, 1340), "Nota: las clases representan el dominio funcional del sistema web; la implementación usa React, BFF Node/Express, Supabase RPC/RLS e IA.", font=F_SMALL, fill=COLORS["muted"])
    save(img, path)


def draw_er(path: Path) -> None:
    img, draw = canvas(
        "Modelo entidad-relación",
        "Tablas principales de Supabase y cardinalidades del sistema web",
        (2300, 1450),
    )
    tables = {
        "usuarios": (80, 175, 435, 405, ["PK uid", "email", "rol", "telefono", "dni_hash"]),
        "productos": (590, 175, 965, 455, ["PK id", "nombre", "categoría", "marca", "precio", "stock", "tallaStock", "activo"]),
        "productoCodigos": (1120, 175, 1510, 385, ["PK id", "FK productoId", "codigo único"]),
        "productoFinanzas": (1665, 175, 2065, 410, ["PK productoId", "costo", "margen", "precioSugerido"]),
        "pedidos": (95, 635, 505, 890, ["PK id", "FK userId", "items JSON", "total", "estado", "metodoPago"]),
        "favoritos": (610, 655, 950, 835, ["PK userId+productoId", "FK userId", "FK productoId"]),
        "ventasDiarias": (1115, 630, 1515, 890, ["PK id", "FK productoId", "cantidad", "total", "fecha", "canal"]),
        "fabricantes": (1665, 650, 2070, 840, ["PK id", "marca", "dni_hash", "dni_masked"]),
        "auditoria": (350, 1040, 760, 1245, ["PK id", "accion", "entidad", "entidadId", "detalle redactado"]),
        "ireHistorial": (945, 1030, 1355, 1265, ["PK id", "score", "nivel", "data_hash", "modeloVersion"]),
        "modeloEstado": (1540, 1035, 1950, 1255, ["PK id", "modeloVersion", "MAE", "MAPE", "fechaCorte"]),
    }
    centers = {}
    for name, (x1, y1, x2, y2, fields) in tables.items():
        rounded_box(draw, (x1, y1, x2, y2), name, fields, COLORS["white"], COLORS["line"])
        centers[name] = ((x1 + x2) // 2, (y1 + y2) // 2)

    relations = [
        ("usuarios", "pedidos", "1 a N"),
        ("usuarios", "favoritos", "1 a N"),
        ("productos", "favoritos", "1 a N"),
        ("productos", "productoCodigos", "1 a 1..N"),
        ("productos", "productoFinanzas", "1 a 1"),
        ("productos", "ventasDiarias", "1 a N"),
        ("fabricantes", "productos", "1 a N"),
        ("pedidos", "auditoria", "N eventos"),
        ("ventasDiarias", "ireHistorial", "base IA"),
        ("modeloEstado", "ireHistorial", "versión"),
    ]
    for a, b, label in relations:
        ax, ay = centers[a]
        bx, by = centers[b]
        arrow(draw, (ax, ay), (bx, by), COLORS["blue_line"], 3)
        lx, ly = (ax + bx) // 2, (ay + by) // 2
        draw.rounded_rectangle((lx - 62, ly - 16, lx + 62, ly + 16), radius=8, fill="#EFF6FF", outline=COLORS["blue_line"])
        center_text(draw, (lx - 62, ly - 16, lx + 62, ly + 16), label, F_SMALL, max_chars=10)

    draw.text((80, 1335), "Controles: RLS/FORCE RLS, REVOKE a anon/authenticated en tablas sensibles y operaciones mutantes vía BFF/RPC.", font=F_SMALL, fill=COLORS["muted"])
    save(img, path)


def sequence_lanes(draw: ImageDraw.ImageDraw, participants: list[str], y_top: int, y_bottom: int) -> list[int]:
    xs = [160 + i * ((2050 - 160) // (len(participants) - 1)) for i in range(len(participants))]
    for x, name in zip(xs, participants):
        draw.rounded_rectangle((x - 120, y_top, x + 120, y_top + 60), radius=16, fill=COLORS["white"], outline=COLORS["line"], width=3)
        center_text(draw, (x - 118, y_top, x + 118, y_top + 60), name, F_SMALL_BOLD, max_chars=18)
        draw.line((x, y_top + 60, x, y_bottom), fill="#94A3B8", width=2)
    return xs


def message(draw: ImageDraw.ImageDraw, x1: int, x2: int, y: int, text: str, dashed: bool = False) -> None:
    arrow(draw, (x1, y), (x2, y), COLORS["line"], 3, dashed=dashed)
    mid = (x1 + x2) // 2
    draw.rounded_rectangle((mid - 185, y - 40, mid + 185, y - 10), radius=8, fill=COLORS["white"], outline="#CBD5E1")
    center_text(draw, (mid - 180, y - 40, mid + 180, y - 10), text, F_SMALL, max_chars=38)


def draw_sale_sequence(path: Path) -> None:
    img, draw = canvas(
        "Diagrama de secuencia: registro de venta física",
        "Flujo trabajador/admin → BFF → RPC Supabase → auditoría",
        (2250, 1350),
    )
    participants = ["Trabajador", "UI Ventas", "BFF", "RPC Supabase", "Productos", "Ventas", "Auditoría"]
    xs = sequence_lanes(draw, participants, 165, 1210)
    steps = [
        (0, 1, 290, "Selecciona producto, talla y cantidad", False),
        (1, 1, 385, "Valida campos y stock visible", False),
        (1, 2, 480, "POST /staff/dailySales/register", False),
        (2, 2, 575, "Verifica Firebase y rol", False),
        (2, 3, 670, "register_daily_sales_atomic", False),
        (3, 4, 765, "Descuenta stock por talla", False),
        (3, 5, 860, "Inserta ventaDiaria", False),
        (3, 6, 955, "Registra evento sin PII", False),
        (3, 2, 1050, "Resultado atómico", True),
        (2, 1, 1145, "Respuesta y toast de éxito", True),
    ]
    for a, b, y, label, dashed in steps:
        message(draw, xs[a], xs[b], y, label, dashed)
    draw.text((75, 1250), "Garantía funcional: si la RPC falla, no se registra venta parcial ni se descuenta stock fuera de la transacción.", font=F_SMALL, fill=COLORS["muted"])
    save(img, path)


def draw_prediction_sequence(path: Path) -> None:
    img, draw = canvas(
        "Diagrama de secuencia: predicción de riesgo empresarial",
        "Flujo de predicción, backtesting, persistencia de IRE y respuesta al administrador",
        (2250, 1350),
    )
    participants = ["Administrador", "UI Predicción", "BFF", "Servicio IA", "Supabase", "Modelo/IRE", "Dashboard"]
    xs = sequence_lanes(draw, participants, 165, 1210)
    steps = [
        (0, 1, 285, "Solicita horizonte de análisis", False),
        (1, 2, 380, "Token Firebase + parámetros", False),
        (2, 4, 475, "Consulta ventas, productos y métricas", False),
        (2, 3, 570, "GET /api/predict/combined", False),
        (3, 5, 665, "Prepara dataset y calcula features", False),
        (5, 5, 760, "Backtesting MAE/MAPE y data_sufficient", False),
        (5, 5, 855, "Calcula IRE y nivel de riesgo", False),
        (3, 4, 950, "Persiste ireHistorial/modeloEstado", False),
        (3, 2, 1045, "Predicción + métricas", True),
        (2, 6, 1140, "Renderiza tabla, IRE y evidencia", True),
    ]
    for a, b, y, label, dashed in steps:
        message(draw, xs[a], xs[b], y, label, dashed)
    draw.text((75, 1250), "Control medible: el resultado incluye suficiencia de datos, versión del modelo, hash del dataset y métricas de error.", font=F_SMALL, fill=COLORS["muted"])
    save(img, path)


def draw_activity(path: Path) -> None:
    img, draw = canvas(
        "Diagrama de actividades: checkout web",
        "Proceso de compra con validación de stock, pago y manejo de errores",
        (2200, 1450),
    )
    nodes = {
        "start": (1040, 165, 1160, 225, "Inicio", COLORS["gray"], COLORS["line"]),
        "catalog": (865, 300, 1335, 380, "Explorar catálogo y ficha de producto", COLORS["blue"], COLORS["blue_line"]),
        "size": (865, 455, 1335, 535, "Seleccionar talla y agregar al carrito", COLORS["blue"], COLORS["blue_line"]),
        "address": (835, 610, 1365, 690, "Confirmar datos de entrega y ubicación", COLORS["blue"], COLORS["blue_line"]),
        "decision": (960, 770, 1240, 885, "¿Método de pago?", COLORS["amber"], COLORS["amber_line"]),
        "stripe": (360, 970, 820, 1065, "Stripe: crear sesión y esperar webhook firmado", COLORS["violet"], COLORS["violet_line"]),
        "cod": (1380, 970, 1840, 1065, "Contraentrega: crear pedido pendiente", COLORS["green"], COLORS["green_line"]),
        "validate": (835, 1145, 1365, 1235, "BFF valida precio, stock, envío e idempotencia", COLORS["green"], COLORS["green_line"]),
        "success": (845, 1315, 1355, 1390, "Pedido visible en confirmación e historial", COLORS["blue"], COLORS["blue_line"]),
        "error": (70, 1145, 620, 1235, "Error controlado: mensaje y reintento", COLORS["red"], COLORS["red_line"]),
    }
    for key, (x1, y1, x2, y2, text, fill, outline) in nodes.items():
        if key == "decision":
            draw.polygon([(1100, y1), (x2, (y1 + y2) // 2), (1100, y2), (x1, (y1 + y2) // 2)], fill=fill, outline=outline)
            center_text(draw, (x1 + 15, y1 + 12, x2 - 15, y2 - 12), text, F_H3, max_chars=18)
        elif key == "start":
            draw.ellipse((x1, y1, x2, y2), fill=COLORS["line"], outline=COLORS["line"], width=3)
            center_text(draw, (x1, y1, x2, y2), text, F_SMALL_BOLD, fill=COLORS["white"], max_chars=10)
        else:
            rounded_box(draw, (x1, y1, x2, y2), "", None, fill, outline, radius=18)
            center_text(draw, (x1 + 10, y1, x2 - 10, y2), text, F_TEXT, max_chars=34)

    def c(key: str) -> tuple[int, int]:
        x1, y1, x2, y2, *_ = nodes[key]
        return (x1 + x2) // 2, (y1 + y2) // 2

    for a, b in [("start", "catalog"), ("catalog", "size"), ("size", "address"), ("address", "decision")]:
        ax, ay = c(a)
        bx, by = c(b)
        arrow(draw, (ax, ay + 40), (bx, by - 45), COLORS["line"], 4)
    arrow(draw, (960, 828), (820, 1015), COLORS["line"], 4)
    arrow(draw, (1240, 828), (1380, 1015), COLORS["line"], 4)
    draw.text((820, 920), "Tarjeta", font=F_SMALL_BOLD, fill=COLORS["muted"])
    draw.text((1260, 920), "Contraentrega", font=F_SMALL_BOLD, fill=COLORS["muted"])
    arrow(draw, c("stripe"), (1090, 1145), COLORS["line"], 4)
    arrow(draw, c("cod"), (1110, 1145), COLORS["line"], 4)
    arrow(draw, (1100, 1235), (1100, 1315), COLORS["line"], 4)
    arrow(draw, (835, 1190), (620, 1190), COLORS["red_line"], 4, dashed=True)
    draw.text((75, 1375), "El flujo evita confirmar pedidos cuando la validación de producto, pago o stock no se completa correctamente.", font=F_SMALL, fill=COLORS["muted"])
    save(img, path)


def draw_deployment(path: Path) -> None:
    img, draw = canvas(
        "Diagrama de despliegue y arquitectura",
        "Distribución de componentes, servicios externos y controles de entrega",
        (2300, 1450),
    )
    components = {
        "Usuario navegador": (80, 210, 440, 405, COLORS["blue"], COLORS["blue_line"], ["React SPA", "Firebase Hosting", "App Check"]),
        "BFF Render": (640, 190, 1040, 430, COLORS["green"], COLORS["green_line"], ["Node/Express", "service_role", "validación de rol", "redacción PII"]),
        "Supabase": (1260, 190, 1695, 455, COLORS["amber"], COLORS["amber_line"], ["PostgreSQL", "RLS/FORCE RLS", "RPC atómicas", "Storage"]),
        "Firebase Auth": (1880, 205, 2220, 390, COLORS["violet"], COLORS["violet_line"], ["Identidad", "tokens", "App Check"]),
        "Stripe": (215, 680, 555, 855, COLORS["violet"], COLORS["violet_line"], ["Checkout", "webhook firmado"]),
        "Cloudinary": (720, 670, 1060, 845, COLORS["blue"], COLORS["blue_line"], ["firma BFF", "imágenes"]),
        "Proveedor DNI": (1230, 660, 1580, 845, COLORS["gray"], COLORS["line"], ["lookup protegido", "proof token"]),
        "Servicio IA": (1780, 650, 2200, 880, COLORS["green"], COLORS["green_line"], ["predicción demanda", "IRE", "backtesting"]),
        "GitHub Actions": (605, 1050, 1065, 1265, COLORS["gray"], COLORS["line"], ["CI", "tests", "Docker", "ZAP", "stress"]),
        "Producción": (1280, 1050, 1730, 1265, COLORS["blue"], COLORS["blue_line"], ["deploy controlado", "secrets requeridos", "readiness"]),
    }
    centers = {}
    for name, (x1, y1, x2, y2, fill, outline, lines) in components.items():
        rounded_box(draw, (x1, y1, x2, y2), name, lines, fill, outline)
        centers[name] = ((x1 + x2) // 2, (y1 + y2) // 2)

    arrows = [
        ("Usuario navegador", "BFF Render", "HTTPS/JSON"),
        ("BFF Render", "Supabase", "SQL/RPC"),
        ("BFF Render", "Firebase Auth", "verifica token"),
        ("BFF Render", "Stripe", "checkout/webhook"),
        ("BFF Render", "Cloudinary", "firma carga"),
        ("BFF Render", "Proveedor DNI", "consulta validada"),
        ("BFF Render", "Servicio IA", "API interna"),
        ("GitHub Actions", "Producción", "quality gate"),
        ("Producción", "Usuario navegador", "Hosting"),
        ("Producción", "BFF Render", "Render deploy"),
        ("Producción", "Servicio IA", "AI deploy"),
    ]
    for a, b, label in arrows:
        ax, ay = centers[a]
        bx, by = centers[b]
        arrow(draw, (ax, ay), (bx, by), COLORS["line"], 4)
        lx, ly = (ax + bx) // 2, (ay + by) // 2
        draw.rounded_rectangle((lx - 90, ly - 18, lx + 90, ly + 18), radius=8, fill=COLORS["white"], outline="#CBD5E1")
        center_text(draw, (lx - 86, ly - 18, lx + 86, ly + 18), label, F_SMALL, max_chars=18)

    draw.text((80, 1340), "Separación clave: el frontend no contiene service_role; las mutaciones sensibles pasan por BFF y RPC controladas.", font=F_SMALL, fill=COLORS["muted"])
    save(img, path)


def write_mermaid_files() -> dict[str, Path]:
    diagrams = {
        "01_casos_uso_general.mmd": """flowchart LR
  Visitante[Visitante / Cliente]
  Admin[Administrador]
  Trabajador[Trabajador]
  Stripe[Stripe]
  IA[Servicio IA]
  subgraph Sistema["Sistema web Calzatura Vilchez"]
    Catalogo((Navegar catálogo))
    Auth((Registrarse e iniciar sesión))
    Checkout((Gestionar carrito y checkout))
    Pedidos((Consultar pedidos))
    Productos((Gestionar productos y stock))
    Ventas((Registrar ventas físicas))
    Excel((Importar y exportar Excel))
    Prediccion((Ejecutar predicción IRE))
    Auditoria((Auditar operaciones))
  end
  Visitante --> Catalogo
  Visitante --> Auth
  Visitante --> Checkout
  Visitante --> Pedidos
  Admin --> Productos
  Admin --> Ventas
  Admin --> Excel
  Admin --> Prediccion
  Admin --> Auditoria
  Trabajador --> Ventas
  Stripe -.webhook.-> Checkout
  IA -.modelo.-> Prediccion
""",
        "02_diagrama_clases.mmd": """classDiagram
  class Usuario {
    +string uid
    +string email
    +string rol
    +actualizarPerfil()
  }
  class Producto {
    +uuid id
    +string nombre
    +decimal precio
    +json tallaStock
    +validarStock()
  }
  class Pedido {
    +uuid id
    +string estado
    +string metodoPago
    +cambiarEstado()
  }
  class PedidoItem {
    +uuid productoId
    +string talla
    +int cantidad
    +subtotal()
  }
  class VentaDiaria {
    +uuid productoId
    +int cantidad
    +registrar()
    +devolver()
  }
  class Auditoria {
    +string accion
    +string entidad
    +json detalle
    +registrarEvento()
  }
  class PrediccionIRE {
    +number score
    +string nivel
    +calcularIRE()
  }
  class BFFService {
    +verificarFirebase()
    +validarRol()
    +invocarRPC()
    +redactarPII()
  }
  Usuario "1" --> "*" Pedido
  Pedido "1" --> "*" PedidoItem
  PedidoItem "*" --> "1" Producto
  Producto "1" --> "*" VentaDiaria
  BFFService --> Pedido
  BFFService --> VentaDiaria
  BFFService --> Auditoria
  BFFService --> PrediccionIRE
""",
        "03_modelo_entidad_relacion.mmd": """erDiagram
  usuarios ||--o{ pedidos : realiza
  usuarios ||--o{ favoritos : guarda
  productos ||--o{ favoritos : aparece
  productos ||--o{ ventasDiarias : registra
  productos ||--|| productoFinanzas : tiene
  productos ||--o{ productoCodigos : identifica
  fabricantes ||--o{ productos : provee
  pedidos ||--o{ auditoria : genera
  ventasDiarias ||--o{ ireHistorial : alimenta
  modeloEstado ||--o{ ireHistorial : versiona
""",
        "04_secuencia_registro_venta.mmd": """sequenceDiagram
  actor Trabajador
  participant UI as UI Ventas
  participant BFF
  participant RPC as RPC Supabase
  participant Productos
  participant Ventas
  participant Auditoria
  Trabajador->>UI: Selecciona producto, talla y cantidad
  UI->>UI: Valida campos y stock visible
  UI->>BFF: POST /staff/dailySales/register
  BFF->>BFF: Verifica Firebase y rol
  BFF->>RPC: register_daily_sales_atomic
  RPC->>Productos: Descuenta stock por talla
  RPC->>Ventas: Inserta ventaDiaria
  RPC->>Auditoria: Registra evento sin PII
  RPC-->>BFF: Resultado atómico
  BFF-->>UI: Respuesta y toast de éxito
""",
        "05_secuencia_prediccion_riesgo.mmd": """sequenceDiagram
  actor Admin as Administrador
  participant UI as UI Predicción
  participant BFF
  participant IA as Servicio IA
  participant DB as Supabase
  participant Modelo as Modelo/IRE
  participant Dash as Dashboard
  Admin->>UI: Solicita horizonte de análisis
  UI->>BFF: Token Firebase + parámetros
  BFF->>DB: Consulta ventas, productos y métricas
  BFF->>IA: GET /api/predict/combined
  IA->>Modelo: Prepara dataset y calcula features
  Modelo->>Modelo: Backtesting MAE/MAPE y data_sufficient
  Modelo->>Modelo: Calcula IRE y nivel de riesgo
  IA->>DB: Persiste ireHistorial/modeloEstado
  IA-->>BFF: Predicción + métricas
  BFF-->>Dash: Renderiza tabla, IRE y evidencia
""",
        "06_actividad_checkout.mmd": """flowchart TD
  A([Inicio]) --> B[Explorar catálogo y ficha de producto]
  B --> C[Seleccionar talla y agregar al carrito]
  C --> D[Confirmar datos de entrega y ubicación]
  D --> E{Método de pago}
  E -->|Tarjeta| F[Crear sesión Stripe y esperar webhook firmado]
  E -->|Contraentrega| G[Crear pedido pendiente]
  F --> H[BFF valida precio, stock, envío e idempotencia]
  G --> H
  H --> I[Pedido visible en confirmación e historial]
  H -->|error| J[Mensaje controlado y reintento]
  J --> D
""",
        "07_despliegue_arquitectura.mmd": """flowchart LR
  Browser[Navegador: React SPA / Firebase Hosting]
  BFF[BFF Render: Node + Express + service_role]
  DB[Supabase: PostgreSQL + RLS + RPC]
  Auth[Firebase Auth + App Check]
  Stripe[Stripe Checkout + webhook]
  Cloudinary[Cloudinary firma BFF]
  DNI[Proveedor DNI protegido]
  AI[Servicio IA]
  CI[GitHub Actions: CI + Docker + ZAP + stress]
  Browser -->|HTTPS| BFF
  BFF -->|SQL/RPC| DB
  BFF -->|verifica token| Auth
  BFF --> Stripe
  BFF --> Cloudinary
  BFF --> DNI
  BFF --> AI
  CI --> Browser
  CI --> BFF
  CI --> AI
""",
    }
    written: dict[str, Path] = {}
    for name, content in diagrams.items():
        target = MMD_DIR / name
        target.write_text(content, encoding="utf-8")
        written[name] = target
    return written


def generate_images() -> dict[str, Path]:
    images = {
        "01_casos_uso_general.png": draw_use_case,
        "02_diagrama_clases.png": draw_class,
        "03_modelo_entidad_relacion.png": draw_er,
        "04_secuencia_registro_venta.png": draw_sale_sequence,
        "05_secuencia_prediccion_riesgo.png": draw_prediction_sequence,
        "06_actividad_checkout.png": draw_activity,
        "07_despliegue_arquitectura.png": draw_deployment,
    }
    written: dict[str, Path] = {}
    for name, fn in images.items():
        target = PNG_DIR / name
        fn(target)
        written[name] = target
    return written


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)
    for name, size, bold_color in [
        ("Title", 18, RGBColor(15, 23, 42)),
        ("Heading 1", 13, RGBColor(30, 64, 175)),
        ("Heading 2", 11, RGBColor(15, 23, 42)),
    ]:
        styles[name].font.name = "Arial"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = bold_color


def add_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(8.5)


def generate_docx(diagrams: list[Diagram]) -> Path:
    doc = Document()
    style_document(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DOCUMENTO 4 — DIAGRAMAS UML Y MODELADO DEL SISTEMA WEB")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(15, 23, 42)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Proyecto: Sistema web de e-commerce, ventas físicas e inteligencia artificial para Calzatura Vilchez").italic = True

    doc.add_paragraph(
        "Este documento resume el diseño del sistema antes de la implementación y lo vincula con artefactos "
        "reales del repositorio: rutas React, BFF Node/Express, Supabase con RLS/RPC, servicio de IA, seguridad, "
        "auditoría, pruebas automatizadas, Docker, ZAP y pruebas de estrés. Se excluyen deliberadamente módulos "
        "fuera del alcance comercial del proyecto."
    )

    doc.add_heading("Resumen de diagramas", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["N.º", "Diagrama", "Para qué sirve", "Evidencia técnica", "Archivo editable"]
    for index, header in enumerate(headers):
        add_cell_text(table.rows[0].cells[index], header, True)
    for idx, item in enumerate(diagrams, start=1):
        row = table.add_row().cells
        add_cell_text(row[0], f"{idx:02d}")
        add_cell_text(row[1], item.title)
        add_cell_text(row[2], item.purpose)
        add_cell_text(row[3], item.evidence)
        add_cell_text(row[4], item.mermaid.name)

    doc.add_heading("Diagramas desarrollados", level=1)
    for idx, item in enumerate(diagrams, start=1):
        doc.add_heading(f"{idx:02d}. {item.title}", level=2)
        meta = doc.add_table(rows=3, cols=2)
        meta.style = "Table Grid"
        add_cell_text(meta.rows[0].cells[0], "Objetivo", True)
        add_cell_text(meta.rows[0].cells[1], item.purpose)
        add_cell_text(meta.rows[1].cells[0], "Elementos incluidos", True)
        add_cell_text(meta.rows[1].cells[1], item.elements)
        add_cell_text(meta.rows[2].cells[0], "Código / evidencia", True)
        add_cell_text(meta.rows[2].cells[1], item.evidence)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(item.image), width=Inches(10.6))
        cap = doc.add_paragraph(f"Figura {idx}. {item.title}. Fuente: elaboración propia con base en el repositorio del sistema web.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True

    doc.add_heading("Criterio de autenticidad", level=1)
    doc.add_paragraph(
        "Los diagramas no son genéricos: cada elemento se trazó contra módulos existentes del repositorio. "
        "Las rutas principales provienen de calzatura-vilchez/src/App.tsx; las operaciones críticas de pedidos, "
        "ventas, Excel, Cloudinary, DNI y auditoría pasan por calzatura-vilchez/bff/server.cjs; las transiciones "
        "de pedido se sostienen en calzatura-vilchez/functions/orderStatusPolicy.js; y la capa de datos se cruza "
        "con calzatura-vilchez/supabase/RLS-MATRIX.md y migraciones Supabase."
    )
    doc.add_paragraph(
        "El paquete también incluye fuentes Mermaid editables y PNG renderizados para anexos, presentación o revisión técnica."
    )

    out = DOC_DIR / DOCX_NAME
    doc.save(out)
    shutil.copy2(out, DOWNLOADS / DOCX_NAME)
    return out


def build_zip(docx_path: Path) -> Path:
    target = OUT_DIR / ZIP_NAME
    with zipfile.ZipFile(target, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.write(docx_path, arcname=DOCX_NAME)
        for path in sorted(PNG_DIR.glob("*.png")):
            zf.write(path, arcname=f"png/{path.name}")
        for path in sorted(MMD_DIR.glob("*.mmd")):
            zf.write(path, arcname=f"mermaid/{path.name}")
    return target


def validate_outputs(docx_path: Path, zip_path: Path) -> None:
    errors: list[str] = []
    if not docx_path.exists() or docx_path.stat().st_size < 100_000:
        errors.append(f"DOCX insuficiente o inexistente: {docx_path}")
    if not (DOWNLOADS / DOCX_NAME).exists():
        errors.append("No se copió el DOCX a Descargas.")
    pngs = sorted(PNG_DIR.glob("*.png"))
    mmds = sorted(MMD_DIR.glob("*.mmd"))
    if len(pngs) != 7:
        errors.append(f"Se esperaban 7 PNG y hay {len(pngs)}.")
    if len(mmds) != 7:
        errors.append(f"Se esperaban 7 Mermaid y hay {len(mmds)}.")
    for path in pngs:
        if path.stat().st_size < 20_000:
            errors.append(f"PNG demasiado pequeño: {path.name}")
    if not zip_path.exists() or zip_path.stat().st_size < 100_000:
        errors.append("ZIP insuficiente o inexistente.")
    forbidden = ["rrhh", "psicolog", "calzatura-vilchez-mobile", "admin_dashboard_page", "build.gradle", "móvil", "movil", "mobile"]
    text_blobs = [path.read_text(encoding="utf-8").lower() for path in mmds]
    if any(term in blob for term in forbidden for blob in text_blobs):
        errors.append("Se detectó término fuera de alcance en fuentes Mermaid.")
    if errors:
        raise SystemExit("\n".join(errors))


def main() -> None:
    ensure_dirs()
    mmd = write_mermaid_files()
    images = generate_images()
    diagrams = [
        Diagram(
            key="casos_uso",
            title="Diagrama de casos de uso general",
            purpose="Muestra actores, límites del sistema y funciones visibles por rol.",
            elements="Visitante/cliente, administrador, trabajador, Stripe, servicio IA; catálogo, autenticación, checkout, pedidos, productos, ventas, Excel, predicción y auditoría.",
            evidence="App.tsx; servicios de pedidos, productos, ventas y predicción; BFF server.cjs.",
            image=images["01_casos_uso_general.png"],
            mermaid=mmd["01_casos_uso_general.mmd"],
        ),
        Diagram(
            key="clases",
            title="Diagrama de clases",
            purpose="Resume entidades de dominio y servicios que coordinan reglas críticas.",
            elements="Usuario, Producto, Pedido, PedidoItem, VentaDiaria, Auditoría, PredicciónIRE, ModeloEstado y BFFService.",
            evidence="src/domains/*; bff/server.cjs; functions/orderStatusPolicy.js.",
            image=images["02_diagrama_clases.png"],
            mermaid=mmd["02_diagrama_clases.mmd"],
        ),
        Diagram(
            key="er",
            title="Modelo entidad-relación",
            purpose="Representa la base de datos y sus relaciones principales.",
            elements="usuarios, productos, productoCodigos, productoFinanzas, pedidos, favoritos, ventasDiarias, fabricantes, auditoria, ireHistorial y modeloEstado.",
            evidence="supabase/RLS-MATRIX.md y migraciones Supabase.",
            image=images["03_modelo_entidad_relacion.png"],
            mermaid=mmd["03_modelo_entidad_relacion.mmd"],
        ),
        Diagram(
            key="venta",
            title="Diagrama de secuencia del registro de venta física",
            purpose="Detalla el paso a paso de una venta presencial registrada desde el panel.",
            elements="Trabajador, UI Ventas, BFF, RPC Supabase, productos, ventasDiarias y auditoría.",
            evidence="BFF /staff/dailySales/register, /admin/dailySales/register y RPC register_daily_sales_atomic.",
            image=images["04_secuencia_registro_venta.png"],
            mermaid=mmd["04_secuencia_registro_venta.mmd"],
        ),
        Diagram(
            key="prediccion",
            title="Diagrama de secuencia de predicción de riesgo empresarial",
            purpose="Explica cómo se consulta, calcula, versiona y muestra el IRE.",
            elements="Administrador, UI Predicción, BFF, servicio IA, Supabase, modelo/IRE y dashboard.",
            evidence="AdminPredictions, proxy IA, ai-service, ireHistorial y modeloEstado.",
            image=images["05_secuencia_prediccion_riesgo.png"],
            mermaid=mmd["05_secuencia_prediccion_riesgo.mmd"],
        ),
        Diagram(
            key="actividad",
            title="Diagrama de actividades del checkout web",
            purpose="Muestra decisiones, validaciones, pagos y reintentos del proceso de compra.",
            elements="Catálogo, carrito, dirección, método de pago, Stripe, contraentrega, BFF y error controlado.",
            evidence="CheckoutPage, CartContext, createOrder, Stripe webhook y OrderSuccessPage.",
            image=images["06_actividad_checkout.png"],
            mermaid=mmd["06_actividad_checkout.mmd"],
        ),
        Diagram(
            key="despliegue",
            title="Diagrama de despliegue y arquitectura del sistema",
            purpose="Ubica dónde se ejecutan frontend, BFF, base de datos, IA e integraciones externas.",
            elements="Firebase Hosting, Render BFF, Supabase, Firebase Auth/App Check, Stripe, Cloudinary, proveedor DNI, servicio IA y GitHub Actions.",
            evidence="firebase.json, docker-compose.yml, deploy-production.yml, server.cjs y configuración IA.",
            image=images["07_despliegue_arquitectura.png"],
            mermaid=mmd["07_despliegue_arquitectura.mmd"],
        ),
    ]
    docx = generate_docx(diagrams)
    zf = build_zip(docx)
    validate_outputs(docx, zf)
    print(f"DOCX={docx}")
    print(f"DOWNLOADS={DOWNLOADS / DOCX_NAME}")
    print(f"ZIP={zf}")
    print(f"PNG={len(list(PNG_DIR.glob('*.png')))}")
    print(f"MERMAID={len(list(MMD_DIR.glob('*.mmd')))}")


if __name__ == "__main__":
    main()
