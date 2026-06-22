#!/usr/bin/env python3
"""
Matriz de Operacionalización de Variables — Calzatura Vilchez.
Alineada con CU-T09 (consistencia), instrumento Likert (24 ítems), doc 07 (IRE) y repo.

Referencias: URP/UNCP · Next Tesista · JP Consultoría Tesis · ISO/IEC 25000 (calidad producto).

Uso: python scripts/generar_matriz_operacionalizacion_variables.py
"""
from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "Matriz_Operacionalizacion_Variables_Calzatura_Vilchez_MEJORADA.docx"
OUT_DOCX_V2 = ROOT / "Matriz_Operacionalizacion_Variables_v2_VALIDADA.docx"
OUT_DOCX_V3 = ROOT / "Matriz_Operacionalizacion_Variables_v3_AUDITADA.docx"
OUT_CSV = ROOT / "documentacion/cuadros-excel/CU-T10-matriz-operacionalizacion-variables.csv"

TITULO = (
    "Sistema web de comercio electrónico con modelo de Inteligencia Artificial "
    "para la predicción del riesgo empresarial en la empresa Calzatura Vilchez, Huancayo, 2026"
)

HEADERS = [
    "Variable",
    "Definición conceptual",
    "Definición operacional",
    "Dimensión",
    "Indicadores",
    "Escala / unidad",
    "Técnica",
    "Instrumento / fuente",
    "Sustento científico",
]

VI_CONCEPT = (
    "Intervención tecnológica compuesta por un sistema web de comercio electrónico y un "
    "microservicio de Inteligencia Artificial que digitaliza ventas, inventario, checkout "
    "y predicción de demanda/riesgo (ISO/IEC 12207 — proceso de implementación)."
)

VD_CONCEPT = (
    "Capacidad de cuantificar y anticipar el riesgo empresarial comercial-operativo de "
    "Calzatura Vilchez mediante el Índice de Riesgo Empresarial (IRE), definido como proxy "
    "de exposición a quiebre de stock, presión de ingresos y cambios de demanda "
    "(documentacion/07-modulo-ia-riesgo-empresarial.md)."
)

# 8 filas = 4 dimensiones VI + 4 dimensiones VD (6 indicadores c/u)
ROWS: list[dict[str, str]] = [
    {
        "var": "VI",
        "concept": VI_CONCEPT,
        "oper": (
            "Medición de madurez y digitalización del canal e-commerce: procesos web, conversión "
            "y tiempos operativos registrados en Supabase y percepción del personal (O₁/O₂)."
        ),
        "dim": "D1: E-commerce y transformación digital del negocio",
        "indicators": (
            "1. Nivel de madurez digital de la empresa.\n"
            "2. Tasa de digitalización de procesos de venta e inventario.\n"
            "3. Número de canales de venta digitales activos (web + panel admin).\n"
            "4. Tasa de conversión digital (visitas → pedidos).\n"
            "5. Tiempo de ciclo de procesamiento de pedidos.\n"
            "6. Disponibilidad del sistema web (%) — trazado CU-T09 ES1-6 / VI-D3-2."
        ),
        "scale": (
            "Ordinal 1–5; razón (%); conteo; razón (%); razón (min); "
            "razón (%)."
        ),
        "technique": (
            "Análisis documental, observación sistemática, entrevista semiestructurada "
            "a empleadores/directivos, encuesta Likert (percepción digitalización)."
        ),
        "instrument": (
            "Ficha documental; logs Supabase (pedidos); guía entrevista empleadores; "
            "Instrumento Likert (ítems 19–21 complemento percepción); CI uptime/monitoring."
        ),
        "refs": "Eje 1: Verhoef 2021; Soto-Acosta 2020; Nambisan 2017; Li 2020; Chatterjee 2020.",
    },
    {
        "var": "VI",
        "concept": VI_CONCEPT,
        "oper": (
            "Medición del desempeño del modelo RandomForestRegressor (ai-service) y fallback "
            "promedio móvil mediante validación walk-forward sobre ventasDiarias/pedidos."
        ),
        "dim": "D2: Inteligencia Artificial y analítica de datos empresariales",
        "indicators": (
            "1. MAE del pronóstico de demanda (unidades).\n"
            "2. RMSE del pronóstico de demanda (unidades).\n"
            "3. sMAPE del pronóstico de ventas (%).\n"
            "4. R² del modelo predictivo (0–1).\n"
            "5. Horizonte de predicción configurado (7, 15 o 30 días).\n"
            "6. Tiempo de inferencia del microservicio IA (ms, p95)."
        ),
        "scale": "Razón (unid.); razón (unid.); razón (%); intervalo 0–1; discreto {7,15,30}; razón (ms).",
        "technique": "Análisis de datos históricos, evaluación ML, validación walk-forward.",
        "instrument": (
            "Supabase ventasDiarias/pedidos; ai-service/evaluate.py; "
            "ai-service/models/demand/features_ml.py; reporte scikit-learn."
        ),
        "refs": "Eje 2: Yuan 2021; Davenport 2020; Makridakis 2018; Fildes 2022; Breiman 2001.",
    },
    {
        "var": "VI",
        "concept": VI_CONCEPT,
        "oper": (
            "Medición de calidad técnica del despliegue: latencia, disponibilidad, integraciones "
            "y cobertura de pruebas automatizadas (CI GitHub Actions, Vitest, Playwright)."
        ),
        "dim": "D3: Arquitectura de software y despliegue del sistema",
        "indicators": (
            "1. Latencia p95 API predicción / catálogo BFF (ms).\n"
            "2. Disponibilidad sistema web y microservicio IA (%).\n"
            "3. Tiempo de despliegue nueva versión (min).\n"
            "4. Número de integraciones API activas (Firebase, Supabase, Stripe, IA, ORS).\n"
            "5. Cobertura pruebas automatizadas Vitest (% líneas scope).\n"
            "6. Tiempo medio recuperación ante incidente documentado (h)."
        ),
        "scale": "Razón (ms); razón (%); razón (min); conteo; razón (%); razón (h).",
        "technique": "Pruebas técnicas, monitoreo CI, revisión repositorio, k6 smoke.",
        "instrument": (
            "ci.yml; vitest.config.ts; playwright test; docs/ops/k6-smoke-evidence.json; "
            "09-implementacion-despliegue-ci.md."
        ),
        "refs": "Eje 4: Jamshidi 2018; Khalid 2022; Paleyes 2022; Martinez-Fernandez 2022.",
    },
    {
        "var": "VI",
        "concept": VI_CONCEPT,
        "oper": (
            "Medición del proceso de desarrollo: sprints, trazabilidad CRISP-ML(Q), MLOps "
            "y retorno estimado de la inversión tecnológica."
        ),
        "dim": "D4: Metodología de desarrollo ágil e híbrida del sistema",
        "indicators": (
            "1. Velocidad del equipo por sprint (puntos completados).\n"
            "2. Tasa de completitud de sprints (%).\n"
            "3. Fases CRISP-ML(Q) cubiertas (conteo 1–6).\n"
            "4. Madurez MLOps (ordinal 1–5).\n"
            "5. Tiempo entrega incrementos funcionales (días).\n"
            "6. ROI estimado inversión tecnológica (%)."
        ),
        "scale": "Razón; razón (%); conteo; ordinal 1–5; razón (días); razón (%).",
        "technique": "Revisión documental, control avance, análisis económico.",
        "instrument": "Backlog GitHub; CHANGELOG.md; documentacion/03-planificacion-proyecto-completa.md.",
        "refs": "Eje 5: Dingsøyr 2012; Salgonde 2021; Kuhrmann 2019; Melville 2004; Choi 2018.",
    },
    {
        "var": "VD",
        "concept": VD_CONCEPT,
        "oper": (
            "Caracterización del modelo de riesgo implementado: IRE compuesto (reglas + datos Supabase) "
            "y RandomForestRegressor para demanda; validación documental y experimental."
        ),
        "dim": "D1: Modelos de predicción del riesgo empresarial",
        "indicators": (
            "1. Tipo de modelo IRE (proxy compuesto 40/35/25 + RF demanda).\n"
            "2. Número de variables operativas (stock, ingresos, demanda, productos).\n"
            "3. Horizonte de predicción (7, 15 o 30 días).\n"
            "4. Técnica agregación IRE (ponderación fija + reglas stock/ingresos).\n"
            "5. Método validación (consistencia interna + walk-forward demanda).\n"
            "6. Comparación vs método empírico baseline (% mejora MAE)."
        ),
        "scale": "Nominal; conteo; discreto {7,15,30}; nominal; nominal; razón (%).",
        "technique": "Análisis documental, evaluación comparativa, entrevista directivos.",
        "instrument": (
            "07-modulo-ia-riesgo-empresarial.md; ai-service; Likert ítems 1–6; guía entrevista."
        ),
        "refs": "Eje 3: Ozbayoglu 2020; Barboza 2017; Altman 1968; Beaver 1966; Fildes 2022.",
    },
    {
        "var": "VD",
        "concept": VD_CONCEPT,
        "oper": (
            "Medición de exactitud y robustez: métricas ML de demanda, invariantes IRE, drift; "
            "AUC-ROC solo si existen eventos de crisis etiquetados (doc 07)."
        ),
        "dim": "D2: Exactitud y validación del IRE y modelos predictivos",
        "indicators": (
            "1. Tests consistencia interna IRE (invariantes, monotonía horizonte).\n"
            "2. MAE pronóstico demanda RF vs baseline (unidades).\n"
            "3. RMSE pronóstico demanda (unidades).\n"
            "4. Estabilidad suma contribuciones IRE (= score mostrado).\n"
            "5. Variación desempeño por data drift mensual (%).\n"
            "6. AUC-ROC clasificador — condicional a eventos etiquetados."
        ),
        "scale": "Nominal (pasa/falla); razón; razón; nominal; razón (%); intervalo 0–1 (si aplica).",
        "technique": "Validación modelo, pytest ai-service, análisis estadístico.",
        "instrument": (
            "ai-service/tests/test_risk.py; evaluate.py; Likert ítems 7–12; "
            "ireHistorial (Supabase)."
        ),
        "refs": "Barboza 2017; Kim 2020; Khalid 2022; Paleyes 2022; Du Jardin 2021.",
    },
    {
        "var": "VD",
        "concept": VD_CONCEPT,
        "oper": (
            "Medición directa del IRE y subíndices en escala 0–100 con umbrales de alerta "
            "operativa definidos en panel AdminPredictions."
        ),
        "dim": "D3: Componentes del Índice de Riesgo Empresarial (IRE)",
        "indicators": (
            "1. Valor compuesto IRE (0–100).\n"
            "2. Componente riesgo_stock (subíndice).\n"
            "3. Componente riesgo_ingresos (subíndice).\n"
            "4. Componente riesgo_demanda (subíndice).\n"
            "5. Alerta temprana: IRE ≥ 51 (alto) o ≥ 76 (crítico).\n"
            "6. Frecuencia actualización automática IRE (por consulta API / día)."
        ),
        "scale": "Intervalo 0–100 (todos los componentes); nominal (umbral); razón (veces/día).",
        "technique": "Cálculo automatizado, análisis documental, juicio expertos.",
        "instrument": (
            "AdminPredictions dashboard; ireHistorial; Likert ítems 13–18; ficha cálculo IRE."
        ),
        "refs": "Altman 1968; Beaver 1966; Barboza 2017; Choi 2018; documentacion/07 §2.",
    },
    {
        "var": "VD",
        "concept": VD_CONCEPT,
        "oper": (
            "Medición de impacto operativo pre/post (O₁/O₂): percepción del personal, quiebres de "
            "stock, cumplimiento pedidos y adopción del dashboard."
        ),
        "dim": "D4: Impacto del sistema en la gestión operativa de la PYME",
        "indicators": (
            "1. Reducción riesgo operativo percibido (Likert).\n"
            "2. Mejora tasa cumplimiento de pedidos (%).\n"
            "3. Reducción quiebres de stock (conteo O₁ vs O₂).\n"
            "4. Ahorro estimado gestión inventario (S/.).\n"
            "5. Tasa adopción dashboard IRE (% sesiones admin).\n"
            "6. Percepción mejora gestión riesgo — ítem 24 Likert."
        ),
        "scale": "Ordinal Likert 1–5; razón (%); conteo; razón (S/.); razón (%); ordinal 1–5.",
        "technique": "Comparación pre/post, encuesta, entrevista, registros Supabase.",
        "instrument": (
            "Instrumento Likert ítems 19–24; guía entrevista; logs pedidos/stock; "
            "CU-T09 matriz consistencia ES3."
        ),
        "refs": "Melville 2004; Choi 2018; Lepenioti 2020; Verhoef 2021; Soto-Acosta 2020.",
    },
]

VALIDATION_ROWS = [
    ("8 dimensiones × 6 indicadores = 48", "Estructura académica estándar (Tesify/UNCP)", "Cumple"),
    ("24 ítems Likert → VD-D1…D4 (6 c/u)", "generar_instrumento_investigacion.py", "Cumple"),
    ("CU-T09 ES1 ↔ VI-D1 (digitalización + disponibilidad)", "validar_coherencia_matrices_tesis.py", "Cumple"),
    ("CU-T09 ES2 ↔ VI-D2 (MAE/RMSE/horizonte 7-15-30)", "idem", "Cumple"),
    ("CU-T09 ES3 ↔ VD-D3/D4 (IRE, quiebres, Likert)", "idem", "Cumple"),
    ("IRE escala 0–100; pesos 40/35/25", "documentacion/07-modulo-ia-riesgo-empresarial.md", "Cumple"),
    ("Alertas IRE ≥51 (alto) y ≥76 (crítico)", "Panel AdminPredictions", "Cumple"),
    ("Sin Firestore; Supabase PostgreSQL", "09-implementacion-despliegue-ci.md", "Cumple"),
    ("RandomForestRegressor + fallback promedio móvil", "ai-service/models/demand/features_ml.py", "Cumple"),
    ("AUC-ROC solo condicional (eventos etiquetados)", "VD-D2 indicador 6 — doc 07", "Cumple"),
    ("SUS no confundido con Likert empleados", "SUS = anexo ISO usuarios finales (aparte)", "Cumple"),
]


def set_cell_text(cell, text: str, bold: bool = False, font_size: float = 6.4) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_widths(table, widths: list[float]) -> None:
    for row in table.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)


def row_to_table_cells(row: dict[str, str]) -> list[str]:
    var_label = (
        "VI: Sistema web de comercio electrónico con modelo de Inteligencia Artificial"
        if row["var"] == "VI"
        else "VD: Predicción del riesgo empresarial mediante el IRE"
    )
    return [
        var_label,
        row["concept"] if row["var"] == "VI" and row["dim"].startswith("D1") else (
            row["concept"] if row["var"] == "VD" and row["dim"].startswith("D1") else "—"
        ),
        row["oper"],
        row["dim"],
        row["indicators"],
        row["scale"],
        row["technique"],
        row["instrument"],
        row["refs"],
    ]


def build_docx() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.top_margin = Inches(0.32)
    sec.bottom_margin = Inches(0.32)
    sec.left_margin = Inches(0.32)
    sec.right_margin = Inches(0.32)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("MATRIZ DE OPERACIONALIZACIÓN DE VARIABLES")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(12)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run(TITULO)
    r.font.name = "Arial"
    r.font.size = Pt(8)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = note.add_run(
        "Traducción de variables abstractas en dimensiones e indicadores observables "
        "(def. conceptual → operacional → dimensión → indicador → escala → técnica → instrumento). "
        "VI y VD: 4 dimensiones × 6 indicadores = 48 unidades. "
        "24 ítems Likert miden percepción del personal sobre VD (6 por dimensión); "
        "indicadores técnicos (MAE, IRE, logs) son complemento objetivo. "
        "Coherencia CU-T09 · CU-T10 · doc 07. Validación: validar_matriz_operacionalizacion_variables.py "
        "y validar_coherencia_matrices_tesis.py."
    )
    r.font.name = "Arial"
    r.font.size = Pt(7)
    r.italic = True

    meta = doc.add_paragraph()
    r = meta.add_run(
        "Tesista: _________________________   Asesor: Dr. Maglioni Arana Caparachin   "
        "Versión: 2.1 auditada   Fecha: junio 2026"
    )
    r.font.name = "Arial"
    r.font.size = Pt(7)

    table = doc.add_table(rows=len(ROWS) + 1, cols=len(HEADERS))
    table.style = "Table Grid"
    table.autofit = False
    set_widths(table, [1.15, 1.35, 1.35, 1.35, 2.05, 1.15, 1.2, 1.45, 1.55])

    for i, h in enumerate(HEADERS):
        set_cell_text(table.cell(0, i), h, bold=True, font_size=6.8)
        shade(table.cell(0, i), "D9EAF7")

    for ri, row in enumerate(ROWS, start=1):
        cells = row_to_table_cells(row)
        # Def conceptual: show full text only on first row of VI and first row of VD
        if row["var"] == "VI" and not row["dim"].startswith("D1"):
            cells[1] = "— (ver fila VI-D1)"
        if row["var"] == "VD" and not row["dim"].startswith("D1"):
            cells[1] = "— (ver fila VD-D1)"
        fill = "F4F9FD" if row["var"] == "VI" else "FFF9EE"
        for ci, val in enumerate(cells):
            set_cell_text(table.cell(ri, ci), val, font_size=6.0)
            shade(table.cell(ri, ci), fill)

    doc.add_paragraph()
    vt_title = doc.add_paragraph()
    vt_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = vt_title.add_run("VALIDACIÓN — COHERENCIA CON CU-T09 E INSTRUMENTO LIKERT")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(9)

    vt = doc.add_table(rows=len(VALIDATION_ROWS) + 1, cols=3)
    vt.style = "Table Grid"
    set_widths(vt, [2.5, 4.5, 1.2])
    for i, h in enumerate(["Criterio", "Evidencia", "Estado"]):
        set_cell_text(vt.cell(0, i), h, bold=True, font_size=7)
        shade(vt.cell(0, i), "D9EAF7")
    for ri, (a, b, c) in enumerate(VALIDATION_ROWS, start=1):
        set_cell_text(vt.cell(ri, 0), a, font_size=6.8)
        set_cell_text(vt.cell(ri, 1), b, font_size=6.8)
        set_cell_text(vt.cell(ri, 2), c, font_size=6.8)
        shade(vt.cell(ri, 2), "E2F0D9")

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = foot.add_run(
        "Nota metodológica: la operacionalización no sustituye la matriz de consistencia (CU-T09) "
        "ni el instrumento de investigación; las tres deben leerse en conjunto. "
        "Riesgo empresarial = comercial-operativo (no quiebra contable). "
        "SUS (documentacion/plantillas/instrumento-sus-calzatura-vilchez.md) mide usabilidad de "
        "clientes finales — anexo ISO, no sustituye Likert del personal ni indicadores técnicos VI."
    )
    r.font.name = "Arial"
    r.font.size = Pt(7)
    r.italic = True

    for target in (OUT_DOCX_V3, OUT_DOCX_V2, OUT_DOCX):
        try:
            doc.save(target)
            print(f"OK: {target}")
            return
        except PermissionError:
            if target == OUT_DOCX:
                print(f"AVISO: {OUT_DOCX} abierto — guardado en {OUT_DOCX_V2}")
            continue
    raise PermissionError("Cierre el DOCX en Word e intente de nuevo.")


def build_csv() -> None:
    OUT_CSV.parent.mkdir(parents=True, exist_ok=True)
    fields = ["Tipo"] + HEADERS
    with OUT_CSV.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=fields, delimiter=";")
        w.writeheader()
        for row in ROWS:
            cells = row_to_table_cells(row)
            w.writerow(dict(zip(["Tipo", *HEADERS], [row["var"], *cells])))


def main() -> None:
    build_docx()
    build_csv()
    print(f"OK: {OUT_CSV}")


if __name__ == "__main__":
    main()
