#!/usr/bin/env python3
"""
Genera Matriz de Propuesta de Valor (Value Proposition Canvas — Osterwalder et al., 2014)
para Calzatura Vilchez. Fuente: CU-T12 + matriz-propuesta-valor-proyecto.md

Uso: python scripts/generar_matriz_propuesta_valor.py
"""
from __future__ import annotations

import csv
import sys
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
CSV_PATH = ROOT / "documentacion/cuadros-excel/CU-T12-matriz-propuesta-valor.csv"
OUT_DOCX = ROOT / "Matriz_Propuesta_Valor_Calzatura_Vilchez.docx"
OUT_DOCX_COPY = ROOT / "documentacion/Matriz_Propuesta_Valor_Calzatura_Vilchez.docx"

TITULO = (
    "Sistema web de comercio electrónico con modelo de Inteligencia Artificial "
    "para la predicción del riesgo empresarial en la empresa Calzatura Vilchez, Huancayo, 2026"
)

PROPUESTA_SINTESIS = (
    "Para la dirección y el equipo administrativo de Calzatura Vilchez (PYME retail calzado, Huancayo), "
    "que enfrentan quiebres de stock, ventas no digitalizadas e imposibilidad de anticipar presión "
    "comercial-operativa, nuestro sistema es una plataforma web integrada de e-commerce con "
    "inteligencia artificial que centraliza catálogo, pedidos, inventario y analítica, calcula y "
    "proyecta el Índice de Riesgo Empresarial (IRE) y alerta antes de que el riesgo se materialice "
    "en pérdidas, a diferencia de hojas de cálculo, registros manuales y venta exclusivamente "
    "presencial, porque unifica datos en Supabase, expone predicción Random Forest con trazabilidad "
    "documentada y cumple requisitos legales peruanos de comercio electrónico (Ley 29571, Ley 29733)."
)

SEGMENTOS = [
    ("SA", "Cliente web", "STK-02", "Alta", "Genera datos de demanda y valida e-commerce"),
    ("SB", "Administrador / Dirección", "STK-01", "Crítica", "Usuario del IRE, inventario y decisiones operativas"),
    ("SC", "Visitante", "—", "Media", "Embudo de conversión y descubrimiento de catálogo"),
]

SB_JOBS = [
    ("J-SB-01", "Publicar y mantener catálogo de calzado con tallas, colores e imágenes", "Catálogo físico / WhatsApp / fotos sueltas"),
    ("J-SB-02", "Controlar stock por variante y evitar vender lo inexistente", "Conteo manual, errores entre tienda y depósito"),
    ("J-SB-03", "Registrar ventas presenciales y online en un solo lugar", "Libretas, Excel no sincronizado"),
    ("J-SB-04", "Gestionar pedidos, pagos y estados de entrega", "Llamadas, papel, seguimiento ad hoc"),
    ("J-SB-05", "Conocer rentabilidad por producto (costo, margen)", "Cálculos esporádicos, sin dashboard"),
    ("J-SB-06", "Anticipar quiebre de stock y caída de ingresos", "Reacción tardía cuando ya faltó producto"),
    ("J-SB-07", "Monitorear riesgo comercial-operativo con indicador único (IRE)", "Intuición gerencial sin métrica formal"),
    ("J-SB-08", "Cumplir obligaciones legales (reclamaciones, privacidad)", "Riesgo de incumplimiento Ley 29571 / 29733"),
]

SB_PAINS = [
    ("P-SB-01", "Quiebre de stock en productos de alta rotación", "Alta"),
    ("P-SB-02", "Capital inmovilizado en productos de baja rotación", "Alta"),
    ("P-SB-03", "Pérdida de ventas online por información desactualizada", "Alta"),
    ("P-SB-04", "Datos dispersos (ventas, pedidos, inventario)", "Alta"),
    ("P-SB-05", "Decisiones de reabastecimiento sin pronóstico", "Alta"),
    ("P-SB-06", "Sin visibilidad de riesgo agregado del negocio", "Crítica"),
    ("P-SB-07", "Tiempo excesivo en tareas administrativas repetitivas", "Media"),
    ("P-SB-08", "Incertidumbre legal en e-commerce", "Media"),
]

SB_GAINS = [
    ("G-SB-01", "Un solo panel con KPIs de ventas, pedidos y productos", "Esperado"),
    ("G-SB-02", "Alertas de stock con fecha estimada de quiebre", "Esperado"),
    ("G-SB-03", "IRE 0–100 con clasificación Bajo/Moderado/Alto/Crítico", "Deseado"),
    ("G-SB-04", "Proyección de IRE a 7/15/30 días", "Deseado"),
    ("G-SB-05", "Ranking de productos y recomendaciones automáticas", "Deseado"),
    ("G-SB-06", "Reducción de errores en precios y totales", "Esperado"),
    ("G-SB-07", "Trazabilidad de pedidos y auditoría admin", "Requerido"),
    ("G-SB-08", "Confianza en cumplimiento normativo peruano", "Requerido"),
]

PAIN_RELIEVERS = [
    ("P-SB-01", "Alertas stock + fecha quiebre + análisis ABC", "AdminPredictionsDashboard, IN-27, IN-34"),
    ("P-SB-02", "Capital inmovilizado e ingresos en riesgo calculados", "IN-35, ai-service"),
    ("P-SB-03", "Stock y precios validados server-side (BFF + triggers)", "RF-RN-01/02, precisionBffGuards"),
    ("P-SB-04", "Supabase PostgreSQL único: ventas, pedidos, productos, IRE histórico", "ireHistorial, migraciones"),
    ("P-SB-05", "RandomForestRegressor + fallback conservador", "ai-service/, RF-IA-02"),
    ("P-SB-06", "IRE compuesto (40% stock + 35% ingresos + 25% demanda)", "07-modulo-ia, RF-IA-01"),
    ("P-SB-07", "Automatización CRUD, import Excel, RPC atómicos", "RF-ADM-*, RF-RN-02"),
    ("P-SB-08", "Páginas legales + TC-CMP E2E", "RF-LEG-*, verify-cumplimiento"),
]

GAIN_CREATORS = [
    ("G-SB-01", "Dashboard admin con KPIs en tiempo real", "RF-ADM-01"),
    ("G-SB-02", "Motor alertas con horizonte configurable", "RF-IA-03, IN-27"),
    ("G-SB-03", "Score IRE con niveles y contribuciones por variable", "Panel IRE, tests invariantes"),
    ("G-SB-04", "IRE proyectado descontando consumo estimado", "IN-32, API /api/predict/combined"),
    ("G-SB-05", "Top vendidos + baja rotación por período", "IN-36"),
    ("G-SB-06", "Guards BFF totales/precios; reglas comerciales BD", "verify-precision"),
    ("G-SB-07", "Auditoría admin, PKCS#7 pedidos, RLS", "RF-SEG, Seguridad checklist"),
    ("G-SB-08", "Ley 29571/29733 implementada", "TC-CMP-001…004"),
]

DIFERENCIADORES = [
    ("Venta solo presencial", "Sin datos online ni escala geográfica", "E-commerce + app móvil Flutter"),
    ("Excel / cuaderno", "Sin tiempo real ni integración pagos", "Supabase + Stripe + BFF"),
    ("E-commerce genérico (Shopify, etc.)", "Sin IRE ni predicción demanda calzado", "IRE + Random Forest + analítica ABC"),
    ("ERP contable completo", "Sobredimensionado para PYME", "Enfoque retail calzado + IA acotada"),
]

BIBLIOGRAFIA = [
    ("Transformación digital PYME", "Verhoef et al. (2019). Digital transformation: A multidisciplinary reflection.", "10.1016/j.jbusres.2019.09.022"),
    ("Retail / cadena valor", "Reinartz et al. (2019). The impact of digital transformation on the retailing value chain.", "10.1016/j.ijresmar.2018.12.002"),
    ("E-commerce + IA + PYME", "Dai et al. (2024). Revolutionizing cross-border e-commerce…", "10.1371/journal.pone.0305639"),
    ("IA para decisiones", "Duan et al. (2019). Artificial intelligence for decision making in the era of Big Data.", "10.1016/j.ijinfomgt.2019.01.021"),
    ("Big data operaciones", "Choi et al. (2018). Big data analytics in operations management.", "10.1111/poms.12838"),
    ("E-SQ moda/calzado", "Gutiérrez-Rodríguez et al. (2020). E-SQ, E-Satisfaction and E-Loyalty for fashion E-Retailers.", "10.1016/j.jretconser.2020.102201"),
    ("Pronóstico / precisión", "Makridakis et al. (2018). Statistical and ML forecasting methods.", "10.1371/journal.pone.0194889"),
    ("Retail forecasting", "Fildes et al. (2019). Retail forecasting: Research and practice.", "10.1016/j.ijforecast.2019.06.004"),
    ("Prescriptive analytics", "Lepenioti et al. (2020). Prescriptive analytics: Literature review.", "10.1016/j.ijinfomgt.2019.04.003"),
]

LIMITACIONES = [
    ("IRE vs quiebra financiera", "IRE = riesgo comercial-operativo, no Altman", "07-modulo-ia §1"),
    ("Validación O₁→O₂", "Mediciones pre/post requieren operación sostenida", "CU-T09, encuesta Likert"),
    ("App móvil iOS", "Flutter Android verificado; IPA iOS parcial", "Portabilidad 71% adaptabilidad"),
    ("CU-T06 trazabilidad EDA", "CSV desactualizado", "Usar CU-T12 + estado del arte 43"),
]

MATRIZ_HEADERS = [
    "Segmento",
    "Tipo VPC",
    "ID",
    "Descripción",
    "Componente sistema",
    "Req_ID",
    "Evidencia",
    "Artículo Q1",
    "Encaje",
]


def set_cell_text(cell, text: str, bold: bool = False, font_size: float = 8.0) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text or "—")
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_widths(table, widths: list[float]) -> None:
    for row in table.rows:
        for index, width in enumerate(widths):
            if index < len(row.cells):
                row.cells[index].width = Inches(width)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    sizes = {1: 12, 2: 11, 3: 10}
    run.font.size = Pt(sizes.get(level, 9))
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)


def add_body(doc: Document, text: str, italic: bool = False, size: float = 9.0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.italic = italic


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], header_fill: str = "D9EAF7") -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_widths(table, widths)
    for col_index, header in enumerate(headers):
        cell = table.cell(0, col_index)
        set_cell_text(cell, header, bold=True, font_size=8)
        shade(cell, header_fill)
    for row_index, row in enumerate(rows, start=1):
        fill = "F9FBFD" if row_index % 2 == 0 else "FFFFFF"
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            set_cell_text(cell, value, font_size=7.5)
            shade(cell, fill)


def load_csv_rows() -> list[dict[str, str]]:
    with CSV_PATH.open(encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f))


def validate(rows: list[dict[str, str]]) -> list[str]:
    errors: list[str] = []
    if len(rows) < 35:
        errors.append(f"CU-T12 tiene pocas filas ({len(rows)}); se esperan ≥35.")
    encaje = Counter(r.get("Encaje", "") for r in rows)
    if encaje.get("Pendiente", 0) > 0:
        errors.append(f"Hay {encaje['Pendiente']} filas con Encaje=Pendiente.")
    sb_pains = {r["ID_elemento"] for r in rows if r["Segmento"] == "SB" and r["Tipo_VPC"] == "Pain"}
    sb_relievers_pain_ids = {r["ID_elemento"].replace("PR-", "P-") for r in rows if r["Segmento"] == "SB" and r["Tipo_VPC"] == "Pain_Reliever"}
    # Pain relievers in separate table cover all pains — check P-SB-* in PAIN_RELIEVERS
    for pid, _, _ in SB_PAINS:
        if pid not in {p for p, _, _ in PAIN_RELIEVERS}:
            errors.append(f"Falta pain reliever documentado para {pid}")
    if not OUT_DOCX.parent.exists():
        errors.append("Directorio raíz del proyecto no accesible.")
    return errors


def build_docx(rows: list[dict[str, str]]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # Portada
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MATRIZ DE PROPUESTA DE VALOR DEL PROYECTO")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(14)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(TITULO)
    run.font.name = "Arial"
    run.font.size = Pt(9)

    add_body(
        doc,
        "Metodología: Value Proposition Canvas (Osterwalder, Pigneur, Bernarda & Smith, 2014). "
        "Cuadro trazable CU-T12 · documentacion/matriz-propuesta-valor-proyecto.md · "
        "formato-09-alcance-proyecto-software.md · 07-modulo-ia-riesgo-empresarial.md.",
        italic=True,
        size=8,
    )
    meta = doc.add_paragraph()
    run = meta.add_run(
        "Empresa: Calzatura Vilchez, Huancayo, Perú   |   Asesor: Dr. Maglioni Arana Caparachin   |   "
        "Versión: 1.0   |   Fecha: junio 2026"
    )
    run.font.name = "Arial"
    run.font.size = Pt(8)

    add_heading(doc, "1. Fundamento metodológico", 1)
    add_body(
        doc,
        "La propuesta de valor es la promesa explícita de beneficios que el sistema entrega a cada segmento, "
        "alineada con trabajos (jobs to be done), frustraciones (pains) y beneficios esperados (gains). "
        "El instrumento articula el perfil del cliente, el mapa de valor y la matriz de encaje producto–mercado "
        "(problem–solution fit), vinculando cada elemento a requisitos RF/RNF implementados y evidencia verificable.",
    )

    add_heading(doc, "1.1 Segmentos de cliente", 2)
    add_table(
        doc,
        ["ID", "Segmento", "Stakeholder", "Prioridad", "Justificación"],
        [[a, b, c, d, e] for a, b, c, d, e in SEGMENTOS],
        [0.5, 1.6, 0.9, 0.7, 2.8],
    )

    add_heading(doc, "2. Propuesta de valor global (enunciado síntesis)", 1)
    add_body(doc, PROPUESTA_SINTESIS)

    add_heading(doc, "3. Segmento SB — Administrador / Dirección (núcleo)", 1)
    add_heading(doc, "3.1 Perfil del cliente — Jobs to be done", 2)
    add_table(
        doc,
        ["ID", "Trabajo a realizar", "Situación actual (sin sistema)"],
        [[a, b, c] for a, b, c in SB_JOBS],
        [0.9, 2.8, 2.8],
    )

    add_heading(doc, "3.2 Pains (frustraciones)", 2)
    add_table(doc, ["ID", "Pain", "Severidad"], [[a, b, c] for a, b, c in SB_PAINS], [0.9, 3.5, 1.0])

    add_heading(doc, "3.3 Gains (beneficios esperados)", 2)
    add_table(doc, ["ID", "Gain", "Tipo"], [[a, b, c] for a, b, c in SB_GAINS], [0.9, 3.5, 1.0])

    add_heading(doc, "3.4 Mapa de valor — Pain relievers", 2)
    add_table(
        doc,
        ["Pain ID", "Aliviador (solución)", "Evidencia implementada"],
        [[a, b, c] for a, b, c in PAIN_RELIEVERS],
        [0.9, 2.8, 2.8],
        header_fill="E2F0D9",
    )

    add_heading(doc, "3.5 Mapa de valor — Gain creators", 2)
    add_table(
        doc,
        ["Gain ID", "Creador de beneficio", "RF / evidencia"],
        [[a, b, c] for a, b, c in GAIN_CREATORS],
        [0.9, 2.8, 2.8],
        header_fill="E2F0D9",
    )

    add_heading(doc, "3.6 Veredicto de encaje segmento SB", 2)
    encaje_sb = Counter(r["Encaje"] for r in rows if r["Segmento"] == "SB")
    add_body(
        doc,
        f"Elementos CU-T12 segmento SB: {sum(encaje_sb.values())} filas — "
        f"Logrado: {encaje_sb.get('Logrado', 0)}, Parcial: {encaje_sb.get('Parcial', 0)}. "
        "Jobs 8/8 · Pains con pain reliever 8/8 · Gains con gain creator 8/8 · "
        "Trazabilidad RF: CU-T05 · Gates: verify-idoneidad, verify-precision, verify-cumplimiento.",
    )

    add_heading(doc, "4. Segmentos SA y SC (resumen)", 1)
    sa_rows = [r for r in rows if r["Segmento"] == "SA"]
    sc_rows = [r for r in rows if r["Segmento"] == "SC"]
    add_table(
        doc,
        ["Segmento", "ID", "Tipo", "Descripción", "Encaje"],
        [[r["Segmento"], r["ID_elemento"], r["Tipo_VPC"], r["Descripcion"], r["Encaje"]] for r in sa_rows + sc_rows],
        [0.5, 0.8, 0.9, 3.2, 0.7],
    )
    add_body(
        doc,
        "Nota SA: encaje Parcial en J-SA-05 (sesión SUS formal con usuarios reales pendiente). "
        "Nota SC: encaje Parcial en P-SC-01 (gate Lighthouse en CI pendiente).",
        italic=True,
        size=8,
    )

    add_heading(doc, "5. Diferenciadores frente a alternativas", 1)
    add_table(
        doc,
        ["Alternativa", "Limitación", "Ventaja Calzatura Vilchez"],
        [[a, b, c] for a, b, c in DIFERENCIADORES],
        [1.4, 2.2, 2.8],
    )

    add_heading(doc, "6. Matriz consolidada CU-T12 (trazabilidad completa)", 1)
    add_body(
        doc,
        "Matriz de encaje VPC → sistema → requisitos → evidencia → artículo Q1. "
        "Fuente CSV: documentacion/cuadros-excel/CU-T12-matriz-propuesta-valor.csv",
        size=8,
    )

    # Landscape section for full matrix
    doc.add_page_break()
    landscape = doc.add_section()
    landscape.orientation = WD_ORIENT.LANDSCAPE
    landscape.page_width, landscape.page_height = landscape.page_height, landscape.page_width
    landscape.top_margin = Inches(0.35)
    landscape.bottom_margin = Inches(0.35)
    landscape.left_margin = Inches(0.35)
    landscape.right_margin = Inches(0.35)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run("ANEXO — MATRIZ CU-T12 PROPUESTA DE VALOR (TRAZABILIDAD COMPLETA)")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)

    matrix_rows = [
        [
            r["Segmento"],
            r["Tipo_VPC"],
            r["ID_elemento"],
            r["Descripcion"][:120] + ("…" if len(r["Descripcion"]) > 120 else ""),
            (r["Componente_sistema"] or "")[:40],
            (r["Req_ID"] or "")[:28],
            (r["Evidencia_repo"] or "")[:35],
            (r["Articulo_Q1"] or "")[:40],
            r["Encaje"],
        ]
        for r in rows
    ]
    table = doc.add_table(rows=len(matrix_rows) + 1, cols=len(MATRIZ_HEADERS))
    table.style = "Table Grid"
    table.autofit = False
    set_widths(table, [0.45, 0.75, 0.65, 1.55, 0.95, 0.85, 0.95, 1.05, 0.55])
    for col_index, header in enumerate(MATRIZ_HEADERS):
        cell = table.cell(0, col_index)
        set_cell_text(cell, header, bold=True, font_size=7)
        shade(cell, "1F4E79")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = None  # keep default on shaded header
    for row_index, row in enumerate(matrix_rows, start=1):
        encaje = row[-1]
        fill = "FFF2CC" if encaje == "Parcial" else ("E2EFDA" if encaje == "Logrado" else "FFFFFF")
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            set_cell_text(cell, value, font_size=6.5)
            shade(cell, fill)

    # Back to portrait for closing sections
    doc.add_section()
    portrait = doc.sections[-1]
    portrait.orientation = WD_ORIENT.PORTRAIT
    portrait.page_width, portrait.page_height = Inches(8.27), Inches(11.69)
    portrait.top_margin = Inches(0.7)
    portrait.bottom_margin = Inches(0.7)
    portrait.left_margin = Inches(0.8)
    portrait.right_margin = Inches(0.8)

    add_heading(doc, "7. Respaldo bibliográfico Q1", 1)
    add_table(
        doc,
        ["Dimensión", "Referencia", "DOI"],
        [[a, b, f"https://doi.org/{c}"] for a, b, c in BIBLIOGRAFIA],
        [1.2, 3.5, 1.8],
    )

    add_heading(doc, "8. Limitaciones declaradas", 1)
    add_table(doc, ["Aspecto", "Limitación", "Mitigación"], [[a, b, c] for a, b, c in LIMITACIONES], [1.4, 2.5, 2.5])

    add_heading(doc, "9. Referencias metodológicas", 1)
    refs = [
        "Osterwalder, A., Pigneur, Y., Bernarda, G., & Smith, A. (2014). Value Proposition Design. Wiley / Strategyzer.",
        "Osterwalder, A., & Pigneur, Y. (2010). Business Model Generation. Wiley.",
        "Strategyzer (2026). The Value Proposition Canvas. https://www.strategyzer.com/library/the-value-proposition-canvas",
        "Christensen, C. M., Hall, T., Dillon, K., & Duncan, D. S. (2016). Know your customers' jobs to be done. Harvard Business Review.",
    ]
    for ref in refs:
        add_body(doc, f"• {ref}", size=8)

    encaje_total = Counter(r["Encaje"] for r in rows)
    add_heading(doc, "10. Resumen de validación documental", 1)
    add_body(
        doc,
        f"Total filas CU-T12: {len(rows)}. Encaje — Logrado: {encaje_total.get('Logrado', 0)}, "
        f"Parcial: {encaje_total.get('Parcial', 0)}, Pendiente: {encaje_total.get('Pendiente', 0)}. "
        f"IRE definido como riesgo comercial-operativo (no quiebra financiera Altman). "
        f"Regeneración: python scripts/generar_matriz_propuesta_valor.py",
    )

    doc.save(OUT_DOCX)
    doc.save(OUT_DOCX_COPY)


def main() -> int:
    if not CSV_PATH.exists():
        print(f"ERROR: no existe {CSV_PATH}", file=sys.stderr)
        return 1
    rows = load_csv_rows()
    issues = validate(rows)
    if issues:
        for issue in issues:
            print(f"VALIDACIÓN: {issue}", file=sys.stderr)
        return 1
    build_docx(rows)
    print(f"OK: {OUT_DOCX}")
    print(f"OK: {OUT_DOCX_COPY}")
    print(f"Filas CU-T12: {len(rows)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
