from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))
from datos_instrumento_likert import (  # noqa: E402
    INSTRUMENT_NAME,
    ITEMS,
    LIKERT_LABELS,
    VD_VARIABLE,
)

OUT = ROOT / "Instrumento_Investigacion_Calzatura_Vilchez_MEJORADO.docx"


def set_run_font(run, size: float, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def add_paragraph(doc: Document, text: str, size: float = 10, bold: bool = False, align=None):
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size, bold)
    return paragraph


def set_cell_text(cell, text: str, bold: bool = False, font_size: float = 8.0, align=None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    set_run_font(run, font_size, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_widths(table, widths: list[float]) -> None:
    for row in table.rows:
        for index, width in enumerate(widths):
            if index < len(row.cells):
                row.cells[index].width = Inches(width)


def dim_fill(n: int) -> str:
    if n <= 6:
        return "F4F9FD"
    if n <= 12:
        return "FFF9EE"
    if n <= 18:
        return "F6F7F2"
    return "FDF4F4"


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.55)
section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.6)
section.right_margin = Inches(0.6)

add_paragraph(doc, "ANEXO 04. INSTRUMENTO DE INVESTIGACION", 12, True, WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, INSTRUMENT_NAME, 10, True, WD_ALIGN_PARAGRAPH.CENTER)

add_paragraph(doc, "I. Ficha tecnica del instrumento", 10, True)
tech_rows = [
    ("Nombre del instrumento", INSTRUMENT_NAME),
    ("Variable evaluada", VD_VARIABLE + "."),
    (
        "Objetivo",
        "Recolectar la percepcion del personal de Calzatura Vilchez sobre la capacidad de "
        "anticipar, validar y gestionar el riesgo empresarial antes y despues de la implementacion del sistema.",
    ),
    ("Tecnica", "Encuesta estructurada."),
    ("Tipo de escala", "Escala ordinal tipo Likert de 5 puntos."),
    ("Numero de items", "24 items distribuidos en 4 dimensiones, con 6 items por dimension."),
    ("Poblacion de aplicacion", "Personal directivo y operativo de Calzatura Vilchez."),
    (
        "Momento de aplicacion",
        "Preprueba y posprueba. El mismo instrumento se aplica antes y despues de la implementacion.",
    ),
    (
        "Validez",
        "Juicio de expertos (min. 3). Coeficiente V de Aiken >= 0.70 (pertinencia, relevancia, claridad).",
    ),
    (
        "Confiabilidad",
        "Alfa de Cronbach por dimension y total. Criterio minimo: alfa >= 0.70.",
    ),
    (
        "Sustento academico",
        "Alineado con CU-T09, CU-T10, CU-T11 y 42 articulos Q1 del estado del arte.",
    ),
]
tech_table = doc.add_table(rows=len(tech_rows), cols=2)
tech_table.style = "Table Grid"
tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_widths(tech_table, [2.2, 4.8])
for row_index, (label, value) in enumerate(tech_rows):
    set_cell_text(tech_table.cell(row_index, 0), label, True, 8)
    set_cell_text(tech_table.cell(row_index, 1), value, False, 8)
    shade(tech_table.cell(row_index, 0), "D9EAF7")

add_paragraph(doc, "II. Datos generales del participante", 10, True)
data_rows = [
    ("Codigo del participante", "____________________________"),
    ("Cargo o area", "____________________________"),
    ("Fecha de aplicacion", "____ / ____ / ______"),
    ("Momento de aplicacion", "[  ] Preprueba     [  ] Posprueba"),
]
data_table = doc.add_table(rows=len(data_rows), cols=2)
data_table.style = "Table Grid"
data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_widths(data_table, [2.2, 4.8])
for row_index, (label, value) in enumerate(data_rows):
    set_cell_text(data_table.cell(row_index, 0), label, True, 8)
    set_cell_text(data_table.cell(row_index, 1), value, False, 8)
    shade(data_table.cell(row_index, 0), "F4F9FD")

add_paragraph(doc, "III. Instrucciones", 10, True)
add_paragraph(
    doc,
    "Estimado(a) participante: marque con una X la alternativa que mejor represente su percepcion. "
    "Si el cuestionario se aplica como preprueba, responda considerando la forma actual de gestion "
    "(metodo actual). Si se aplica como posprueba, responda considerando el sistema web con modelo "
    "de Inteligencia Artificial ya implementado. La informacion sera utilizada solo con fines academicos.",
    9,
)

scale_table = doc.add_table(rows=2, cols=5)
scale_table.style = "Table Grid"
scale_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_widths(scale_table, [1.3, 1.3, 1.3, 1.3, 1.3])
for index, value in enumerate(["1", "2", "3", "4", "5"]):
    set_cell_text(scale_table.cell(0, index), value, True, 8, WD_ALIGN_PARAGRAPH.CENTER)
    shade(scale_table.cell(0, index), "D9EAF7")
for index, value in enumerate(LIKERT_LABELS):
    set_cell_text(scale_table.cell(1, index), value, False, 7.5, WD_ALIGN_PARAGRAPH.CENTER)

add_paragraph(doc, "IV. Cuestionario", 10, True)

questionnaire = doc.add_table(rows=len(ITEMS) + 1, cols=8)
questionnaire.style = "Table Grid"
questionnaire.alignment = WD_TABLE_ALIGNMENT.CENTER
set_widths(questionnaire, [0.35, 1.6, 3.35, 0.35, 0.35, 0.35, 0.35, 0.35])
headers = ["N", "Dimension", "Item", "1", "2", "3", "4", "5"]
for col_index, header in enumerate(headers):
    set_cell_text(questionnaire.cell(0, col_index), header, True, 7.3, WD_ALIGN_PARAGRAPH.CENTER)
    shade(questionnaire.cell(0, col_index), "D9EAF7")

for item in ITEMS:
    row_index = item["n"]
    values = [str(row_index), item["dim"], item["text"], "[  ]", "[  ]", "[  ]", "[  ]", "[  ]"]
    fill = dim_fill(row_index)
    for col_index, value in enumerate(values):
        align = WD_ALIGN_PARAGRAPH.CENTER if col_index in [0, 3, 4, 5, 6, 7] else WD_ALIGN_PARAGRAPH.JUSTIFY
        set_cell_text(questionnaire.cell(row_index, col_index), value, False, 6.7, align)
        shade(questionnaire.cell(row_index, col_index), fill)

add_paragraph(doc, "V. Observaciones", 10, True)
for _ in range(3):
    add_paragraph(doc, "______________________________________________________________________________________________", 9)

add_paragraph(doc, "Gracias por su colaboracion.", 9, True, WD_ALIGN_PARAGRAPH.CENTER)

doc.save(OUT)
print(OUT)
