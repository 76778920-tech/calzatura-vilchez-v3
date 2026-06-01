from __future__ import annotations

import math
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "artifacts" / "diagramas"
MERMAID_DIR = OUT_DIR / "mermaid"
IMG_DIR = OUT_DIR / "imagenes"
DOC_FILE = ROOT / "artifacts" / "documentos" / "Documento_4_Diagramas_UML_Modelado_Calzatura_Vilchez_COMPLETO.docx"
DOWNLOAD_DOC = Path.home() / "Downloads" / DOC_FILE.name

W, H = 1800, 1150
BLUE = "#1F4E79"
LIGHT_BLUE = "#D9EAF7"
GREEN = "#B7E1CD"
LIGHT_GREEN = "#E2F0D9"
YELLOW = "#FFF2CC"
ORANGE = "#FCE4D6"
PURPLE = "#EADCF8"
GRAY = "#F2F2F2"
DARK = "#1F1F1F"
LINE = "#4D4D4D"
WHITE = "#FFFFFF"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


F_TITLE = font(42, True)
F_H = font(28, True)
F_B = font(22)
F_S = font(18)
F_XS = font(16)
F_MONO = font(16)


def new_canvas(title: str, subtitle: str = "") -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (W, H), WHITE)
    d = ImageDraw.Draw(img)
    d.rectangle([0, 0, W, 90], fill=BLUE)
    d.text((40, 22), title, fill=WHITE, font=F_TITLE)
    if subtitle:
        d.text((42, 95), subtitle, fill=BLUE, font=F_B)
    return img, d


def text_size(d: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont) -> tuple[int, int]:
    box = d.textbbox((0, 0), text, font=fnt)
    return box[2] - box[0], box[3] - box[1]


def wrap_text(d: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = f"{current} {word}".strip()
        if text_size(d, trial, fnt)[0] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def centered_text(d: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fnt=F_B, fill=DARK, max_lines: int = 4):
    x1, y1, x2, y2 = box
    lines = wrap_text(d, text, fnt, x2 - x1 - 20)[:max_lines]
    total_h = len(lines) * (fnt.size + 5)
    y = y1 + ((y2 - y1) - total_h) // 2
    for line in lines:
        tw, th = text_size(d, line, fnt)
        d.text((x1 + ((x2 - x1) - tw) // 2, y), line, fill=fill, font=fnt)
        y += fnt.size + 5


def round_box(d: ImageDraw.ImageDraw, box, fill=GRAY, outline=LINE, radius=22, width=3, label: str | None = None, fnt=F_B):
    d.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)
    if label:
        centered_text(d, box, label, fnt=fnt)


def ellipse(d: ImageDraw.ImageDraw, box, fill=LIGHT_BLUE, outline=BLUE, width=3, label: str | None = None):
    d.ellipse(box, fill=fill, outline=outline, width=width)
    if label:
        centered_text(d, box, label, fnt=F_S, max_lines=3)


def arrow(d: ImageDraw.ImageDraw, start, end, fill=LINE, width=3, label: str | None = None, fnt=F_XS):
    d.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 16
    left = (end[0] - length * math.cos(angle - math.pi / 6), end[1] - length * math.sin(angle - math.pi / 6))
    right = (end[0] - length * math.cos(angle + math.pi / 6), end[1] - length * math.sin(angle + math.pi / 6))
    d.polygon([end, left, right], fill=fill)
    if label:
        mx = (start[0] + end[0]) // 2
        my = (start[1] + end[1]) // 2 - 22
        d.rectangle([mx - 170, my - 2, mx + 170, my + 24], fill=WHITE)
        centered_text(d, (mx - 165, my, mx + 165, my + 24), label, fnt=fnt, max_lines=1)


def actor(d: ImageDraw.ImageDraw, x: int, y: int, label: str):
    d.ellipse([x - 18, y, x + 18, y + 36], outline=BLUE, width=4)
    d.line([x, y + 36, x, y + 115], fill=BLUE, width=4)
    d.line([x - 45, y + 58, x + 45, y + 58], fill=BLUE, width=4)
    d.line([x, y + 115, x - 42, y + 170], fill=BLUE, width=4)
    d.line([x, y + 115, x + 42, y + 170], fill=BLUE, width=4)
    centered_text(d, (x - 95, y + 178, x + 95, y + 245), label, fnt=F_S)


def class_box(d, x, y, w, title, attrs, methods=None, fill=LIGHT_BLUE):
    methods = methods or []
    h = 58 + max(2, len(attrs)) * 24 + max(1, len(methods)) * 24 + 30
    d.rounded_rectangle([x, y, x + w, y + h], radius=12, fill=fill, outline=BLUE, width=3)
    d.rectangle([x, y, x + w, y + 50], fill=BLUE)
    centered_text(d, (x, y + 4, x + w, y + 46), title, fnt=F_S, fill=WHITE)
    yy = y + 60
    for attr in attrs:
        d.text((x + 14, yy), attr, fill=DARK, font=F_XS)
        yy += 24
    d.line([x, yy + 4, x + w, yy + 4], fill=BLUE, width=2)
    yy += 12
    for method in methods:
        d.text((x + 14, yy), method, fill=DARK, font=F_XS)
        yy += 24
    return (x, y, x + w, y + h)


def table_box(d, x, y, w, title, fields, fill=LIGHT_BLUE):
    h = 50 + len(fields) * 24 + 20
    d.rounded_rectangle([x, y, x + w, y + h], radius=12, fill=fill, outline=BLUE, width=3)
    d.rectangle([x, y, x + w, y + 44], fill=BLUE)
    centered_text(d, (x, y + 4, x + w, y + 40), title, fnt=F_S, fill=WHITE)
    yy = y + 54
    for field in fields:
        d.text((x + 12, yy), field, fill=DARK, font=F_XS)
        yy += 24
    return (x, y, x + w, y + h)


def sequence_diagram(title: str, subtitle: str, participants: list[str], messages: list[tuple[int, int, str, str]], filename: str):
    img, d = new_canvas(title, subtitle)
    top = 170
    left = 120
    lane_w = (W - 2 * left) // (len(participants) - 1)
    xs = [left + i * lane_w for i in range(len(participants))]
    for x, name in zip(xs, participants, strict=True):
        round_box(d, (x - 105, top, x + 105, top + 70), fill=LIGHT_BLUE, label=name, fnt=F_XS)
        d.line([x, top + 72, x, H - 70], fill="#A6A6A6", width=2)
    y = top + 120
    for src, dst, label, note in messages:
        start = (xs[src], y)
        end = (xs[dst], y)
        arrow(d, start, end, fill=BLUE, width=3, label=label)
        if note:
            d.text((min(start[0], end[0]) + 20, y + 16), note, fill=DARK, font=F_XS)
        y += 82
    save(img, filename)


def save(img: Image.Image, filename: str):
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    path = IMG_DIR / filename
    img.save(path)


def write_mermaid(name: str, content: str):
    MERMAID_DIR.mkdir(parents=True, exist_ok=True)
    (MERMAID_DIR / name).write_text(content.strip() + "\n", encoding="utf-8")


def diagram_use_case():
    img, d = new_canvas("Diagrama de casos de uso general", "Actores y funciones principales del sistema web")
    d.rounded_rectangle([300, 160, 1510, 1030], radius=30, outline=BLUE, width=4, fill="#FBFDFF")
    d.text((335, 180), "Sistema web Calzatura Vilchez", fill=BLUE, font=F_H)
    actor(d, 135, 250, "Visitante / Cliente")
    actor(d, 135, 670, "Administrador")
    actor(d, 1660, 310, "Trabajador")
    actor(d, 1660, 710, "Servicios externos")
    use_cases = [
        (410, 260, 710, 340, "Explorar catálogo y campañas"),
        (760, 260, 1060, 340, "Registrarse / iniciar sesión"),
        (1110, 260, 1410, 340, "Comprar: carrito, entrega y pago"),
        (410, 430, 710, 510, "Consultar pedidos y favoritos"),
        (760, 430, 1060, 510, "Gestionar productos, stock y precios"),
        (1110, 430, 1410, 510, "Administrar pedidos y usuarios"),
        (410, 620, 710, 700, "Importar / exportar Excel"),
        (760, 620, 1060, 700, "Registrar ventas físicas"),
        (1110, 620, 1410, 700, "Gestionar devoluciones"),
        (410, 810, 710, 890, "Auditoría y privacidad"),
        (760, 810, 1060, 890, "Predicción IA / IRE"),
        (1110, 810, 1410, 890, "Libro de reclamaciones"),
    ]
    for uc in use_cases:
        ellipse(d, uc[:4], label=uc[4])
    for point in [(410, 300), (760, 300), (1110, 300), (410, 470)]:
        arrow(d, (230, 365), point, width=2)
    for point in [(760, 470), (1110, 470), (410, 660), (760, 850), (1110, 850)]:
        arrow(d, (230, 785), point, width=2)
    for point in [(760, 660), (1110, 660)]:
        arrow(d, (1560, 430), point, width=2)
    for point in [(1110, 300), (760, 850), (1110, 850)]:
        arrow(d, (1560, 820), point, width=2)
    save(img, "01_casos_uso_general.png")
    write_mermaid(
        "01_casos_uso_general.mmd",
        """
flowchart LR
  Cliente[Visitante / Cliente] --> UC1((Explorar catálogo y campañas))
  Cliente --> UC2((Registrarse / iniciar sesión))
  Cliente --> UC3((Comprar: carrito, entrega y pago))
  Cliente --> UC4((Consultar pedidos y favoritos))
  Admin[Administrador] --> UC5((Gestionar productos, stock y precios))
  Admin --> UC6((Administrar pedidos y usuarios))
  Admin --> UC7((Importar / exportar Excel))
  Admin --> UC10((Auditoría y privacidad))
  Admin --> UC11((Predicción IA / IRE))
  Trabajador[Trabajador] --> UC8((Registrar ventas físicas))
  Trabajador --> UC9((Gestionar devoluciones))
  Externos[Servicios externos] --> UC3
  Externos --> UC11
  Admin --> UC12((Libro de reclamaciones))
""",
    )


def diagram_class():
    img, d = new_canvas("Diagrama de clases", "Entidades principales, responsabilidades y relaciones de dominio")
    boxes = {
        "Usuario": class_box(d, 70, 160, 270, "Usuario", ["uid", "email", "rol", "telefono", "dni_hash"], ["validarRol()", "actualizarPerfil()"], LIGHT_BLUE),
        "Producto": class_box(d, 410, 160, 300, "Producto", ["id", "nombre", "categoria", "precio", "stock", "tallaStock"], ["activar()", "actualizarStock()"], LIGHT_GREEN),
        "Pedido": class_box(d, 780, 160, 300, "Pedido", ["id", "userId", "items", "total", "estado", "metodoPago"], ["crear()", "cambiarEstado()"], YELLOW),
        "VentaDiaria": class_box(d, 1150, 160, 300, "VentaDiaria", ["id", "productoId", "cantidad", "precio", "canal"], ["registrar()", "devolver()"], ORANGE),
        "ProductoCodigo": class_box(d, 410, 480, 300, "ProductoCodigo", ["productoId", "codigo", "actualizadoEn"], ["validarUnicidad()"], LIGHT_GREEN),
        "ProductoFinanza": class_box(d, 70, 480, 270, "ProductoFinanza", ["productId", "costo", "margen", "precioSugerido"], ["calcularRango()"], LIGHT_GREEN),
        "Favorito": class_box(d, 780, 480, 300, "Favorito", ["userId", "productId", "creadoEn"], ["agregar()", "quitar()"], LIGHT_BLUE),
        "Fabricante": class_box(d, 1150, 480, 300, "Fabricante", ["id", "marca", "estado", "dni_hash"], ["crear()", "actualizar()"], LIGHT_BLUE),
        "Auditoria": class_box(d, 70, 810, 270, "Auditoria", ["accion", "entidad", "entidadId", "detalle"], ["registrarEvento()"], PURPLE),
        "IreHistorial": class_box(d, 410, 810, 300, "IreHistorial", ["score", "nivel", "version", "variables"], ["registrarPrediccion()"], PURPLE),
        "ModeloEstado": class_box(d, 780, 810, 300, "ModeloEstado", ["data_hash", "training_meta", "version"], ["actualizarModelo()"], PURPLE),
    }
    lines = [
        ("Usuario", "Pedido", "1", "0..*"),
        ("Usuario", "Favorito", "1", "0..*"),
        ("Producto", "Pedido", "1", "0..* ítems"),
        ("Producto", "VentaDiaria", "1", "0..*"),
        ("Producto", "ProductoCodigo", "1", "1"),
        ("Producto", "ProductoFinanza", "1", "0..1"),
        ("Producto", "Favorito", "1", "0..*"),
        ("Fabricante", "Producto", "1", "0..*"),
        ("Auditoria", "Usuario", "registra", ""),
        ("IreHistorial", "ModeloEstado", "usa", ""),
    ]
    centers = {k: ((v[0] + v[2]) // 2, (v[1] + v[3]) // 2) for k, v in boxes.items()}
    for a, b, la, lb in lines:
        d.line([centers[a], centers[b]], fill=LINE, width=2)
        mx, my = (centers[a][0] + centers[b][0]) // 2, (centers[a][1] + centers[b][1]) // 2
        d.rectangle([mx - 55, my - 14, mx + 55, my + 14], fill=WHITE)
        d.text((mx - 48, my - 10), f"{la} {lb}".strip(), fill=BLUE, font=F_XS)
    save(img, "02_diagrama_clases.png")
    write_mermaid(
        "02_diagrama_clases.mmd",
        """
classDiagram
  class Usuario { uid; email; rol; telefono; dni_hash }
  class Producto { id; nombre; categoria; precio; stock; tallaStock }
  class Pedido { id; userId; items; total; estado; metodoPago }
  class VentaDiaria { id; productoId; cantidad; precio; canal }
  class ProductoCodigo { productoId; codigo; actualizadoEn }
  class ProductoFinanza { productId; costo; margen; precioSugerido }
  class Favorito { userId; productId; creadoEn }
  class Fabricante { id; marca; estado; dni_hash }
  class Auditoria { accion; entidad; entidadId; detalle }
  class IreHistorial { score; nivel; version; variables }
  class ModeloEstado { data_hash; training_meta; version }
  Usuario "1" --> "0..*" Pedido
  Usuario "1" --> "0..*" Favorito
  Producto "1" --> "0..*" Pedido
  Producto "1" --> "0..*" VentaDiaria
  Producto "1" --> "1" ProductoCodigo
  Producto "1" --> "0..1" ProductoFinanza
  Producto "1" --> "0..*" Favorito
  Fabricante "1" --> "0..*" Producto
  IreHistorial --> ModeloEstado
""",
    )


def diagram_er():
    img, d = new_canvas("Modelo entidad-relación", "Base de datos Supabase: tablas, claves y relaciones críticas")
    boxes = {
        "usuarios": table_box(d, 60, 150, 300, "usuarios", ["PK uid", "email", "rol", "dni_hash", "telefono"], LIGHT_BLUE),
        "productos": table_box(d, 430, 150, 330, "productos", ["PK id", "nombre", "categoria", "precio", "stock", "activo"], LIGHT_GREEN),
        "pedidos": table_box(d, 830, 150, 330, "pedidos", ["PK id", "FK userId", "items jsonb", "total", "estado"], YELLOW),
        "ventasDiarias": table_box(d, 1230, 150, 330, "ventasDiarias", ["PK id", "FK productoId", "cantidad", "precio", "canal"], ORANGE),
        "productoCodigos": table_box(d, 430, 450, 330, "productoCodigos", ["PK productoId", "codigo UNIQUE", "actualizadoEn"], LIGHT_GREEN),
        "productoFinanzas": table_box(d, 60, 450, 300, "productoFinanzas", ["PK productId", "costo", "margen", "precioSugerido"], LIGHT_GREEN),
        "favoritos": table_box(d, 830, 450, 330, "favoritos", ["PK userId + productId", "FK userId", "FK productId"], LIGHT_BLUE),
        "fabricantes": table_box(d, 1230, 450, 330, "fabricantes", ["PK id", "marca", "estado", "dni_hash"], LIGHT_BLUE),
        "auditoria": table_box(d, 60, 760, 300, "auditoria", ["PK id", "accion", "entidad", "entidadId", "detalle"], PURPLE),
        "ireHistorial": table_box(d, 430, 760, 330, "ireHistorial", ["PK id", "score", "nivel", "version", "variables"], PURPLE),
        "modeloEstado": table_box(d, 830, 760, 330, "modeloEstado", ["PK id", "data_hash", "training_meta", "version"], PURPLE),
        "libro_reclamaciones": table_box(d, 1230, 760, 330, "libro_reclamaciones", ["PK id", "codigo", "estado", "cliente", "detalle"], PURPLE),
    }
    centers = {k: ((v[0] + v[2]) // 2, (v[1] + v[3]) // 2) for k, v in boxes.items()}
    rels = [
        ("usuarios", "pedidos", "1:N"),
        ("usuarios", "favoritos", "1:N"),
        ("productos", "favoritos", "1:N"),
        ("productos", "ventasDiarias", "1:N"),
        ("productos", "productoCodigos", "1:1"),
        ("productos", "productoFinanzas", "1:0..1"),
        ("fabricantes", "productos", "1:N"),
        ("modeloEstado", "ireHistorial", "1:N"),
    ]
    for a, b, label in rels:
        d.line([centers[a], centers[b]], fill=LINE, width=2)
        mx, my = (centers[a][0] + centers[b][0]) // 2, (centers[a][1] + centers[b][1]) // 2
        d.rectangle([mx - 38, my - 14, mx + 38, my + 14], fill=WHITE)
        centered_text(d, (mx - 36, my - 12, mx + 36, my + 12), label, fnt=F_XS)
    d.text((60, 1070), "Controles: RLS + FORCE RLS, REVOKE a anon/authenticated en tablas sensibles, BFF con service_role, auditoría con redacción PII.", fill=BLUE, font=F_S)
    save(img, "03_modelo_entidad_relacion.png")
    write_mermaid(
        "03_modelo_entidad_relacion.mmd",
        """
erDiagram
  usuarios ||--o{ pedidos : realiza
  usuarios ||--o{ favoritos : guarda
  productos ||--o{ favoritos : aparece
  productos ||--o{ ventasDiarias : vendido_en
  productos ||--|| productoCodigos : tiene
  productos ||--o| productoFinanzas : tiene
  fabricantes ||--o{ productos : provee
  modeloEstado ||--o{ ireHistorial : genera
""",
    )


def diagram_sequence_sale():
    sequence_diagram(
        "Diagrama de secuencia — registro de venta física",
        "Flujo atómico: trabajador/admin registra venta, BFF valida y Supabase descuenta stock",
        ["Trabajador", "AdminSales/StaffSales", "BFF Node", "Supabase RPC", "productos", "ventasDiarias", "Auditoría"],
        [
            (0, 1, "Selecciona producto/talla", "Cantidad y precio validado en UI"),
            (1, 2, "POST dailySales/register", "Firebase ID token"),
            (2, 2, "Verifica rol", "trabajador/admin"),
            (2, 3, "register_daily_sales_atomic", "Transacción controlada"),
            (3, 4, "Descontar stock", "por producto y talla"),
            (3, 5, "Insertar venta", "venta física registrada"),
            (2, 6, "Registrar evento", "detalle sin PII"),
            (2, 1, "200 OK", "venta confirmada"),
            (1, 0, "Toast + actualización", "stock y tabla refrescados"),
        ],
        "04_secuencia_registro_venta.png",
    )
    write_mermaid(
        "04_secuencia_registro_venta.mmd",
        """
sequenceDiagram
  actor Trabajador
  participant UI as AdminSales/StaffSales
  participant BFF as BFF Node
  participant RPC as Supabase RPC
  participant Productos as productos
  participant Ventas as ventasDiarias
  participant Aud as auditoria
  Trabajador->>UI: Selecciona producto, talla y cantidad
  UI->>BFF: POST /dailySales/register + Firebase token
  BFF->>BFF: Verificar rol trabajador/admin
  BFF->>RPC: register_daily_sales_atomic(p_sales)
  RPC->>Productos: Descontar stock por talla
  RPC->>Ventas: Insertar venta diaria
  BFF->>Aud: Registrar evento sin PII
  BFF-->>UI: 200 OK
  UI-->>Trabajador: Venta confirmada
""",
    )


def diagram_sequence_ai():
    sequence_diagram(
        "Diagrama de secuencia — predicción de riesgo empresarial",
        "Flujo IA: admin solicita predicción, servicio calcula IRE y devuelve métricas auditables",
        ["Administrador", "AdminPredictions", "BFF / Cliente IA", "Servicio IA", "Supabase", "Modelo ML", "Dashboard"],
        [
            (0, 1, "Solicita predicción", "Panel /admin/predicciones"),
            (1, 2, "Request con token", "Firebase ID token"),
            (2, 3, "POST predict/IRE", "features agregadas"),
            (3, 4, "Leer datos", "productos, ventas, finanzas, historial"),
            (3, 5, "Entrenar/evaluar", "data_hash, métricas, backtesting"),
            (5, 4, "Persistir IRE", "ireHistorial / modeloEstado"),
            (3, 2, "Respuesta IA", "score, nivel, variables, recomendaciones"),
            (2, 1, "Normalizar respuesta", "errores/cold start controlados"),
            (1, 6, "Renderizar dashboard", "IRE, sparkline y recomendaciones"),
        ],
        "05_secuencia_prediccion_riesgo.png",
    )
    write_mermaid(
        "05_secuencia_prediccion_riesgo.mmd",
        """
sequenceDiagram
  actor Administrador
  participant UI as AdminPredictions
  participant BFF as BFF / Cliente IA
  participant IA as Servicio IA
  participant DB as Supabase
  participant ML as Modelo ML
  participant Dash as Dashboard
  Administrador->>UI: Solicita predicción
  UI->>BFF: Request con token
  BFF->>IA: POST predict/IRE
  IA->>DB: Leer ventas, productos, finanzas e historial
  IA->>ML: Entrenar/evaluar con data_hash
  ML->>DB: Persistir IRE y métricas
  IA-->>BFF: score, nivel, variables y recomendaciones
  BFF-->>UI: Respuesta normalizada
  UI->>Dash: Mostrar IRE, tendencia y recomendaciones
""",
    )


def diagram_activity():
    img, d = new_canvas("Diagrama de actividades — compra web", "Flujo de negocio: catálogo, carrito, checkout, pago y pedido")
    steps = [
        (750, 150, 1050, 220, "Inicio"),
        (700, 280, 1100, 350, "Cliente explora catálogo y detalle"),
        (700, 410, 1100, 480, "Selecciona talla y agrega al carrito"),
        (700, 540, 1100, 610, "Ingresa dirección y método de pago"),
        (700, 670, 1100, 740, "BFF valida stock, precio, entrega e idempotencia"),
        (735, 800, 1065, 900, "¿Método Stripe?"),
        (300, 950, 650, 1030, "Crear sesión Stripe y esperar webhook"),
        (1140, 950, 1490, 1030, "Contraentrega: descontar stock al crear pedido"),
    ]
    for i, step in enumerate(steps):
        if i == 5:
            d.polygon([(900, step[1]), (step[2], 850), (900, step[3]), (step[0], 850)], fill=YELLOW, outline=BLUE)
            centered_text(d, (step[0] + 30, step[1] + 10, step[2] - 30, step[3] - 10), step[4], F_S)
        else:
            round_box(d, step[:4], fill=LIGHT_BLUE if i < 4 else LIGHT_GREEN, label=step[4], fnt=F_S)
    for a, b in zip(steps[:5], steps[1:6], strict=False):
        arrow(d, ((a[0]+a[2])//2, a[3]), ((b[0]+b[2])//2, b[1]), fill=BLUE)
    arrow(d, (760, 850), (520, 950), fill=BLUE, label="Sí")
    arrow(d, (1040, 850), (1320, 950), fill=BLUE, label="No")
    round_box(d, (690, 1050, 1110, 1125), fill=GREEN, label="Pedido registrado, auditable y visible en historial", fnt=F_S)
    arrow(d, (520, 1030), (730, 1050), fill=BLUE)
    arrow(d, (1320, 1030), (1070, 1050), fill=BLUE)
    save(img, "06_actividad_compra_web.png")
    write_mermaid(
        "06_actividad_compra_web.mmd",
        """
flowchart TD
  A([Inicio]) --> B[Cliente explora catálogo y detalle]
  B --> C[Selecciona talla y agrega al carrito]
  C --> D[Ingresa dirección y método de pago]
  D --> E[BFF valida stock, precio, entrega e idempotencia]
  E --> F{¿Método Stripe?}
  F -- Sí --> G[Crear sesión Stripe y esperar webhook]
  F -- No --> H[Contraentrega: descontar stock al crear pedido]
  G --> I([Pedido registrado, auditable y visible en historial])
  H --> I
""",
    )


def diagram_deployment():
    img, d = new_canvas("Diagrama de arquitectura y despliegue", "Dónde se ejecuta cada componente y cómo se comunican")
    nodes = [
        (80, 180, 420, 320, "Navegador cliente/admin\nReact SPA"),
        (520, 160, 900, 330, "Firebase Hosting\nassets web estáticos"),
        (1010, 160, 1390, 330, "BFF Render / Node\nAPI segura + service_role"),
        (1480, 160, 1740, 330, "Firebase Auth\nApp Check"),
        (1010, 430, 1390, 600, "Supabase PostgreSQL\nRLS, RPC, triggers"),
        (1480, 430, 1740, 600, "Stripe\nCheckout + webhook"),
        (1010, 700, 1390, 870, "Servicio IA\npredicción + IRE"),
        (1480, 700, 1740, 870, "Cloudinary / ORS / DNI\nproveedores externos"),
        (520, 760, 900, 930, "GitHub Actions + Docker\nCI, tests, ZAP, deploy"),
    ]
    for node in nodes:
        fill = LIGHT_BLUE if node[0] < 1000 else LIGHT_GREEN if "Supabase" in node[4] else YELLOW if "Stripe" in node[4] else PURPLE if "IA" in node[4] else GRAY
        round_box(d, node[:4], fill=fill, label=node[4], fnt=F_S)
    arrows = [
        ((420, 250), (520, 250), "GET web"),
        ((900, 250), (1010, 250), "API HTTPS"),
        ((1390, 240), (1480, 240), "token/auth"),
        ((1200, 330), (1200, 430), "SQL/RPC"),
        ((1390, 300), (1480, 500), "webhook/pago"),
        ((1200, 600), (1200, 700), "datos IA"),
        ((1390, 790), (1480, 790), "medios/rutas/DNI"),
        ((710, 760), (710, 330), "build/deploy"),
        ((900, 845), (1010, 790), "smoke/ZAP"),
    ]
    for s, e, label in arrows:
        arrow(d, s, e, fill=BLUE, label=label)
    d.text((85, 1010), "Controles: secretos fuera del frontend, Firebase token, App Check, Stripe signature, RLS, service_role en BFF, CI/CD, Docker, ZAP y stress.", fill=BLUE, font=F_S)
    save(img, "07_arquitectura_despliegue.png")
    write_mermaid(
        "07_arquitectura_despliegue.mmd",
        """
flowchart LR
  Browser[Navegador React SPA] --> Hosting[Firebase Hosting]
  Hosting --> BFF[BFF Render / Node]
  BFF --> Auth[Firebase Auth / App Check]
  BFF --> DB[(Supabase PostgreSQL)]
  BFF --> Stripe[Stripe Checkout + Webhook]
  BFF --> AI[Servicio IA]
  BFF --> Providers[Cloudinary / ORS / DNI]
  CI[GitHub Actions + Docker] --> Hosting
  CI --> BFF
  CI --> AI
  CI --> ZAP[OWASP ZAP + stress]
""",
    )


def generate_all_diagrams():
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    MERMAID_DIR.mkdir(parents=True, exist_ok=True)
    diagram_use_case()
    diagram_class()
    diagram_er()
    diagram_sequence_sale()
    diagram_sequence_ai()
    diagram_activity()
    diagram_deployment()


def configure_doc(document: Document):
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Arial"
        styles[name].font.color.rgb = RGBColor(31, 78, 121)


def add_title(document: Document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DOCUMENTO 4 — DIAGRAMAS UML Y MODELADO")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(31, 78, 121)
    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("Sistema web Calzatura Vilchez — diseño de casos de uso, clases, datos, secuencias, actividades y despliegue").italic = True
    p3 = document.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run("Fecha: 30/05/2026").bold = True


def add_doc_table(document: Document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell._tc.get_or_add_tcPr().append(OxmlElement("w:tcPr"))
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(31, 78, 121)
                run.font.size = Pt(8)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    document.add_paragraph()


def build_document():
    doc = Document()
    configure_doc(doc)
    add_title(doc)
    doc.add_heading("1. Criterio de selección", level=1)
    doc.add_paragraph(
        "Se presentan diagramas suficientes para sustentar el diseño del sistema web antes y durante la implementación: "
        "actores, entidades, base de datos, procesos críticos, flujo de compra, IA y despliegue. "
        "El formato es conciso, pero cada diagrama contiene componentes reales del proyecto y evidencia trazable."
    )
    add_doc_table(
        doc,
        ["Diagrama", "Para qué sirve", "Archivo generado"],
        [
            ["Casos de uso general", "Muestra actores y funciones principales del sistema.", "01_casos_uso_general.png / .mmd"],
            ["Clases", "Muestra entidades principales y relaciones del dominio.", "02_diagrama_clases.png / .mmd"],
            ["Entidad-relación", "Muestra base de datos Supabase y relaciones críticas.", "03_modelo_entidad_relacion.png / .mmd"],
            ["Secuencia: registro de venta", "Muestra paso a paso la venta física con RPC atómica.", "04_secuencia_registro_venta.png / .mmd"],
            ["Secuencia: predicción de riesgo", "Muestra interacción con IA, datos, modelo e IRE.", "05_secuencia_prediccion_riesgo.png / .mmd"],
            ["Actividad: compra web", "Muestra flujo de negocio desde catálogo hasta pedido.", "06_actividad_compra_web.png / .mmd"],
            ["Arquitectura/despliegue", "Muestra dónde se ejecuta cada componente.", "07_arquitectura_despliegue.png / .mmd"],
        ],
    )
    diagrams = [
        ("2. Diagrama de casos de uso general", "01_casos_uso_general.png"),
        ("3. Diagrama de clases", "02_diagrama_clases.png"),
        ("4. Modelo entidad-relación", "03_modelo_entidad_relacion.png"),
        ("5. Diagrama de secuencia: registro de venta", "04_secuencia_registro_venta.png"),
        ("6. Diagrama de secuencia: predicción de riesgo empresarial", "05_secuencia_prediccion_riesgo.png"),
        ("7. Diagrama de actividades: compra web", "06_actividad_compra_web.png"),
        ("8. Diagrama de arquitectura y despliegue", "07_arquitectura_despliegue.png"),
    ]
    for heading, image_name in diagrams:
        doc.add_heading(heading, level=1)
        doc.add_picture(str(IMG_DIR / image_name), width=Inches(7.3))
        doc.add_paragraph()
    doc.add_heading("9. Evidencia editable", level=1)
    doc.add_paragraph(
        f"Los archivos Mermaid editables se encuentran en: {MERMAID_DIR}. "
        "Estos archivos permiten regenerar o ajustar los diagramas en herramientas compatibles sin rehacer el modelado."
    )
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Diagramas UML y modelado — Calzatura Vilchez — 30/05/2026")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)
    DOC_FILE.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOC_FILE)
    shutil.copy2(DOC_FILE, DOWNLOAD_DOC)


def main():
    generate_all_diagrams()
    build_document()
    print(DOC_FILE)
    print(DOWNLOAD_DOC)
    print(IMG_DIR)
    print(MERMAID_DIR)


if __name__ == "__main__":
    main()
