# -*- coding: utf-8 -*-
"""Genera Auditoria-Panel-Administrativo-Calzatura-Vilchez.docx en docs/. Ejecutar: python scripts/generate_admin_audit_docx.py"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_matrix(doc: Document, title: str, rows: list[tuple[str, str]]) -> None:
    doc.add_paragraph(title, style="Heading 4")
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (rol, texto) in enumerate(rows):
        table.rows[i].cells[0].text = rol
        table.rows[i].cells[1].text = texto
    doc.add_paragraph()


def p(doc: Document, text: str, bold: bool = False) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    if bold:
        run.bold = True


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out = root / "docs" / "Auditoria-Panel-Administrativo-Calzatura-Vilchez.docx"
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    t = doc.add_heading("Auditoría del panel administrativo", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Calzatura Vilchez — Aplicación web (React / Vite). "
        "Área auditada: rutas bajo /admin y componentes asociados."
    )
    doc.add_paragraph(
        "Enfoques considerados en cada sección: ISO (calidad de software y gestión de datos), "
        "Gerente (operación y riesgo), Cliente final (impacto indirecto en tienda), "
        "Testeador (pruebas y regresiones), Diseñador UX/UI, Publicidad/Comercial."
    )
    doc.add_paragraph(f"Documento generado automáticamente. Ruta de salida: {out}")

    # --- Marco 0
    doc.add_heading("Marco 0 — Rutas, área administradores y control de acceso", level=1)
    doc.add_heading("Alcance en código", level=2)
    p(
        doc,
        "Rutas (App.tsx): /admin con AreaRoute area=\"administradores\" y AdminLayout; hijos: "
        "Dashboard (índice), productos, pedidos, ventas, usuarios, fabricantes, predicciones, datos (paths.ts).",
    )
    p(
        doc,
        "RouteGuards.tsx: AreaRoute usa canAccessArea(area, userProfile?.rol, user?.email). "
        "accessControl.ts: administradores solo para rol admin; isSuperAdminEmail(email) abre el área.",
    )
    p(
        doc,
        "AdminLayout.tsx: segunda guarda — sin usuario → login con redirect; con usuario pero no isAdmin → inicio. "
        "Sidebar con NavLink, colapso persistido en localStorage, tema claro/oscuro, ver tienda y cerrar sesión.",
    )
    doc.add_heading("Fortalezas", level=2)
    for x in [
        "Doble comprobación (ruta + layout) frente a acceso por URL.",
        "Carga diferida (lazy) del layout y páginas admin.",
        "Navegación lateral clara con estado activo y persistencia de colapso.",
        "Salida explícita a tienda y logout con feedback (toast).",
    ]:
        doc.add_paragraph(x, style="List Bullet")
    doc.add_heading("Riesgos y observaciones", level=2)
    for x in [
        "La seguridad real depende de Supabase RLS, Functions y políticas; el panel es capa de presentación.",
        "Superadmin por lista de emails en cliente: documentar como convenio de producto, no como único control ISO.",
        "Comentario en App.tsx aún menciona Firestore como control; desalineado con stack actual en parte del dominio.",
        "Sin skip link \"ir al contenido\" típico en layouts admin para accesibilidad.",
    ]:
        doc.add_paragraph(x, style="List Bullet")
    add_matrix(
        doc,
        "Matriz por rol — Marco 0",
        [
            ("ISO", "Documentar separación UI vs autorización en servidor; alinear comentarios con arquitectura."),
            ("Gerente", "Claro quién puede operar el panel; delegación de cuentas admin."),
            ("Cliente", "No ve el admin; riesgo indirecto si acceso admin se filtra."),
            ("Testeador", "Casos: cliente en /admin, admin, superadmin, sesión expirada."),
            ("Diseñador", "Revisar contraste y área clic del sidebar colapsado."),
            ("Publicidad", "Poco directo salvo coordinación de quién actualiza campañas en catálogo."),
        ],
    )

    # --- Dashboard
    doc.add_heading("Sección 1 — Dashboard (AdminDashboard.tsx)", level=1)
    doc.add_heading("Alcance", level=2)
    p(
        doc,
        "KPIs (productos, pedidos, ingresos pedidos, usuarios, ventas hoy, ganancia estimada hoy, pendientes), "
        "gráfico últimos 7 días (ventas manuales + pedidos completados), pedidos recientes, resumen de estados, "
        "accesos rápidos a módulos, tabla de auditoría reciente (fetchRecentAudit), modal de detalle de pedido.",
    )
    doc.add_heading("Fuentes de datos y fallos", level=2)
    p(
        doc,
        "Promise.all de cinco fetch sin .catch: si cualquiera rechaza, el then no corre pero finally pone loading=false; "
        "stats permanecen en ceros sin toast — degradación poco comunicada (ISO 25010 usabilidad/fiabilidad). "
        "fetchRecentAudit tiene .catch silencioso (auditoría vacía sin explicación).",
    )
    doc.add_heading("Fortalezas", level=2)
    for x in [
        "Visión 360° en una pantalla con enlaces a profundidad.",
        "Bloque de trazabilidad etiquetado como ISO 9001.",
        "Combinación ventas manuales + pedidos para \"hoy\" y gráfico.",
    ]:
        doc.add_paragraph(x, style="List Bullet")
    doc.add_heading("Riesgos", level=2)
    for x in [
        "Escalabilidad: carga todos los productos, pedidos, ventas, finanzas y usuarios — posible lentitud con volumen alto.",
        "Ganancia estimada: fallback de costo al precio de venta si falta finanza puede distorsionar margen.",
        "Zona horaria del navegador vs servidor para \"hoy\".",
        "Filas de tabla clicables sin rol de botón explícito — teclado/lector.",
    ]:
        doc.add_paragraph(x, style="List Bullet")
    add_matrix(
        doc,
        "Matriz por rol — Dashboard",
        [
            ("ISO", "Tratar KPIs como indicativos; mejorar manejo explícito de error en carga."),
            ("Gerente", "Reunión diaria rápida; riesgo de decisión con dashboard vacío por fallo de red."),
            ("Cliente", "Indirecto: mal lectura de pendientes → mala comunicación de plazos."),
            ("Testeador", "Mock 500 en un fetch y verificar UX y mensajes."),
            ("Diseñador", "Jerarquía visual y diferenciar skeleton de \"sin datos\"."),
            ("Publicidad", "KPI hoy útil; sin vínculo campaña→SKU en esta vista."),
        ],
    )

    # --- Productos
    doc.add_heading("Sección 2 — Productos (AdminProducts + services + migraciones)", level=1)
    doc.add_heading("Fortalezas", level=2)
    for x in [
        "Variantes por color (hasta 5), familia, código, finanzas, stock por talla acoplado a categoría.",
        "RPC create_product_variants_atomic y update_product_atomic; triggers/CHECK comerciales; índice único de códigos.",
        "Imágenes: imageRules (dimensiones, ratio, peso post-compresión); URL con validateImageUrlDimensions; slot vacío si falla.",
        "Mensajes de permiso: 42501 y row-level security. E2E: admin-code-guards, admin-commercial-guards, admin-stock-tallas.",
    ]:
        doc.add_paragraph(x, style="List Bullet")
    doc.add_heading("Riesgos", level=2)
    for x in [
        "Componente muy grande (mantenibilidad ISO 25010).",
        "Campo/etiqueta de campaña explícita recomendable si marketing necesita attribution.",
        "E2E no cubren todas las combinaciones de filtros/errores.",
    ]:
        doc.add_paragraph(x, style="List Bullet")
    add_matrix(
        doc,
        "Matriz por rol — Productos",
        [
            ("ISO", "Referencia de defensa en profundidad (BD + pruebas) para el resto del admin."),
            ("Gerente", "Menos riesgo de catálogo inconsistente."),
            ("Cliente", "Impacto directo en tienda."),
            ("Testeador", "Mayor cobertura E2E relativa aquí."),
            ("Diseñador", "Modal complejo; revisar scroll, foco, móvil."),
            ("Publicidad", "Destacado/descuento; considerar metadato de campaña."),
        ],
    )

    # --- Pedidos
    doc.add_heading("Sección 3 — Pedidos (AdminOrders + orders.ts)", level=1)
    p(
        doc,
        "fetchAllOrders desde Supabase; updateOrderStatus actualiza solo estado con logAudit. "
        "UI: filtro, tarjetas expandibles, select de estado sin confirmación, detalle con dirección y Stripe session id.",
    )
    doc.add_heading("Riesgos destacados", level=2)
    for x in [
        "useEffect sin .catch si fetchAllOrders falla — promesa rechazada no manejada.",
        "Cualquier transición de estado permitida en UI si BD lo permite — riesgo operativo y de servicio al cliente.",
        "Toast genérico en error de actualización de estado.",
        "Sin E2E dedicado en carpeta e2e para esta pantalla.",
    ]:
        doc.add_paragraph(x, style="List Bullet")
    add_matrix(
        doc,
        "Matriz por rol — Pedidos",
        [
            ("ISO", "Definir transiciones válidas en BD o RPC; UI acorde."),
            ("Gerente", "Reclamos y logística vs estado mostrado."),
            ("Cliente", "Alto impacto en confianza."),
            ("Testeador", "Matriz transiciones × API; race en select."),
            ("Diseñador", "Touch targets; confirmación para cancelado/entregado."),
            ("Publicidad", "Alinear estados con comunicaciones externas."),
        ],
    )

    # --- Ventas
    doc.add_heading("Sección 4 — Ventas diarias (AdminSales)", level=1)
    p(
        doc,
        "Flujo de mostrador: líneas, DNI lookup, documentos, impresión; addDailySale; devolución markSaleReturned "
        "con restauración de stock vía updateProduct (parcial), distinto del patrón RPC de edición de catálogo.",
    )
    doc.add_heading("Riesgos", level=2)
    for x in [
        "Archivo muy extenso (~1000 líneas) — mantenibilidad.",
        "Posible competencia de escrituras de stock sin control de versión.",
        "PII (DNI). Sin E2E en e2e/.",
    ]:
        doc.add_paragraph(x, style="List Bullet")
    add_matrix(
        doc,
        "Matriz por rol — Ventas",
        [
            ("ISO", "Auditoría de ventas/devoluciones y retención de datos personales."),
            ("Gerente", "Caja e inventario físico vs online."),
            ("Cliente", "Stock mal devuelto → overselling web."),
            ("Testeador", "Prioridad alta por dinero y stock."),
            ("Diseñador", "Considerar wizard por pasos."),
            ("Publicidad", "Coherencia precio tienda vs campaña web."),
        ],
    )

    # --- Usuarios
    doc.add_heading("Sección 5 — Usuarios (AdminUsers)", level=1)
    p(
        doc,
        "Carga fetchAllUsers + fetchAllOrders para contar pedidos por usuario. Cambio de rol con updateUserRole; "
        "solo superadmin por email gestiona roles admin.",
    )
    doc.add_heading("Riesgos", level=2)
    for x in [
        "Complejidad O(usuarios × pedidos) al filtrar si el volumen crece.",
        "Sin E2E específico.",
    ]:
        doc.add_paragraph(x, style="List Bullet")
    add_matrix(
        doc,
        "Matriz por rol — Usuarios",
        [
            ("ISO", "RLS en tabla usuarios; superadmin no solo en frontend."),
            ("Gerente", "Gobernanza de cuentas admin."),
            ("Cliente", "Rol incorrecto afecta acceso y datos."),
            ("Testeador", "Matriz de roles y RLS."),
            ("Diseñador", "Paginación/virtualización si crece la base."),
            ("Publicidad", "Indirecto."),
        ],
    )

    # --- Fabricantes
    doc.add_heading("Sección 6 — Fabricantes (AdminManufacturers)", level=1)
    for x in [
        "Fortalezas: CRUD, documentos (límite 8), DNI, imágenes Cloudinary, filtros.",
        "Riesgos: archivo grande, PII, sin E2E.",
    ]:
        doc.add_paragraph(x)
    add_matrix(
        doc,
        "Matriz por rol — Fabricantes",
        [
            ("ISO", "Datos de proveedor y documentación."),
            ("Gerente", "Cadena de suministro."),
            ("Cliente", "Confianza indirecta."),
            ("Testeador", "Smoke / E2E recomendable."),
            ("Diseñador", "Formularios largos en móvil."),
            ("Publicidad", "Storytelling de origen si aplica."),
        ],
    )

    # --- Predicciones
    doc.add_heading("Sección 7 — Predicciones IA (AdminPredictions)", level=1)
    for x in [
        "Fortalezas: métricas ricas, servicio configurable VITE_AI_SERVICE_URL, UI avanzada.",
        "Riesgos críticos: archivo muy grande (~2600+ líneas); VITE_AI_SERVICE_BEARER_TOKEN expuesto en bundle si se usa — "
        "preferible BFF o función con secreto. Dependencia de servicio externo. Sin E2E.",
    ]:
        doc.add_paragraph(x)
    add_matrix(
        doc,
        "Matriz por rol — Predicciones IA",
        [
            ("ISO 27001", "Token y datos enviados al modelo; retención de logs en servicio Python."),
            ("Gerente", "Decisiones de compra con revisión humana documentada."),
            ("Cliente", "Ruptura/sobrestock por mala predicción."),
            ("Testeador", "Contratos API mock, timeouts."),
            ("Diseñador", "Simplificar por persona de uso."),
            ("Publicidad", "Vincular con etiquetas de campaña en catálogo."),
        ],
    )

    # --- Datos Excel
    doc.add_heading("Sección 8 — Datos Excel (AdminData)", level=1)
    for x in [
        "Fortalezas: plantillas xlsx, importValidate/importTransform, export, ejemplo ID_PRODUCTO sin sufijo Firestore.",
        "Riesgos: import masivo puede saltar parte de la UX si BD no cubre todas las combinaciones; acciones destructivas; "
        "posible uso de mismas variables de servicio IA que en predicciones.",
    ]:
        doc.add_paragraph(x)
    add_matrix(
        doc,
        "Matriz por rol — Datos Excel",
        [
            ("ISO", "Procedimiento de import, staging, auditoría."),
            ("Gerente", "Muy potente y muy peligroso; segregación de funciones."),
            ("Cliente", "Catálogo roto visible de inmediato."),
            ("Testeador", "Archivos inválidos, columnas faltantes, UTF-8."),
            ("Diseñador", "Confirmación en dos pasos para destructivos."),
            ("Publicidad", "Datos de campaña mal importados arruinan attribution."),
        ],
    )

    # --- Cierre
    doc.add_heading("Checklist transversal de cierre", level=1)
    for x in [
        "¿Cada acción sensible tiene equivalente en RLS o función en servidor?",
        "¿Hay .catch / toast en todas las cargas Promise.all del admin (revisar Dashboard y Pedidos)?",
        "¿E2E mínimo por módulo con dinero o stock (pedidos, ventas, datos)?",
        "¿Token de servicios IA solo en backend?",
        "¿Documentación y comentarios alineados con Supabase + Functions vs Firestore legacy?",
    ]:
        doc.add_paragraph(x, style="List Number")

    doc.add_paragraph()
    p(doc, "Fin del documento.", bold=True)

    doc.save(out)
    print(f"OK: {out}")


if __name__ == "__main__":
    main()
