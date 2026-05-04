"""
Genera AdminPanel-auditoria-total.docx desde el markdown homónimo.
Ejecutar: python scripts/build_admin_panel_audit_docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "15-modulos" / "AdminPanel-auditoria-total.md"
OUT_PATH = ROOT / "docs" / "15-modulos" / "AdminPanel-auditoria-total.docx"


def strip_md(s: str) -> str:
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"`([^`]+)`", r"\1", s)
    s = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", s)
    return s


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, rows: list[list[str]], header: bool = True) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.rows[i].cells[j]
            text = row[j] if j < len(row) else ""
            cell.text = text
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            if header and i == 0:
                set_cell_shading(cell, "D9E2F3")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
    doc.add_paragraph()


def parse_md_table(lines: list[str]) -> list[list[str]] | None:
    if not lines or not lines[0].strip().startswith("|"):
        return None
    rows = []
    for line in lines:
        s = line.strip()
        if not s.startswith("|"):
            break
        if re.match(r"^\|[\s\-:|]+\|$", s):
            continue
        cells = [strip_md(c.strip()) for c in s.strip("|").split("|")]
        rows.append(cells)
    return rows if rows else None


def main() -> None:
    text = MD_PATH.read_text(encoding="utf-8")
    doc = Document()
    core = doc.core_properties
    core.title = "Auditoría total — Panel administrativo"
    core.subject = "Calzatura Vilchez — /admin"
    core.keywords = "auditoría, admin, ISO, E2E"

    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("# "):
            doc.add_heading(strip_md(line[2:].strip()), level=0)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(strip_md(line[3:].strip()), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(strip_md(line[4:].strip()), level=2)
            i += 1
            continue
        if line.strip() == "---":
            p = doc.add_paragraph()
            p.add_run("—" * 40).italic = True
            i += 1
            continue
        if line.strip().startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            rows = parse_md_table(block)
            if rows:
                add_table(doc, rows, header=True)
            continue
        if line.strip():
            p = doc.add_paragraph()
            p.add_run(strip_md(line.strip()))
            p.paragraph_format.space_after = Pt(3)
        i += 1

    doc.add_paragraph()
    foot = doc.add_paragraph(
        "Documento generado desde AdminPanel-auditoria-total.md — "
        "Regenerar con: python scripts/build_admin_panel_audit_docx.py"
    )
    foot.italic = True
    foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in foot.runs:
        run.font.size = Pt(9)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(OUT_PATH))
        print(f"Escrito: {OUT_PATH}")
    except PermissionError:
        alt = OUT_PATH.with_name(f"{OUT_PATH.stem}-generado{OUT_PATH.suffix}")
        doc.save(str(alt))
        print(f"El archivo original esta abierto en otro programa. Escrito: {alt}")


if __name__ == "__main__":
    main()
